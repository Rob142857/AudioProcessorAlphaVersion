import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="webrtcvad")

import sys
import os
import shutil
import subprocess
import tempfile
import argparse
import whisper
import torch
from docx import Document
from moviepy import VideoFileClip, AudioFileClip
import imageio_ffmpeg as iio_ffmpeg
from deepmultilingualpunctuation import PunctuationModel
try:
    import webrtcvad
    WEBRTCVAD_AVAILABLE = True
except ImportError:
    WEBRTCVAD_AVAILABLE = False
    print("⚠️  webrtcvad not available - Voice Activity Detection disabled")
import re
import time
import json

try:
    import torch_directml
except Exception:
    torch_directml = None


def format_duration(seconds):
    """Convert seconds to a readable time format (HH:MM:SS or MM:SS)."""
    if seconds is None:
        return None
    
    seconds = round(seconds, 2)
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:05.2f}"
    else:
        return f"{minutes:02d}:{secs:05.2f}"


def split_into_paragraphs(text, max_length=500):
    """Richer paragraphing heuristics:
    - Treat blank lines as hard paragraph breaks.
    - Within a block, split into sentences (., !, ?) while respecting common abbreviations and initials.
    - Group 2–5 sentences or until ~800 chars per paragraph.
    - Preserve bullet/numbered lines as their own paragraphs.
    """
    if not text:
        return []

    # Hard breaks by blank lines
    blocks = re.split(r"\n\s*\n+", text)

    # Common abbreviations to avoid splitting on
    abbreviations = {
        "Mr.", "Mrs.", "Ms.", "Dr.", "Prof.", "Sr.", "Jr.", "St.",
        "vs.", "etc.", "e.g.", "i.e.", "cf.", "Co.", "Corp.", "Inc.", "Ltd.",
        "U.S.", "U.K.", "No.", "Mt.", "Rd.", "Ave.", "Jan.", "Feb.", "Mar.",
        "Apr.", "Aug.", "Sept.", "Oct.", "Nov.", "Dec."
    }

    bullet_re = re.compile(r"^\s*(?:[-*•]\s+|\d+[\.)]\s+)")
    # Sentence split (rough): break after ., !, ? followed by whitespace and a capital/quote/open paren
    sentence_split_re = re.compile(r"(?<=[.!?])[\)]?\"?'?\s+(?=[\"'\(A-Z0-9])")

    def is_abbrev_end(s: str) -> bool:
        s = s.strip()
        if not s:
            return False
        last = s.split()[-1]
        if last in abbreviations:
            return True
        # Single-letter initial (e.g., "J.")
        if re.match(r"^[A-Z]\.$", last):
            return True
        # Ellipses
        if s.endswith("..."):
            return True
        return False

    paragraphs = []

    for block in blocks:
        if not block.strip():
            continue

        # If block looks like bullets or numbered list, keep each line as a paragraph
        lines = block.splitlines()
        if any(bullet_re.match(ln) for ln in lines):
            for ln in lines:
                ln = ln.strip()
                if ln:
                    paragraphs.append(ln)
            continue

        # Collapse internal newlines to single spaces for sentence processing
        block_text = re.sub(r"\s+", " ", block.strip())
        # Initial naive split
        parts = sentence_split_re.split(block_text)
        # Merge sentences that were split after abbreviations or initials
        sentences = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if sentences and (is_abbrev_end(sentences[-1]) or len(sentences[-1]) <= 2):
                sentences[-1] = (sentences[-1] + " " + part).strip()
            else:
                sentences.append(part)

        # Group into paragraphs by sentence count and char budget
        cur = []
        cur_len = 0
        max_chars_per_para = 800
        max_sentences_per_para = 5

        for s in sentences:
            # Start new para if adding would exceed limits
            if cur and (len(s) + cur_len > max_chars_per_para or len(cur) >= max_sentences_per_para):
                paragraphs.append(" ".join(cur).strip())
                cur = []
                cur_len = 0
            cur.append(s)
            cur_len += len(s) + 1
        if cur:
            paragraphs.append(" ".join(cur).strip())

    return paragraphs


def convert_to_mp3(input_path, bitrate="192k"):
    ext = os.path.splitext(input_path)[1].lower()
    mp3_path = os.path.splitext(input_path)[0] + "_converted.mp3"
    if ext in [".mp4", ".mov", ".avi", ".mkv"]:
        clip = VideoFileClip(input_path)
        try:
            if clip.audio is None:
                raise ValueError("No audio track found in the video file")
            clip.audio.write_audiofile(mp3_path, codec="libmp3lame", bitrate=bitrate)
        finally:
            clip.close()
        return mp3_path
    elif ext in [".mp3", ".wav", ".flac", ".ogg", ".m4a"]:
        clip = AudioFileClip(input_path)
        try:
            clip.write_audiofile(mp3_path, codec="libmp3lame", bitrate=bitrate)
        finally:
            clip.close()
        return mp3_path
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def ensure_ffmpeg_in_path():
    if shutil.which("ffmpeg"):
        return
    ffexe = None
    if hasattr(iio_ffmpeg, "get_ffmpeg_exe"):
        try:
            ffexe = iio_ffmpeg.get_ffmpeg_exe()
        except Exception:
            ffexe = None
    if not ffexe and hasattr(iio_ffmpeg, "get_exe"):
        try:
            ffexe = iio_ffmpeg.get_exe()
        except Exception:
            ffexe = None
    if ffexe and os.path.isfile(ffexe):
        # Create a small wrapper batch file named ffmpeg.bat in the script directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        bat_path = os.path.join(script_dir, "ffmpeg.bat")
        try:
            if not os.path.exists(bat_path):
                with open(bat_path, "w", encoding="utf-8") as bat:
                    # Use quotes around ffexe and forward all args
                    bat.write(f'@"{ffexe}" %*\n')
            # Prepend script dir to PATH so 'ffmpeg' resolves to our batch wrapper
            os.environ["PATH"] = script_dir + os.pathsep + os.environ.get("PATH", "")
            return
        except Exception:
            # Fall back to adding the ffmpeg folder directly
            ffdir = os.path.dirname(ffexe)
            os.environ["PATH"] = ffdir + os.pathsep + os.environ.get("PATH", "")
            return
    raise RuntimeError("ffmpeg not found. Install ffmpeg or ensure imageio_ffmpeg is available.")


def preprocess_audio(input_path, out_path, bitrate="128k"):
    """Run ffmpeg with simple filters to improve audio quality before transcription.
    Applies highpass, lowpass, spectral denoise and normalization, and resamples to 16k mono.
    """
    ff = shutil.which("ffmpeg") or None
    if not ff:
        # Try imageio_ffmpeg
        ffexe = None
        if hasattr(iio_ffmpeg, "get_ffmpeg_exe"):
            try:
                ffexe = iio_ffmpeg.get_ffmpeg_exe()
            except Exception:
                ffexe = None
        if not ffexe and hasattr(iio_ffmpeg, "get_exe"):
            try:
                ffexe = iio_ffmpeg.get_exe()
            except Exception:
                ffexe = None
        ff = ffexe
    if not ff:
        raise RuntimeError("ffmpeg not found for preprocessing")

    # Filter chain: remove low rumble, limit highs, denoise, normalize
    filters = "highpass=f=80, lowpass=f=8000, afftdn, dynaudnorm=g=5"
    cmd = [
        ff,
        "-y",
        "-i",
        input_path,
        "-ac",
        "1",
        "-ar",
        "16000",
        "-af",
        filters,
        "-codec:a",
        "libmp3lame",
        "-b:a",
        bitrate,
        out_path,
    ]
    subprocess.run(cmd, check=True)


def choose_device(preferred: str = "auto"):
    """Return a torch device string based on preferred choice and availability.
    preferred: 'auto'|'cpu'|'cuda'|'dml'
    Behavior: prefer CPU + CUDA combination for best performance and compatibility.
    """
    preferred = (preferred or "auto").lower()
    if preferred == "auto":
        # Prefer CUDA for x64 Windows systems with NVIDIA GPUs, with CPU fallback
        if torch.cuda.is_available():
            return "cuda"
        # CPU fallback - always available
        return "cpu"

    if preferred in ("dml", "directml"):
        if torch_directml is not None:
            try:
                _ = torch_directml.device()
                return "dml"
            except Exception:
                print("torch-directml importable but initialization failed; falling back to CPU")
                return "cpu"
        print("torch-directml not installed or not importable; falling back to CPU")
        return "cpu"

    if preferred == "cuda" and torch.cuda.is_available():
        return "cuda"
    return "cpu"


def get_media_duration(path):
    """Return media duration in seconds. Try moviepy then ffprobe fallback."""
    try:
        ext = os.path.splitext(path)[1].lower()
        if ext in [".mp4", ".mov", ".mkv", ".avi"]:
            clip = VideoFileClip(path)
            try:
                return float(clip.duration)
            finally:
                clip.close()
        else:
            clip = AudioFileClip(path)
            try:
                return float(clip.duration)
            finally:
                clip.close()
    except Exception:
        # fallback to ffprobe
        ff = shutil.which("ffprobe") or shutil.which("ffmpeg")
        if not ff:
            return None
        # use ffprobe if available
        try:
            cmd = [ff, "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", path]
            out = subprocess.check_output(cmd, stderr=subprocess.DEVNULL)
            return float(out.decode().strip())
        except Exception:
            return None


def transcribe_file(input_path, model_name="large", preprocess=True, keep_temp=False, bitrate="192k", device_preference="auto", vad=True, punctuate=True, output_dir=None):
    """
    Transcribe a single audio/video file with optimized settings for best quality.
    
    Args:
        input_path: Path to input audio/video file
        model_name: Whisper model ("large" or "medium")
        preprocess: Enable preprocessing (always True for quality)
        keep_temp: Keep temporary files
        bitrate: Audio bitrate for preprocessing
        device_preference: Device preference ("auto", "cpu", "cuda", "dml")
        vad: Enable VAD segmentation (always True for quality)  
        punctuate: Enable punctuation restoration (always True for quality)
        output_dir: Output directory (defaults to Downloads)
    """
    ensure_ffmpeg_in_path()

    start_time = time.time()

    temp_files = []
    log_outputs = []
    try:
        ext = os.path.splitext(input_path)[1].lower()
        # If user requested preprocessing, always create a temp file with processed audio
        if preprocess:
            tf = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
            tf.close()
            pre_path = tf.name
            print(f"Preprocessing audio -> {pre_path}...")
            preprocess_audio(input_path, pre_path, bitrate=bitrate)
            audio_path = pre_path
            temp_files.append(pre_path)
        else:
            # Convert non-mp3 inputs to mp3
            if ext != ".mp3":
                print(f"Converting {input_path} to MP3 (compressed)...")
                audio_path = convert_to_mp3(input_path, bitrate=bitrate)
                temp_files.append(audio_path)
            else:
                audio_path = input_path

        media_length = get_media_duration(input_path)

        print(f"Loading Whisper model '{model_name}' (this may take a while)...")
        # Select device
        device = choose_device(device_preference)
        print(f"Selected device: {device}")
        print(f"CUDA available: {torch.cuda.is_available()}")
        if torch.cuda.is_available():
            print(f"GPU: {torch.cuda.get_device_name(0)}")
        
        # Whisper's load_model will place parameters on the detected device; load then ensure placement
        model = whisper.load_model(model_name, device=device)
        print(f"Model loaded on device: {next(model.parameters()).device}")
        
        # Verify the correct model was loaded by checking model size
        model_size = sum(p.numel() for p in model.parameters())
        print(f"Model size: {model_size:,} parameters")
        
        # Estimate model type based on size (approximate)
        if model_size > 700000000:  # Large model has ~1.5B parameters
            loaded_model_type = "large"
        elif model_size > 350000000:  # Large model has ~770M parameters (when cached medium loads as large)
            loaded_model_type = "large"
        elif model_size > 70000000:  # Small model has ~244M parameters
            loaded_model_type = "small"
        elif model_size > 35000000:  # Base model has ~74M parameters
            loaded_model_type = "base"
        else:
            loaded_model_type = "tiny"
        
        print(f"Detected loaded model type: {loaded_model_type}")
        
        if loaded_model_type != model_name:
            print(f"⚠️  WARNING: Requested model '{model_name}' but loaded '{loaded_model_type}' (cached model used)")
            print("   To force loading the correct model, delete the Whisper cache and try again")
        else:
            print(f"✓ Correct model '{model_name}' loaded successfully")
        
        # Verify model is on correct device
        if device == "cuda" and torch.cuda.is_available():
            if next(model.parameters()).device.type != "cuda":
                print(f"WARNING: Model not on CUDA! Attempting to move...")
                model = model.to("cuda")
                print(f"Model moved to: {next(model.parameters()).device}")

        outputs = []
        segment_files_used = []
        if vad:
            print("Running VAD segmentation...")
            segments = vad_segment_times(audio_path, aggressiveness=1, frame_duration_ms=30, padding_ms=300)
            print(f"Raw segments found: {len(segments)}")
            segments = post_process_segments(segments, min_duration=0.3, merge_gap=0.5, max_segments=1000)
            print(f"Segments after post-processing: {len(segments)}")
            seg_files = []
            for idx, (s, e) in enumerate(segments):
                seg_tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
                seg_tmp.close()
                print(f"Extracting segment {idx+1}/{len(segments)}: {s:.2f}-{e:.2f}s")
                extract_segment(audio_path, s, e, seg_tmp.name, bitrate=bitrate)
                seg_files.append(seg_tmp.name)
                temp_files.append(seg_tmp.name)
            # transcribe each segment with optimized settings
            for seg in seg_files:
                print(f"Transcribing segment {seg}")
                # Optimized settings for best transcription quality
                res = model.transcribe(seg, language=None, 
                                     compression_ratio_threshold=float('inf'),  # Disable repetitive content detection
                                     logprob_threshold=-1.0,                    # Disable low-confidence filtering
                                     no_speech_threshold=0.1,                   # Lower threshold for "no speech"
                                     condition_on_previous_text=False,          # Disable context dependency
                                     temperature=0.0)                           # Use deterministic decoding
                text = res.get("text", "").strip()
                outputs.append(text)
            segment_files_used = seg_files
        else:
            # Single file transcription with optimized settings
            res = model.transcribe(audio_path, language=None,
                                 compression_ratio_threshold=float('inf'),  # Disable repetitive content detection  
                                 logprob_threshold=-1.0,                    # Disable low-confidence filtering
                                 no_speech_threshold=0.1,                   # Lower threshold for "no speech"
                                 condition_on_previous_text=False,          # Disable context dependency
                                 temperature=0.0)                           # Use deterministic decoding
            outputs = [res.get("text", "").strip()]

        full_text = "\n\n".join(outputs)

        if punctuate:
            print("Running punctuation model to restore punctuation...")
            pm = PunctuationModel()
            full_text = pm.restore_punctuation(full_text)

        # clean fillers
        full_text = clean_fillers(full_text)

        # Determine output directory and base filename  
        if output_dir:
            # Use the specified output directory
            output_base = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0])
        else:
            # Use Downloads folder as default
            downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            output_base = os.path.join(downloads_dir, os.path.splitext(os.path.basename(input_path))[0])

        # Save transcription as text file
        out_txt_path = output_base + "_transcription.txt"
        with open(out_txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        print(f"Transcription saved to {out_txt_path}")
        log_outputs.append(os.path.abspath(out_txt_path))

        doc = Document()
        doc.add_heading("Transcription", 0)
        for para in split_into_paragraphs(full_text):
            doc.add_paragraph(para)
        out_docx = output_base + ".docx"
        doc.save(out_docx)
        print(f"Word document saved to {out_docx}")
        log_outputs.append(os.path.abspath(out_docx))

        # include segments if kept
        if keep_temp and segment_files_used:
            log_outputs.extend([os.path.abspath(p) for p in segment_files_used])

        end_time = time.time()
        elapsed = end_time - start_time

        # write log file next to the output txt (or next to input if outputs were temp)
        orig_base = os.path.splitext(os.path.basename(input_path))[0]
        log_name = f"{orig_base}_transcription_log.txt"
        # prefer writing next to project folder or same dir as output
        log_dir = os.path.dirname(os.path.abspath(out_txt_path)) if out_txt_path else os.path.dirname(os.path.abspath(input_path))
        log_path = os.path.join(log_dir, log_name)
        
        # Convert paths to forward slashes for readability
        input_path_clean = os.path.abspath(input_path).replace('\\', '/')
        output_paths_clean = [os.path.abspath(p).replace('\\', '/') for p in log_outputs]
        
        try:
            with open(log_path, "w", encoding="utf-8") as lf:
                lf.write("TRANSCRIPTION LOG\n")
                lf.write("=" * 50 + "\n\n")
                lf.write(f"Input File: {input_path_clean}\n")
                lf.write(f"Model: {model_name}\n")
                lf.write(f"Device: {device}\n")
                lf.write(f"Preprocess: {bool(preprocess)}\n")
                lf.write(f"VAD Segmentation: {bool(vad)}\n")
                lf.write(f"Punctuation Restore: {bool(punctuate)}\n")
                lf.write(f"Processing Time: {format_duration(elapsed)}\n")
                if media_length is not None:
                    lf.write(f"Original Length: {format_duration(media_length)}\n")
                lf.write(f"Temp Files Removed: {not keep_temp}\n\n")
                lf.write("Output Files:\n")
                for output_path in output_paths_clean:
                    lf.write(f"  - {output_path}\n")
            print(f"Log written to {log_path}")
        except Exception as e:
            print(f"Failed to write log file: {e}")

        # Memory cleanup between files to prevent accumulation
        print("🧹 Cleaning up memory and GPU cache...")
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        import gc
        gc.collect()
        
        # Monitor memory after cleanup
        import psutil
        memory_after = psutil.virtual_memory()
        if torch.cuda.is_available():
            try:
                gpu_memory_after = torch.cuda.memory_allocated() / (1024**3)
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available, GPU {gpu_memory_after:.1f}GB used")
            except:
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")
        else:
            print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")

    finally:
        if not keep_temp:
            for p in temp_files:
                try:
                    os.remove(p)
                except Exception:
                    pass

    return out_txt_path


def frames_from_pcm(pcm_bytes, frame_duration_ms=30, sample_rate=16000):
    bytes_per_frame = int(sample_rate * 2 * (frame_duration_ms / 1000.0))
    for i in range(0, len(pcm_bytes), bytes_per_frame):
        yield pcm_bytes[i:i+bytes_per_frame]


def get_pcm_from_file(input_path):
    # Use ffmpeg to output raw PCM s16le 16k mono
    ff = shutil.which("ffmpeg") or None
    if not ff:
        ffexe = None
        if hasattr(iio_ffmpeg, "get_ffmpeg_exe"):
            try:
                ffexe = iio_ffmpeg.get_ffmpeg_exe()
            except Exception:
                ffexe = None
        if not ffexe and hasattr(iio_ffmpeg, "get_exe"):
            try:
                ffexe = iio_ffmpeg.get_exe()
            except Exception:
                ffexe = None
        ff = ffexe
    cmd = [ff, "-i", input_path, "-f", "s16le", "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000", "-"]
    p = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.DEVNULL)
    data = p.stdout.read()
    p.stdout.close()
    return data


def vad_segment_times(input_path, aggressiveness=3, frame_duration_ms=30, padding_ms=300):
    """Voice Activity Detection to segment audio. Falls back to simple duration-based segments if webrtcvad unavailable."""
    
    if not WEBRTCVAD_AVAILABLE:
        print("⚠️  webrtcvad not available - using simple duration-based segmentation")
        # Fallback: create segments based on duration (every 30 seconds)
        try:
            audio_clip = AudioFileClip(input_path)
            duration = audio_clip.duration
            audio_clip.close()
            
            segments = []
            segment_length = 30.0  # 30 second segments
            for i in range(0, int(duration), int(segment_length)):
                start = float(i)
                end = min(float(i + segment_length), duration)
                segments.append((start, end))
            
            print(f"📊 Created {len(segments)} duration-based segments ({segment_length}s each)")
            return segments
            
        except Exception as e:
            print(f"❌ Error creating fallback segments: {e}")
            # Last resort: single segment for entire audio
            return [(0.0, 60.0)]  # Assume 60s max, will be clipped later
    
    # Original webrtcvad implementation
    pcm = get_pcm_from_file(input_path)
    vad = webrtcvad.Vad(aggressiveness)
    frames = list(frames_from_pcm(pcm, frame_duration_ms=frame_duration_ms))
    sample_rate = 16000
    in_speech = False
    segments = []
    speech_start = 0
    for i, frame in enumerate(frames):
        is_speech = False
        if len(frame) == int(sample_rate * 2 * (frame_duration_ms/1000.0)):
            is_speech = vad.is_speech(frame, sample_rate)
        t = (i * frame_duration_ms) / 1000.0
        if is_speech and not in_speech:
            in_speech = True
            speech_start = t
        elif not is_speech and in_speech:
            in_speech = False
            speech_end = t
            # expand by padding
            start = max(0, speech_start - (padding_ms/1000.0))
            end = speech_end + (padding_ms/1000.0)
            segments.append((start, end))
    # if file ends while in speech
    if in_speech:
        speech_end = (len(frames) * frame_duration_ms) / 1000.0
        start = max(0, speech_start - (padding_ms/1000.0))
        end = speech_end + (padding_ms/1000.0)
        segments.append((start, end))
    return segments


def extract_segment(input_path, start, end, out_path, bitrate="128k"):
    ff = shutil.which("ffmpeg") or None
    if not ff:
        ffexe = None
        if hasattr(iio_ffmpeg, "get_ffmpeg_exe"):
            try:
                ffexe = iio_ffmpeg.get_ffmpeg_exe()
            except Exception:
                ffexe = None
        if not ffexe and hasattr(iio_ffmpeg, "get_exe"):
            try:
                ffexe = iio_ffmpeg.get_exe()
            except Exception:
                ffexe = None
        ff = ffexe
    cmd = [
        ff,
        "-y",
        "-i",
        input_path,
        "-ss",
        str(start),
        "-to",
        str(end),
        "-ac",
        "1",
        "-ar",
        "16000",
        "-codec:a",
        "libmp3lame",
        "-b:a",
        bitrate,
        out_path,
    ]
    subprocess.run(cmd, check=True)


def clean_fillers(text):
    # simple filler removal
    fillers = [r"\bumm?\b", r"\buh\b", r"\bhello\b", r"\bhi\b", r"\bsorry\b", r"\byou know\b"]
    pattern = re.compile("|".join(fillers), flags=re.IGNORECASE)
    cleaned = pattern.sub("", text)
    # collapse extra spaces
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def transcribe_file_no_vad(input_path, model_name="large", preprocess=True, keep_temp=False, bitrate="192k", device_preference="auto", punctuate=True, output_dir=None):
    """
    Transcribe a single audio/video file WITHOUT VAD segmentation - transcribes entire file as one piece.
    This is useful for troubleshooting when VAD is cutting out too much content.
    
    Args:
        input_path: Path to input audio/video file
        model_name: Whisper model ("large" or "medium")
        preprocess: Enable preprocessing (always True for quality)
        keep_temp: Keep temporary files
        bitrate: Audio bitrate for preprocessing
        device_preference: Device preference ("auto", "cpu", "cuda", "dml")
        punctuate: Enable punctuation restoration (always True for quality)
        output_dir: Output directory (defaults to Downloads)
    """
    ensure_ffmpeg_in_path()

    start_time = time.time()

    temp_files = []
    log_outputs = []
    try:
        ext = os.path.splitext(input_path)[1].lower()
        # If user requested preprocessing, always create a temp file with processed audio
        if preprocess:
            tf = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
            tf.close()
            pre_path = tf.name
            print(f"Preprocessing audio -> {pre_path}...")
            preprocess_audio(input_path, pre_path, bitrate=bitrate)
            audio_path = pre_path
            temp_files.append(pre_path)
        else:
            # Convert non-mp3 inputs to mp3
            if ext != ".mp3":
                print(f"Converting {input_path} to MP3 (compressed)...")
                audio_path = convert_to_mp3(input_path, bitrate=bitrate)
                temp_files.append(audio_path)
            else:
                audio_path = input_path

        media_length = get_media_duration(input_path)

        print(f"Loading Whisper model '{model_name}' (this may take a while)...")
        # Select device
        device = choose_device(device_preference)
        print(f"Selected device: {device}")
        print(f"CUDA available: {torch.cuda.is_available()}")
        if torch.cuda.is_available():
            print(f"GPU: {torch.cuda.get_device_name(0)}")
        
        # Whisper's load_model will place parameters on the detected device; load then ensure placement
        model = whisper.load_model(model_name, device=device)
        print(f"Model loaded on device: {next(model.parameters()).device}")
        
        # Verify the correct model was loaded by checking model size
        model_size = sum(p.numel() for p in model.parameters())
        print(f"Model size: {model_size:,} parameters")
        
        # Estimate model type based on size (approximate)
        if model_size > 700000000:  # Large model has ~1.5B parameters
            loaded_model_type = "large"
        elif model_size > 350000000:  # Large model has ~770M parameters (when cached medium loads as large)
            loaded_model_type = "large"
        elif model_size > 70000000:  # Small model has ~244M parameters
            loaded_model_type = "small"
        elif model_size > 35000000:  # Base model has ~74M parameters
            loaded_model_type = "base"
        else:
            loaded_model_type = "tiny"
        
        print(f"Detected loaded model type: {loaded_model_type}")
        
        if loaded_model_type != model_name:
            print(f"⚠️  WARNING: Requested model '{model_name}' but loaded '{loaded_model_type}' (cached model used)")
            print("   To force loading the correct model, delete the Whisper cache and try again")
        else:
            print(f"✓ Correct model '{model_name}' loaded successfully")
        
        # Verify model is on correct device
        if device == "cuda" and torch.cuda.is_available():
            if next(model.parameters()).device.type != "cuda":
                print(f"WARNING: Model not on CUDA! Attempting to move...")
                model = model.to("cuda")
                print(f"Model moved to: {next(model.parameters()).device}")

        print("Transcribing entire file (no VAD segmentation)...")
        # Single file transcription with optimized settings
        res = model.transcribe(audio_path, language=None,
                             compression_ratio_threshold=float('inf'),  # Disable repetitive content detection  
                             logprob_threshold=-1.0,                    # Disable low-confidence filtering
                             no_speech_threshold=0.1,                   # Lower threshold for "no speech"
                             condition_on_previous_text=False,          # Disable context dependency
                             temperature=0.0)                           # Use deterministic decoding
        full_text = res.get("text", "").strip()

        if punctuate:
            print("Running punctuation model to restore punctuation...")
            pm = PunctuationModel()
            full_text = pm.restore_punctuation(full_text)

        # clean fillers
        full_text = clean_fillers(full_text)

        # Determine output directory and base filename  
        if output_dir:
            # Use the specified output directory
            output_base = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0])
        else:
            # Use Downloads folder as default
            downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            output_base = os.path.join(downloads_dir, os.path.splitext(os.path.basename(input_path))[0])

        # Save transcription as text file
        out_txt_path = output_base + "_no_vad_transcription.txt"
        with open(out_txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        print(f"Transcription saved to {out_txt_path}")
        log_outputs.append(os.path.abspath(out_txt_path))

        doc = Document()
        doc.add_heading("Transcription (No VAD)", 0)
        for para in split_into_paragraphs(full_text):
            doc.add_paragraph(para)
        out_docx = output_base + "_no_vad.docx"
        doc.save(out_docx)
        print(f"Word document saved to {out_docx}")
        log_outputs.append(os.path.abspath(out_docx))

        end_time = time.time()
        elapsed = end_time - start_time

        # write log file next to the output txt
        orig_base = os.path.splitext(os.path.basename(input_path))[0]
        log_name = f"{orig_base}_no_vad_transcription_log.txt"
        log_dir = os.path.dirname(os.path.abspath(out_txt_path))
        log_path = os.path.join(log_dir, log_name)
        
        # Convert paths to forward slashes for readability
        input_path_clean = os.path.abspath(input_path).replace('\\', '/')
        output_paths_clean = [os.path.abspath(p).replace('\\', '/') for p in log_outputs]
        
        try:
            with open(log_path, "w", encoding="utf-8") as lf:
                lf.write("TRANSCRIPTION LOG (No VAD)\n")
                lf.write("=" * 50 + "\n\n")
                lf.write(f"Input File: {input_path_clean}\n")
                lf.write(f"Model: {model_name}\n")
                lf.write(f"Device: {device}\n")
                lf.write(f"Preprocess: {bool(preprocess)}\n")
                lf.write(f"VAD Segmentation: DISABLED\n")
                lf.write(f"Punctuation Restore: {bool(punctuate)}\n")
                lf.write(f"Processing Time: {format_duration(elapsed)}\n")
                if media_length is not None:
                    lf.write(f"Original Length: {format_duration(media_length)}\n")
                lf.write(f"Temp Files Removed: {not keep_temp}\n\n")
                lf.write("Output Files:\n")
                for output_path in output_paths_clean:
                    lf.write(f"  - {output_path}\n")
            print(f"Log written to {log_path}")
        except Exception as e:
            print(f"Failed to write log file: {e}")

        # Memory cleanup between files to prevent accumulation
        print("🧹 Cleaning up memory and GPU cache...")
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        import gc
        gc.collect()
        
        # Monitor memory after cleanup
        import psutil
        memory_after = psutil.virtual_memory()
        if torch.cuda.is_available():
            try:
                gpu_memory_after = torch.cuda.memory_allocated() / (1024**3)
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available, GPU {gpu_memory_after:.1f}GB used")
            except:
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")
        else:
            print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")

    finally:
        if not keep_temp:
            for p in temp_files:
                try:
                    os.remove(p)
                except Exception:
                    pass

    return out_txt_path


def vad_segment_times_lecture(input_path, aggressiveness=0, frame_duration_ms=30, padding_ms=500):
    """Voice Activity Detection optimized for lectures with less optimised settings."""
    if not WEBRTCVAD_AVAILABLE:
        print("⚠️  webrtcvad not available - using simple duration-based segmentation for lectures")
        # Fallback: create segments based on duration (every 60 seconds for lectures)
        try:
            audio_clip = AudioFileClip(input_path)
            duration = audio_clip.duration
            audio_clip.close()

            segments = []
            segment_length = 60.0  # 60 second segments for lectures
            for i in range(0, int(duration), int(segment_length)):
                start = float(i)
                end = min(float(i + segment_length), duration)
                segments.append((start, end))

            print(f"📊 Created {len(segments)} lecture-duration-based segments ({segment_length}s each)")
            return segments

        except Exception as e:
            print(f"❌ Error creating lecture fallback segments: {e}")
            # Last resort: single segment for entire audio
            return [(0.0, 60.0)]  # Assume 60s max, will be clipped later

    # Original webrtcvad implementation with lecture-optimized settings
    pcm = get_pcm_from_file(input_path)
    vad = webrtcvad.Vad(aggressiveness)  # Less optimised for lectures
    frames = list(frames_from_pcm(pcm, frame_duration_ms=frame_duration_ms))
    sample_rate = 16000
    in_speech = False
    segments = []
    speech_start = 0
    for i, frame in enumerate(frames):
        is_speech = False
        if len(frame) == int(sample_rate * 2 * (frame_duration_ms/1000.0)):
            is_speech = vad.is_speech(frame, sample_rate)
        t = (i * frame_duration_ms) / 1000.0
        if is_speech and not in_speech:
            in_speech = True
            speech_start = t
        elif not is_speech and in_speech:
            in_speech = False
            speech_end = t
            # Expand by more padding for lectures (speakers may pause)
            start = max(0, speech_start - (padding_ms/1000.0))
            end = speech_end + (padding_ms/1000.0)
            segments.append((start, end))
    # If file ends while in speech
    if in_speech:
        speech_end = (len(frames) * frame_duration_ms) / 1000.0
        start = max(0, speech_start - (padding_ms/1000.0))
        end = speech_end + (padding_ms/1000.0)
        segments.append((start, end))
    return segments


def post_process_segments_lecture(segments, min_duration=0.1, merge_gap=1.0, max_segments=500):
    """Post-process segments optimized for lectures - less restrictive than standard VAD."""
    if not segments:
        return []
    # Sort
    segs = sorted(segments, key=lambda s: s[0])
    merged = [list(segs[0])]
    for s, e in segs[1:]:
        last = merged[-1]
        # Merge with larger gap for lectures (speakers pause more)
        if s - last[1] <= merge_gap:
            last[1] = max(last[1], e)
        else:
            merged.append([s, e])
    # Filter short segments (but less restrictive for lectures)
    filtered = []
    for s, e in merged:
        if (e - s) >= min_duration:
            filtered.append((s, e))
    # Allow more segments for lectures
    if len(filtered) > max_segments:
        # Keep longest segments
        filtered = sorted(filtered, key=lambda t: t[1]-t[0], reverse=True)[:max_segments]
        filtered = sorted(filtered, key=lambda t: t[0])
    return filtered


def post_process_segments(segments, min_duration=0.5, merge_gap=0.3, max_segments=1000):
    """Merge nearby segments, drop very short ones, and cap total segments.
    segments: list of (start,end) tuples sorted by start
    """
    if not segments:
        return []
    # sort
    segs = sorted(segments, key=lambda s: s[0])
    merged = [list(segs[0])]
    for s, e in segs[1:]:
        last = merged[-1]
        # if gap between last end and this start is small, merge
        if s - last[1] <= merge_gap:
            last[1] = max(last[1], e)
        else:
            merged.append([s, e])
    # filter short segments
    filtered = []
    for s, e in merged:
        if (e - s) >= min_duration:
            filtered.append((s, e))
    # cap number of segments - use merging instead of dropping for better coverage
    if len(filtered) > max_segments:
        print(f"⚠️  Found {len(filtered)} segments, merging to reduce to {max_segments}...")
        # Instead of dropping segments, merge smaller ones together
        # Sort by duration (shortest first) and merge adjacent segments
        filtered.sort(key=lambda x: x[1] - x[0])  # Sort by duration ascending

        # Keep the longest segments as-is, merge the shortest ones
        keep_count = min(max_segments, len(filtered) // 2)  # Keep at least half
        longest_segments = filtered[-keep_count:]  # Keep longest segments
        segments_to_merge = filtered[:-keep_count]  # Shortest segments to merge

        # Merge shortest segments by combining adjacent ones
        merged_short = []
        i = 0
        while i < len(segments_to_merge):
            current = list(segments_to_merge[i])
            # Try to merge with next segment if they're close
            if i + 1 < len(segments_to_merge):
                next_seg = segments_to_merge[i + 1]
                if next_seg[0] - current[1] <= merge_gap * 2:  # More lenient gap for merging
                    current[1] = max(current[1], next_seg[1])
                    i += 2  # Skip next segment since we merged it
                else:
                    i += 1
            else:
                i += 1
            merged_short.append(tuple(current))

        # Combine longest segments with merged short segments
        final_segments = longest_segments + merged_short
        final_segments.sort(key=lambda x: x[0])  # Sort by start time

        print(f"✅ Merged to {len(final_segments)} segments (kept {len(longest_segments)} longest, merged {len(merged_short)} from {len(segments_to_merge)} short segments)")
        return final_segments
    else:
        return filtered


def transcribe_lecture(input_path, model_name="large", preprocess=True, keep_temp=False, bitrate="192k", device_preference="auto", punctuate=True, output_dir=None):
    """
    Specialized transcription for lectures with less optimised VAD settings.
    Lectures often have pauses, background noise, and varying audio quality that standard VAD misses.
    """
    ensure_ffmpeg_in_path()

    start_time = time.time()

    temp_files = []
    log_outputs = []
    try:
        ext = os.path.splitext(input_path)[1].lower()
        # If user requested preprocessing, always create a temp file with processed audio
        if preprocess:
            tf = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
            tf.close()
            pre_path = tf.name
            print(f"Preprocessing audio -> {pre_path}...")
            preprocess_audio(input_path, pre_path, bitrate=bitrate)
            audio_path = pre_path
            temp_files.append(pre_path)
        else:
            # Convert non-mp3 inputs to mp3
            if ext != ".mp3":
                print(f"Converting {input_path} to MP3 (compressed)...")
                audio_path = convert_to_mp3(input_path, bitrate=bitrate)
                temp_files.append(audio_path)
            else:
                audio_path = input_path

        media_length = get_media_duration(input_path)

        print(f"Loading Whisper model '{model_name}' (this may take a while)...")
        # Select device
        device = choose_device(device_preference)
        print(f"Selected device: {device}")
        print(f"CUDA available: {torch.cuda.is_available()}")
        if torch.cuda.is_available():
            print(f"GPU: {torch.cuda.get_device_name(0)}")

        # Whisper's load_model will place parameters on the detected device; load then ensure placement
        model = whisper.load_model(model_name, device=device)
        print(f"Model loaded on device: {next(model.parameters()).device}")

        # Verify the correct model was loaded by checking model size
        model_size = sum(p.numel() for p in model.parameters())
        print(f"Model size: {model_size:,} parameters")

        # Estimate model type based on size (approximate)
        if model_size > 700000000:  # Large model has ~1.5B parameters
            loaded_model_type = "large"
        elif model_size > 350000000:  # Large model has ~770M parameters (when cached medium loads as large)
            loaded_model_type = "large"
        elif model_size > 70000000:  # Small model has ~244M parameters
            loaded_model_type = "small"
        elif model_size > 35000000:  # Base model has ~74M parameters
            loaded_model_type = "base"
        else:
            loaded_model_type = "tiny"

        print(f"Detected loaded model type: {loaded_model_type}")

        if loaded_model_type != model_name:
            print(f"⚠️  WARNING: Requested model '{model_name}' but loaded '{loaded_model_type}' (cached model used)")
            print("   To force loading the correct model, delete the Whisper cache and try again")
        else:
            print(f"✓ Correct model '{model_name}' loaded successfully")

        # Verify model is on correct device
        if device == "cuda" and torch.cuda.is_available():
            if next(model.parameters()).device.type != "cuda":
                print(f"WARNING: Model not on CUDA! Attempting to move...")
                model = model.to("cuda")
                print(f"Model moved to: {next(model.parameters()).device}")

        outputs = []
        segment_files_used = []

        print("Running LECTURE-OPTIMIZED VAD segmentation...")
        print("Using less optimised settings for lecture audio...")

        # Less optimised VAD settings for lectures
        segments = vad_segment_times_lecture(audio_path, aggressiveness=0, frame_duration_ms=30, padding_ms=500)
        print(f"Raw segments found: {len(segments)}")

        # Less restrictive post-processing for lectures
        segments = post_process_segments_lecture(segments, min_duration=0.1, merge_gap=1.0, max_segments=500)
        print(f"Segments after lecture-optimized post-processing: {len(segments)}")

        if len(segments) == 0:
            print("⚠️  No segments found with VAD, falling back to transcribing entire file...")
            # Fallback to no-VAD if no segments found
            res = model.transcribe(audio_path, language=None,
                                 compression_ratio_threshold=float('inf'),  # Disable repetitive content detection
                                 logprob_threshold=-1.0,                    # Disable low-confidence filtering
                                 no_speech_threshold=0.05,                  # Very low threshold for lectures
                                 condition_on_previous_text=False,          # Disable context dependency
                                 temperature=0.0)                           # Use deterministic decoding
            # Handle both dict and list return types from Whisper
            if isinstance(res, dict):
                outputs = [res.get("text", "").strip()]
            elif isinstance(res, list) and len(res) > 0:
                outputs = [str(res[0]).strip()]
            else:
                outputs = [""]
        else:
            seg_files = []
            for idx, (s, e) in enumerate(segments):
                seg_tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
                seg_tmp.close()
                print(f"Extracting segment {idx+1}/{len(segments)}: {s:.2f}-{e:.2f}s ({e-s:.2f}s duration)")
                extract_segment(audio_path, s, e, seg_tmp.name, bitrate=bitrate)
                seg_files.append(seg_tmp.name)
                temp_files.append(seg_tmp.name)

            # Transcribe each segment with lecture-optimized settings
            for seg in seg_files:
                print(f"Transcribing segment {seg}")
                # Lecture-optimized Whisper settings
                res = model.transcribe(seg, language=None,
                                     compression_ratio_threshold=float('inf'),  # Disable repetitive content detection
                                     logprob_threshold=-2.0,                    # Even more permissive confidence threshold
                                     no_speech_threshold=0.05,                  # Very low threshold for lectures
                                     condition_on_previous_text=False,          # Disable context dependency
                                     temperature=0.0,                           # Use deterministic decoding
                                     beam_size=5,                              # Higher quality beam search
                                     patience=2.0)                             # More thorough processing
                # Handle both dict and list return types from Whisper
                if isinstance(res, dict):
                    text = res.get("text", "").strip()
                elif isinstance(res, list) and len(res) > 0:
                    text = str(res[0]).strip()
                else:
                    text = ""
                if text:  # Only add non-empty transcriptions
                    outputs.append(text)

            segment_files_used = seg_files

        full_text = "\n\n".join(outputs)

        if punctuate:
            print("Running punctuation model to restore punctuation...")
            pm = PunctuationModel()
            full_text = pm.restore_punctuation(full_text)

        # Clean fillers but be less optimised for lectures
        full_text = clean_fillers_lecture(full_text)

        # Determine output directory and base filename
        if output_dir:
            # Use the specified output directory
            output_base = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0])
        else:
            # Use Downloads folder as default
            downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            output_base = os.path.join(downloads_dir, os.path.splitext(os.path.basename(input_path))[0])

        # Save transcription as text file
        out_txt_path = output_base + "_lecture_transcription.txt"
        with open(out_txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        print(f"Lecture transcription saved to {out_txt_path}")
        log_outputs.append(os.path.abspath(out_txt_path))

        doc = Document()
        doc.add_heading("Lecture Transcription", 0)
        doc.add_paragraph(f"Source: {os.path.basename(input_path)}")
        if media_length:
            doc.add_paragraph(f"Duration: {format_duration(media_length)}")
        doc.add_paragraph(f"Model: {model_name} (Lecture-Optimized)")
        doc.add_paragraph(f"Segments Processed: {len(outputs)}")
        doc.add_paragraph("")

        for para in split_into_paragraphs(full_text):
            doc.add_paragraph(para)
        out_docx = output_base + "_lecture.docx"
        doc.save(out_docx)
        print(f"Word document saved to {out_docx}")
        log_outputs.append(os.path.abspath(out_docx))

        # Include segments if kept
        if keep_temp and segment_files_used:
            log_outputs.extend([os.path.abspath(p) for p in segment_files_used])

        end_time = time.time()
        elapsed = end_time - start_time

        # Write log file
        orig_base = os.path.splitext(os.path.basename(input_path))[0]
        log_name = f"{orig_base}_lecture_transcription_log.txt"
        log_dir = os.path.dirname(os.path.abspath(out_txt_path)) if out_txt_path else os.path.dirname(os.path.abspath(input_path))
        log_path = os.path.join(log_dir, log_name)

        # Convert paths to forward slashes for readability
        input_path_clean = os.path.abspath(input_path).replace('\\', '/')
        output_paths_clean = [os.path.abspath(p).replace('\\', '/') for p in log_outputs]

        try:
            with open(log_path, "w", encoding="utf-8") as lf:
                lf.write("LECTURE TRANSCRIPTION LOG\n")
                lf.write("=" * 50 + "\n\n")
                lf.write(f"Input File: {input_path_clean}\n")
                lf.write(f"Model: {model_name} (Lecture-Optimized)\n")
                lf.write(f"Device: {device}\n")
                lf.write(f"Preprocess: {bool(preprocess)}\n")
                lf.write(f"VAD Segmentation: LECTURE-OPTIMIZED\n")
                lf.write(f"Punctuation Restore: {bool(punctuate)}\n")
                lf.write(f"Processing Time: {format_duration(elapsed)}\n")
                if media_length is not None:
                    lf.write(f"Original Length: {format_duration(media_length)}\n")
                lf.write(f"Segments Found: {len(segments) if 'segments' in locals() else 0}\n")
                lf.write(f"Segments Transcribed: {len(outputs)}\n")
                lf.write(f"Temp Files Removed: {not keep_temp}\n\n")
                lf.write("Output Files:\n")
                for output_path in output_paths_clean:
                    lf.write(f"  - {output_path}\n")
            print(f"Log written to {log_path}")
        except Exception as e:
            print(f"Failed to write log file: {e}")

        # Memory cleanup between files to prevent accumulation
        print("🧹 Cleaning up memory and GPU cache...")
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        import gc
        gc.collect()
        
        # Monitor memory after cleanup
        import psutil
        memory_after = psutil.virtual_memory()
        if torch.cuda.is_available():
            try:
                gpu_memory_after = torch.cuda.memory_allocated() / (1024**3)
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available, GPU {gpu_memory_after:.1f}GB used")
            except:
                print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")
        else:
            print(f"📊 Memory after cleanup: RAM {memory_after.available / (1024**3):.1f}GB available")

    finally:
        if not keep_temp:
            for p in temp_files:
                try:
                    os.remove(p)
                except Exception:
                    pass

    return out_txt_path


def clean_fillers_lecture(text):
    """Clean fillers but be less optimised for lectures (preserve academic language)."""
    # Less optimised filler removal for lectures
    fillers = [r"\bumm?\b", r"\buh\b", r"\bhello\b", r"\bhi\b"]
    pattern = re.compile("|".join(fillers), flags=re.IGNORECASE)
    cleaned = pattern.sub("", text)
    # Collapse extra spaces
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def main(argv=None):
    parser = argparse.ArgumentParser(description="Transcribe audio/video to text and docx with optimized settings")
    parser.add_argument("input", help="input audio or video file")
    parser.add_argument("--model", default="large", help="whisper model: large, medium")
    parser.add_argument("--keep-temp", action="store_true", help="keep temporary files")
    parser.add_argument("--bitrate", default="192k", help="mp3 bitrate for converted audio")
    parser.add_argument("--device", default="auto", help="device preference: auto/cpu/cuda/dml")
    parser.add_argument("--output-dir", dest="output_dir", help="output directory for txt and docx files (default: Downloads)")
    args = parser.parse_args(argv)

    if not os.path.isfile(args.input):
        print(f"File not found: {args.input}")
        sys.exit(1)

    transcribe_file(args.input, model_name=args.model, keep_temp=args.keep_temp, bitrate=args.bitrate, device_preference=args.device, output_dir=args.output_dir)


if __name__ == "__main__":
    main()
