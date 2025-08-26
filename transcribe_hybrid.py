"""
Hybrid GPU+CPU transcription for maximum performance.

This module implements parallel processing using both CUDA GPU and CPU cores
simultaneously for optimal transcription speed on systems with both capabilities.
"""
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="webrtcvad")

import sys
import os
import shutil
import subprocess
import tempfile
import argparse
import whisper
import torch
from docx import Document
from moviepy import VideoFileClip, AudioFileClip
import imageio_ffmpeg as iio_ffmpeg
from deepmultilingualpunctuation import PunctuationModel
import webrtcvad
import re
import time
import json
import multiprocessing
import threading
import queue
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed


def choose_hybrid_devices():
    """Determine optimal device configuration for hybrid processing."""
    devices = {"primary": "cpu", "secondary": None}
    
    if torch.cuda.is_available():
        devices["primary"] = "cuda"
        # Use CPU as secondary for preprocessing/postprocessing
        devices["secondary"] = "cpu"
        print(f"Hybrid mode: Primary=CUDA ({torch.cuda.get_device_name(0)}), Secondary=CPU ({multiprocessing.cpu_count()} cores)")
    else:
        devices["primary"] = "cpu"
        print(f"CPU-only mode: {multiprocessing.cpu_count()} cores available")
    
    return devices


def load_models_parallel(model_name="medium"):
    """Load Whisper models on both GPU and CPU if available."""
    devices = choose_hybrid_devices()
    models = {}
    
    # Load primary model (usually CUDA)
    print(f"Loading primary model on {devices['primary']}...")
    models['primary'] = whisper.load_model(model_name, device=devices['primary'])
    
    # Load secondary model if we have hybrid setup
    if devices['secondary'] and devices['secondary'] != devices['primary']:
        try:
            print(f"Loading secondary model on {devices['secondary']}...")
            models['secondary'] = whisper.load_model(model_name, device=devices['secondary'])
        except Exception as e:
            print(f"Failed to load secondary model: {e}")
            models['secondary'] = None
    
    return models, devices


def transcribe_segment_optimized(model, audio_path, segment_info=None):
    """Transcribe a single segment with optimized parameters."""
    start_time = None
    end_time = None
    
    if segment_info:
        segment_path, start_time, end_time = segment_info
        audio_to_transcribe = segment_path
    else:
        audio_to_transcribe = audio_path
    
    try:
        # Optimized Whisper parameters for quality
        result = model.transcribe(
            audio_to_transcribe,
            language=None,
            compression_ratio_threshold=float('inf'),  # Bypass content filtering
            logprob_threshold=-1.0,                    # Accept lower confidence
            no_speech_threshold=0.1,                   # Lower speech detection threshold
            condition_on_previous_text=False,          # Disable context dependency
            temperature=0.0                            # Deterministic decoding
        )
        
        text = result.get("text", "").strip()
        if segment_info:
            return {"text": text, "start": start_time, "end": end_time}
        else:
            return {"text": text}
            
    except Exception as e:
        print(f"Transcription error: {e}")
        if segment_info:
            return {"text": "", "start": start_time, "end": end_time}
        else:
            return {"text": ""}


def parallel_transcribe_segments(models, devices, segments, max_workers=None):
    """Transcribe segments in parallel using multiple models/devices."""
    if not max_workers:
        # Use number of available cores, but cap to reasonable limit
        max_workers = min(multiprocessing.cpu_count(), 8)
    
    results = []
    segment_queue = queue.Queue()
    
    # Add all segments to queue
    for i, segment in enumerate(segments):
        segment_queue.put((i, segment))
    
    def worker(worker_id, model_key):
        """Worker function for parallel transcription."""
        model = models.get(model_key)
        if not model:
            return
            
        local_results = []
        device_name = devices[model_key] if model_key in devices else "unknown"
        
        while not segment_queue.empty():
            try:
                seg_id, segment_info = segment_queue.get_nowait()
                print(f"Worker {worker_id} (on {device_name}): Processing segment {seg_id+1}")
                
                result = transcribe_segment_optimized(model, None, segment_info)
                result["segment_id"] = seg_id
                local_results.append(result)
                
            except queue.Empty:
                break
            except Exception as e:
                print(f"Worker {worker_id} error: {e}")
        
        return local_results
    
    # Start workers
    futures = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Primary model workers (GPU if available)
        primary_workers = max_workers // 2 if models.get('secondary') else max_workers
        for i in range(primary_workers):
            future = executor.submit(worker, f"P{i}", "primary")
            futures.append(future)
        
        # Secondary model workers (CPU)
        if models.get('secondary'):
            secondary_workers = max_workers - primary_workers
            for i in range(secondary_workers):
                future = executor.submit(worker, f"S{i}", "secondary")
                futures.append(future)
        
        # Collect results
        all_results = []
        for future in as_completed(futures):
            try:
                worker_results = future.result()
                if worker_results:
                    all_results.extend(worker_results)
            except Exception as e:
                print(f"Worker completion error: {e}")
    
    # Sort results by segment_id
    all_results.sort(key=lambda x: x.get('segment_id', 0))
    return all_results


def extract_segment_parallel(audio_path, segments, output_dir, bitrate="192k"):
    """Extract audio segments in parallel using CPU multiprocessing."""
    segment_files = []
    
    def extract_single_segment(args):
        idx, start_time, end_time = args
        output_path = os.path.join(output_dir, f"segment_{idx:04d}.mp3")
        
        try:
            # Use ffmpeg for fast segment extraction
            cmd = [
                "ffmpeg", "-y", "-i", audio_path,
                "-ss", str(start_time),
                "-t", str(end_time - start_time),
                "-acodec", "libmp3lame",
                "-b:a", bitrate,
                "-ar", "16000",
                "-ac", "1",
                output_path
            ]
            subprocess.run(cmd, capture_output=True, check=True)
            return (idx, output_path, start_time, end_time)
        except Exception as e:
            print(f"Segment extraction error for segment {idx}: {e}")
            return None
    
    # Prepare arguments for parallel processing
    segment_args = [(i, start, end) for i, (start, end) in enumerate(segments)]
    
    # Use multiprocessing for CPU-intensive segment extraction
    with ProcessPoolExecutor(max_workers=min(multiprocessing.cpu_count(), 16)) as executor:
        results = list(executor.map(extract_single_segment, segment_args))
    
    # Filter successful extractions and return in order
    successful_segments = []
    for result in results:
        if result:
            idx, path, start, end = result
            successful_segments.append((path, start, end))
    
    return successful_segments


def transcribe_file_hybrid(input_path, model_name="medium", device_preference="auto", 
                          output_dir=None, max_workers=None):
    """
    Main hybrid transcription function using GPU+CPU parallel processing.
    """
    import tempfile
    from transcribe import (vad_segment_times, post_process_segments, 
                          preprocess_audio, get_media_duration, 
                          split_into_paragraphs, format_duration)
    
    start_time = time.time()
    temp_files = []
    
    try:
        print(f"🚀 Starting HYBRID GPU+CPU transcription of: {input_path}")
        
        # Set up output directory
        if not output_dir:
            output_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        
        # Create temp directory for segments
        temp_dir = tempfile.mkdtemp(prefix="hybrid_transcribe_")
        temp_files.append(temp_dir)
        
        # Get media duration
        duration = get_media_duration(input_path)
        if duration:
            print(f"Media duration: {format_duration(duration)}")
        
        # Preprocess audio
        print("🔄 Preprocessing audio...")
        audio_path = preprocess_audio(input_path, temp_files, bitrate="192k")
        
        # Load models in parallel
        print("🧠 Loading Whisper models...")
        models, devices = load_models_parallel(model_name)
        
        # VAD segmentation
        print("✂️  Running VAD segmentation...")
        segments = vad_segment_times(audio_path, aggressiveness=2, 
                                   frame_duration_ms=30, padding_ms=200)
        print(f"Raw segments found: {len(segments)}")
        
        segments = post_process_segments(segments, min_duration=0.6, 
                                       merge_gap=0.35, max_segments=120)
        print(f"Segments after post-processing: {len(segments)}")
        
        if len(segments) > 1:
            # Extract segments in parallel
            print("🎵 Extracting segments in parallel...")
            segment_files = extract_segment_parallel(audio_path, segments, temp_dir)
            
            # Transcribe segments in parallel using hybrid GPU+CPU
            print(f"🗣️  Transcribing {len(segment_files)} segments in parallel...")
            print(f"Using {max_workers or 'auto'} workers across available devices")
            
            results = parallel_transcribe_segments(models, devices, segment_files, max_workers)
            
            # Combine results
            full_text = " ".join([r.get("text", "") for r in results if r.get("text")])
        else:
            # Single file transcription with primary model
            print("🗣️  Single file transcription...")
            result = transcribe_segment_optimized(models['primary'], audio_path)
            full_text = result.get("text", "")
        
        # Post-processing
        if not full_text.strip():
            print("⚠️  No text transcribed!")
            full_text = "[No speech detected]"
        
        # Apply punctuation
        print("📝 Applying punctuation...")
        try:
            punctuation_model = PunctuationModel()
            full_text = punctuation_model.restore_punctuation(full_text)
        except Exception as e:
            print(f"Punctuation restoration failed: {e}")
        
        # Format into paragraphs
        formatted_text = split_into_paragraphs(full_text, max_length=500)
        if isinstance(formatted_text, list):
            formatted_text = '\n\n'.join(formatted_text)
        
        # Save outputs
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        txt_path = os.path.join(output_dir, f"{base_name}_hybrid.txt")
        docx_path = os.path.join(output_dir, f"{base_name}_hybrid.docx")
        
        # Save text file
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(formatted_text)
        
        # Save Word document
        doc = Document()
        doc.add_heading(f'Transcription: {base_name}', 0)
        
        if duration:
            doc.add_paragraph(f'Duration: {format_duration(duration)}')
        doc.add_paragraph(f'Processing time: {format_duration(time.time() - start_time)}')
        doc.add_paragraph(f'Model: {model_name} (Hybrid GPU+CPU)')
        doc.add_paragraph('')
        
        for para in formatted_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(docx_path)
        
        elapsed = time.time() - start_time
        print(f"✅ Hybrid transcription complete in {format_duration(elapsed)}")
        print(f"📄 Text file: {txt_path}")
        print(f"📄 Word document: {docx_path}")
        
        if duration:
            speedup = duration / elapsed
            print(f"⚡ Processing speed: {speedup:.1f}x realtime")
        
        return txt_path
        
    except Exception as e:
        print(f"❌ Hybrid transcription failed: {e}")
        raise
        
    finally:
        # Cleanup temp files
        for temp_file in temp_files:
            try:
                if os.path.isdir(temp_file):
                    shutil.rmtree(temp_file)
                elif os.path.isfile(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Cleanup warning: {e}")


def main():
    parser = argparse.ArgumentParser(description="Hybrid GPU+CPU speech transcription")
    parser.add_argument("--input", required=True, help="Input audio/video file")
    parser.add_argument("--model", default="medium", choices=["tiny", "base", "small", "medium", "large"],
                       help="Whisper model size")
    parser.add_argument("--output-dir", help="Output directory (default: Downloads)")
    parser.add_argument("--workers", type=int, help="Max parallel workers (default: auto)")
    
    args = parser.parse_args()
    
    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        return 1
    
    try:
        transcribe_file_hybrid(
            args.input,
            model_name=args.model,
            output_dir=args.output_dir,
            max_workers=args.workers
        )
        return 0
    except Exception as e:
        print(f"Error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
