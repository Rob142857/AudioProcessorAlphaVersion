"""
Optimised transcription utilities with safe PyTorch lifecycle management.
This module provides a high-quality single-file transcription path that avoids
re-importing torch and aggressively clearing only model-level caches between runs.
"""
import warnings
import re
import json
warnings.filterwarnings("ignore", category=UserWarning, module="webrtcvad")

# Suppress verbose tqdm progress bars from transformers/huggingface
import os
os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "0")  # Keep minimal progress
os.environ.setdefault("TRANSFORMERS_VERBOSITY", "warning")  # Reduce transformer logs
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")  # Avoid tokenizer warnings

import sys
import time
import gc
import psutil
import argparse
import multiprocessing
import threading
from typing import Any, cast, Optional, Dict

# IMPORTANT: Import torch once at module import time. Do NOT delete torch.* from sys.modules.
try:
    import torch  # type: ignore
    from torch.utils.data import Dataset, DataLoader
    import numpy as np
    _torch_available = True
except Exception as e:  # pragma: no cover
    torch = None  # type: ignore
    Dataset = type(None)  # type: ignore
    DataLoader = type(None)  # type: ignore
    np = None  # type: ignore
    _torch_import_error = e
    _torch_available = False

def _ensure_torch_available():
    if torch is None:
        raise RuntimeError(f"PyTorch is required but failed to import: {_torch_import_error}")


def is_verbatim() -> bool:
    """Return True if verbatim/faithful output is requested (default True).

    Disable by setting TRANSCRIBE_VERBATIM to 0/false/no.
    """
    val = os.environ.get("TRANSCRIBE_VERBATIM", "1").strip()
    return val.lower() not in ("0", "false", "no")


def _is_verbatim_mode() -> bool:
    """Return True when verbatim/faithful mode is enabled. Default: True.

    Controlled by env TRANSCRIBE_VERBATIM: "1"/"true" to enable; "0"/"false" to disable.
    """
    val = str(os.environ.get("TRANSCRIBE_VERBATIM", "1")).strip().lower()
    return val not in ("0", "false", "off")


def _apply_recommended_env_defaults() -> None:
    """Set optimized default environment variables using setdefault.

    We prioritize faithful transcription, safe GPU usage (VRAM margin), deterministic
    behavior (quality/beam off by default), and conservative preprocessing.
    Explicit user overrides always take precedence.
    """
    defaults = {
        # Fidelity & formatting
        "TRANSCRIBE_VERBATIM": "0",            # Enable post-processing for better punctuation/formatting
        "TRANSCRIBE_PARAGRAPH_GAP": "1.5",     # Silence gap (seconds) for paragraph breaks (slightly longer for clearer breaks)
        # Model selection
        "TRANSCRIBE_MODEL_NAME": "large-v3-turbo",   # Turbo model for good speed/accuracy balance
        # GPU safety / fragmentation mitigation
        "TRANSCRIBE_GPU_FRACTION": "0.92",     # Leave headroom instead of 0.99 to reduce OOM risk
        # Processing feature toggles - QUALITY MODE ENABLED BY DEFAULT
        "TRANSCRIBE_QUALITY_MODE": "1",        # Beam search ENABLED for better accuracy
        "TRANSCRIBE_PREPROC_STRONG_FILTERS": "0", # Conservative audio preprocessing
        "TRANSCRIBE_USE_DATASET": "0",         # Disable external segmentation by default
        "TRANSCRIBE_VAD": "0",                 # Disable VAD unless needed
        # Optional force flags disabled
        "TRANSCRIBE_FORCE_GPU": "0",           # Respect preflight memory heuristic
        "TRANSCRIBE_FORCE_FP16": "0",          # Stability over memory unless user demands
        # Perf mode off (user can enable for aggressive thread tweaks)
        "TRANSCRIBE_MAX_PERF": "0",
        # Allow domain-specific prompts for better word recognition
        "TRANSCRIBE_ALLOW_PROMPT": "1",        # Enable initial_prompt from special_words.txt
    }
    alloc_conf_default = "expandable_segments:True,max_split_size_mb:64"
    for k, v in defaults.items():
        os.environ.setdefault(k, v)
    os.environ.setdefault("PYTORCH_CUDA_ALLOC_CONF", alloc_conf_default)

# Apply defaults at import time without overriding user choices
_apply_recommended_env_defaults()

# Compatibility shim for calling model.transcribe() across backends and versions
# Filters out kwargs that the installed backend doesn't support to avoid
# "unexpected keyword argument" errors (e.g., older faster-whisper builds)
import inspect as _inspect  # kept module-level for reuse across call sites

def _compatible_transcribe_call(_model, _audio, _kwargs):
    try:
        sig = _inspect.signature(_model.transcribe)
        allowed = set(sig.parameters.keys())
        # Apply lightweight aliasing for known parameter name differences
        kw = dict(_kwargs or {})
        if 'logprob_threshold' in kw and 'log_prob_threshold' in allowed:
            kw['log_prob_threshold'] = kw.pop('logprob_threshold')
        # Keep only supported keys; leave values unchanged for the rest
        filtered = {k: v for k, v in kw.items() if k in allowed}
    except Exception:
        # If we can’t introspect, pass kwargs through unchanged
        filtered = dict(_kwargs or {})
    return _model.transcribe(_audio, **filtered)

def _as_result_dict(res: Any) -> Dict[str, Any]:
    """Normalize backend-specific transcribe() outputs to a common dict.

    Expected shape:
      { 'text': str, 'segments': [ { 'text': str, 'start': float, 'end': float }, ... ], ... }

    Supports:
      - OpenAI whisper: dict with 'segments'/'text'
      - faster-whisper: (segments_iterable, info) tuple
    """
    try:
        # Native whisper-like result
        if isinstance(res, dict):
            return res

        # faster-whisper returns (segments, info)
        if isinstance(res, tuple) and len(res) == 2:
            segments_iter, info = res
            seg_list = []
            parts = []
            try:
                for s in segments_iter or []:
                    txt = (getattr(s, 'text', '') or '').strip()
                    start = float(getattr(s, 'start', 0.0) or 0.0)
                    end = float(getattr(s, 'end', 0.0) or 0.0)
                    if txt:
                        seg_list.append({'text': txt, 'start': start, 'end': end})
                        parts.append(txt)
            except Exception:
                pass
            out: Dict[str, Any] = {
                'segments': seg_list,
                'text': ' '.join(parts).strip(),
            }
            try:
                lang = getattr(info, 'language', None) or getattr(info, 'language_code', None)
                if lang:
                    out['language'] = lang
            except Exception:
                pass
            return out
    except Exception:
        pass
    return {'segments': [], 'text': ''}


def preprocess_audio_with_padding(input_path: str, temp_dir: str = None) -> str:
    """
    Preprocess audio/video file to high-quality MP3 with silence padding.
    
    Args:
        input_path: Path to input audio/video file
        temp_dir: Directory for temporary files (default: system temp)
        
    Returns:
        Path to the preprocessed MP3 file with padding
        
    Features:
    - Converts any audio/video format to high-quality MP3 (320kbps)
    - Adds 1 second of silence at the beginning
    - Adds 1 second of silence at the end  
    - Normalizes audio levels
    - Prevents missed words at start/end of recordings
    """
    import tempfile
    import subprocess
    import shutil
    
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()
    
    # Generate unique temp filename
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_output = os.path.join(temp_dir, f"preprocessed_{base_name}_{int(time.time())}.mp3")
    
    print(f"🔄 Preprocessing audio with silence padding...")
    print(f"📁 Input: {os.path.basename(input_path)}")
    print(f"📁 Temp output: {os.path.basename(temp_output)}")
    
    try:
        # Check if ffmpeg is available
        ffmpeg_cmd = shutil.which("ffmpeg")
        if not ffmpeg_cmd:
            # Try to use the bundled ffmpeg.exe if available
            bundled_ffmpeg = os.path.join(os.path.dirname(__file__), "ffmpeg.exe")
            if os.path.exists(bundled_ffmpeg):
                ffmpeg_cmd = bundled_ffmpeg
            else:
                raise FileNotFoundError("ffmpeg not found. Please install ffmpeg or ensure it's in PATH.")
        
        # Choose conservative vs strong filters based on env (default: conservative)
        strong_filters = str(os.environ.get("TRANSCRIBE_PREPROC_STRONG_FILTERS", "")).strip() in ("1", "true", "True")
        if strong_filters:
            # Legacy/strong filtering for very noisy tapes
            afilters = "adelay=1000|1000,highpass=f=80,lowpass=f=8000,afftdn=nf=-25,loudnorm,apad=pad_len=48000"
        else:
            # Conservative: light padding + loudness normalization only (preserves fidelity)
            afilters = "adelay=500|500,loudnorm,apad=pad_len=24000"

        # FFmpeg command to:
        # 1) Add brief silence padding to prevent start/end truncation
        # 2) Normalize loudness
        # 3) Optionally apply stronger denoise/EQ if explicitly enabled
        # 4) Convert to high-quality MP3 (320kbps)
        cmd = [
            ffmpeg_cmd,
            "-i", input_path,
            # Audio processing filters (selected above)
            "-af", afilters,
            # High quality MP3 encoding
            "-codec:a", "libmp3lame",
            "-b:a", "320k",
            "-ar", "48000",  # 48kHz sample rate for best quality
            "-ac", "2",      # Stereo output
            # Overwrite output file
            "-y",
            temp_output
        ]
        
        print(f"🔧 Running ffmpeg preprocessing...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        
        if result.returncode != 0:
            print(f"❌ FFmpeg preprocessing failed:")
            print(f"Error: {result.stderr}")
            # Return original file if preprocessing fails
            return input_path
        
        # Verify the output file was created and has reasonable size
        if os.path.exists(temp_output) and os.path.getsize(temp_output) > 1024:  # > 1KB
            print(f"✅ Audio preprocessing completed successfully")
            print(f"📊 Original size: {os.path.getsize(input_path) / (1024*1024):.1f} MB")
            print(f"📊 Preprocessed size: {os.path.getsize(temp_output) / (1024*1024):.1f} MB")
            return temp_output
        else:
            print(f"⚠️  Preprocessing produced invalid output, using original file")
            return input_path
            
    except subprocess.TimeoutExpired:
        print(f"⚠️  FFmpeg preprocessing timed out after 5 minutes, using original file")
        return input_path
    except Exception as e:
        print(f"⚠️  Audio preprocessing failed: {e}")
        print(f"🔄 Continuing with original file...")
        return input_path


class AudioTranscriptionDataset(Dataset):
    """
    PyTorch Dataset for efficient audio transcription with GPU pipeline optimization.

    This dataset enables:
    - Batch processing of audio segments
    - Efficient GPU memory usage
    - Parallel data loading and preprocessing
    - Better utilization of GPU pipelines
    """

    def __init__(self, audio_path: str, segment_length: int = 30, overlap: int = 5):
        """
        Initialize the dataset.

        Args:
            audio_path: Path to the audio file
            segment_length: Length of each audio segment in seconds
            overlap: Overlap between segments in seconds
        """
        self.audio_path = audio_path
        self.segment_length = segment_length
        self.overlap = overlap
        self.segments = []

        # Load audio and create segments
        self._load_and_segment_audio()

    def _load_and_segment_audio(self):
        """Load audio file and create overlapping segments for efficient processing."""
        try:
            import whisper
            from whisper.audio import load_audio

            # Load the audio file
            audio = load_audio(self.audio_path)
            sample_rate = whisper.audio.SAMPLE_RATE
            total_samples = len(audio)

            # Calculate segment parameters
            segment_samples = self.segment_length * sample_rate
            overlap_samples = self.overlap * sample_rate
            step_samples = segment_samples - overlap_samples

            # Create overlapping segments
            start_sample = 0

            while start_sample < total_samples:
                end_sample = min(start_sample + segment_samples, total_samples)
                segment_audio = audio[start_sample:end_sample]

                # Pad short segments if needed
                if len(segment_audio) < segment_samples and np is not None:
                    padding = np.zeros(segment_samples - len(segment_audio))
                    segment_audio = np.concatenate([segment_audio, padding])

                self.segments.append({
                    'audio': segment_audio,
                    'start_time': start_sample / sample_rate,
                    'end_time': end_sample / sample_rate,
                    'segment_id': len(self.segments)
                })

                start_sample += step_samples

        except Exception as e:
            print(f"⚠️  Failed to create audio segments: {e}")
            # Fallback: treat entire file as single segment
            self.segments = [{
                'audio': [],
                'start_time': 0.0,
                'end_time': 0.0,
                'segment_id': 0
            }]

    def __len__(self) -> int:
        """Return the number of segments."""
        return len(self.segments)

    def __getitem__(self, idx: int) -> Dict[str, Any]:
        """Get a segment by index."""
        return self.segments[idx]


def create_efficient_dataloader(audio_path: str, batch_size: int = 4, num_workers: int = 2) -> DataLoader:
    """
    Create an efficient DataLoader for audio transcription.

    Args:
        audio_path: Path to the audio file
        batch_size: Number of segments to process in parallel
        num_workers: Number of worker processes for data loading

    Returns:
        DataLoader configured for efficient GPU processing
    """
    dataset = AudioTranscriptionDataset(audio_path)

    # Configure DataLoader for GPU efficiency
    dataloader = DataLoader(
        dataset,
        batch_size=batch_size,
        shuffle=False,  # Maintain temporal order
        num_workers=num_workers,
        pin_memory=torch.cuda.is_available() if torch and torch.cuda else False,  # Faster GPU transfer
        prefetch_factor=2 if num_workers > 0 else None,  # Prefetch batches
        persistent_workers=num_workers > 0  # Keep workers alive
    )

    return dataloader


# --- Special words support (prompt biasing) ---------------------------------
def _read_lines(path: str) -> list:
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return [ln.strip() for ln in f.readlines()]
    except Exception:
        return []


def load_awkward_terms(input_path: str) -> list:
    """Load user-provided domain terms from common locations or env.

    Priority:
      1) TRANSCRIBE_AWKWARD_TERMS env (comma-separated)
      2) TRANSCRIBE_AWKWARD_FILE env (path to .txt/.md)
      3) special_words.txt / special_words.md in the input file's folder
      4) special_words.txt / special_words.md in the repo root (this module's dir)
    """
    terms = []
    try:
        # 1) Inline env list
        env_list = os.environ.get('TRANSCRIBE_AWKWARD_TERMS', '')
        if env_list.strip():
            for t in env_list.split(','):
                t = t.strip()
                if t:
                    terms.append(t)

        # 2) Env file
        env_file = os.environ.get('TRANSCRIBE_AWKWARD_FILE', '').strip()
        if env_file and os.path.isfile(env_file):
            terms.extend(_read_lines(env_file))

        # 3) Local folder files
        in_dir = os.path.dirname(input_path)
        for fname in ('special_words.txt', 'special_words.md'):
            p = os.path.join(in_dir, fname)
            if os.path.isfile(p):
                terms.extend(_read_lines(p))
                break

        # 4) Repo root
        repo_dir = os.path.dirname(__file__)
        for fname in ('special_words.txt', 'special_words.md'):
            p = os.path.join(repo_dir, fname)
            if os.path.isfile(p):
                terms.extend(_read_lines(p))
                break
    except Exception:
        pass

    # Normalize simple bullet formats and filter empties/comments
    cleaned = []
    for ln in terms:
        if not ln:
            continue
        s = ln.lstrip('-•*\t >').strip()
        if not s or s.startswith('#'):
            continue
        if s not in cleaned:
            cleaned.append(s)

    # Cap length to keep prompt small and focused
    return cleaned[:40]


def build_initial_prompt(terms: list, max_chars: int = 400) -> Optional[str]:
    """Build a concise initial_prompt string to bias Whisper.

    Keeps capitalization as provided; trims to max_chars.
    """
    if not terms:
        return None
    try:
        # Best practice: provide only a short neutral vocabulary list to reduce prompt leakage
        # Avoid meta-instructions that can be transcribed verbatim.
        payload = '; '.join(terms)
        prompt = payload.strip()
        if len(prompt) > max_chars:
            prompt = prompt[: max_chars - 3].rstrip() + '...'
        return prompt
    except Exception:
        return None

# --- Artifact mitigation: collapse excessive exact repetitions ----------------
def _collapse_repetitions(text: str, max_repeats: int = 3) -> str:
    """Collapse excessive immediate repetitions of the same phrase.

    This targets simple loops like "to grow, to grow, to grow, ..." and reduces
    them to at most `max_repeats` consecutive occurrences.
    """
    try:
        # Check for environment override for max repeats
        env_max_repeats = os.environ.get("TRANSCRIBE_MAX_REPEAT_CAP", "").strip()
        if env_max_repeats:
            try:
                max_repeats = max(1, min(10, int(env_max_repeats)))  # Cap between 1-10
            except Exception:
                pass  # Use default if invalid
        
        # Normalize spaces around commas for matching
        t = re.sub(r"\s*,\s*", ", ", text)
        # Build a regex that captures a short phrase (1-6 words) repeated many times
        # Words may include apostrophes; keep phrases modest to avoid over-collapsing
        pattern = r"\b((?:[A-Za-z']+\s+){0,5}[A-Za-z']+)\b(?:,\s*\1\b){" + str(max_repeats) + ",}"

        def repl(m):
            phrase = m.group(1)
            return ", ".join([phrase] * max_repeats)

        # Apply repeatedly a few times to catch nested patterns
        for _ in range(2):
            new_t = re.sub(pattern, repl, t)
            if new_t == t:
                break
            t = new_t
        return t
    except Exception:
        return text


def _remove_prompt_artifacts(text: str) -> str:
    """Remove prompt text artifacts that sometimes appear in transcriptions.
    
    Whisper can accidentally transcribe the initial_prompt as actual speech,
    especially phrases like "Maintain capitalization" which appear in our prompts.
    This function aggressively removes these artifacts.
    """
    try:
        # Phrases to completely remove (case-insensitive)
        artifact_phrases = [
            r'Maintain capitalization[,\s]*',
            r'maintain capitalization[,\s]*',
            r'Maintain capitalization or overuse these terms[,\.\s]*',
            r'Do not force, repeat, or overuse these terms[,\.\s]*',
            r'otherwise ignore them[,\.\s]*',
        ]
        
        # Remove each artifact phrase globally
        for phrase_pattern in artifact_phrases:
            text = re.sub(phrase_pattern, '', text, flags=re.IGNORECASE)
        
        # Clean up multiple spaces, commas, and periods left behind
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r',\s*,+', ',', text)
        text = re.sub(r'\.\s*\.+', '.', text)
        text = re.sub(r'^\s*[,\.]+\s*', '', text)  # Remove leading punctuation
        
        # Clean up paragraphs that became empty or whitespace-only
        lines = text.split('\n')
        lines = [line.strip() for line in lines if line.strip() and not re.match(r'^[,\.\s]+$', line.strip())]
        text = '\n'.join(lines)
        
        return text.strip()
    except Exception:
        return text


def _collapse_sentence_repetitions(text: str, max_repeats: int = 3) -> str:
    """Collapse full-sentence repetitions like 'It was a lie.' repeated many times.

    Runs after punctuation/capitalization so sentences end with .?!
    Keeps up to `max_repeats` identical consecutive sentences.
    """
    try:
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())
        out = []
        last = None
        count = 0
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            if s == last:
                count += 1
                if count <= max_repeats:
                    out.append(s)
            else:
                last = s
                count = 1
                out.append(s)
        return ' '.join(out)
    except Exception:
        return text


# --- Extended artifact & repetition mitigation (new) -------------------------
_ARTIFACT_DROP_LINE_PATTERNS = [
    r"subtitles by the ",  # YouTube subtitles watermark
    r"copyright .* all rights reserved",  # generic copyright blocks
    r"mooji media ltd",  # known stray source watermark
]

_ARTIFACT_INLINE_PATTERNS = [
    r"repeat or overuse these terms",
    r"as written",
    r"if and only if these domain terms are clearly spoken.*?maintain capitalization",
    r"do not force, repeat, or overuse these terms",
    r"maintain capitalization",
    # Common subtitle/watermark leftovers (case-insensitive, tolerant of punctuation)
    r"\bsubtitles?\s+by\s+the\s+amara\.?org\s+community\b[:\-]?",
    r"\bamara\.?org\s+community\b[:\-]?",
    # Generic recording disclaimers that sometimes leak into transcripts
    r"no part of this recording may be reproduced[^\n\.]*",
]

# Music/hallucination patterns - repetitive vocalisations during music sections
_MUSIC_HALLUCINATION_PATTERNS = [
    # "Oh, oh, oh" or "Oh, Oh, Oh" patterns (common during music)
    r"\bOh,?\s*oh,?\s*oh[,\s]*(?:oh[,\s]*)*",
    # "I'm here" repeated (hallucination during silence/music)
    r"(?:I'm here,?\s*)+I'm here",
    # "Let's go" repeated
    r"(?:Let's go\.?\s*)+(?:Let's go\.?)?",
    # "my God" repetitions
    r"(?:my God,?\s*)+my God",
    # Generic la-la-la vocalisation
    r"(?:la,?\s*)+la\b",
    # "da da da" patterns
    r"(?:da,?\s*)+da\b",
    # Thank you repeated
    r"(?:Thank you\.?\s*)+(?:Thank you\.?)?",
    # End-of-recording hallucinations
    r"\bThank you,?\s*the end\.?",
    r"\bThe end,?\s*I'm sorry[^.]*\.",
    r"\bI\s*'?\s*m going to go\.?",
    r"\bI\s+I\s*'?\s*m\b",  # "I I 'm" stutter patterns
    r"\bI\s+ca\s*n'?t\s+worry[^.]*\.",
    r"(?:I'm sorry[,.]?\s*)+",  # repeated "I'm sorry"
    # "the end I" pattern (common artifact)
    r"\bthe end I\.?\s*",
    r"-\s*the end I\.?\s*",
    r"\binto that-\s*the end I\.?",  # specific pattern from this recording
    # End-of-recording nonsense patterns
    r"\bThe end\.\s*Namajipa[^.]*\.",  # gibberish names
    r"\bThe end\.\s*Let me go\.?",
    r"\bI can just be all right tonight\.?\s*The end\.?",
    r"\bNamajipa\s+jirapare\.?",  # gibberish
    r"\bLet me go\.\s*$",  # trailing "Let me go"
]

def _remove_music_hallucinations(text: str) -> tuple[str, int]:
    """Remove common Whisper hallucinations during music/silence sections.
    
    Returns (cleaned_text, count_of_removals).
    """
    removed_count = 0
    for pattern in _MUSIC_HALLUCINATION_PATTERNS:
        new_text, n = re.subn(pattern, "", text, flags=re.IGNORECASE)
        if n:
            removed_count += n
            text = new_text
    
    # Clean up any double spaces left behind
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s+([,.!?])', r'\1', text)  # Fix space before punctuation
    
    return text.strip(), removed_count

def _remove_extended_artifacts(text: str) -> tuple[str, dict]:
    """Remove broader watermark/prompt/copyright artifact lines.

    Returns (cleaned_text, stats_dict).
    """
    removed_counts = {"patterns": {}, "lines_removed": 0, "inline_removed": {}}
    # Inline removals for soft artifacts
    for pat in _ARTIFACT_INLINE_PATTERNS:
        new_text, n = re.subn(pat, "", text, flags=re.IGNORECASE)
        if n:
            removed_counts["inline_removed"][pat] = n
            text = new_text

    # Also remove watermark-style patterns inline to avoid single-line wipeouts
    for pat in _ARTIFACT_DROP_LINE_PATTERNS:
        new_text, n = re.subn(pat, "", text, flags=re.IGNORECASE)
        if n:
            removed_counts["patterns"][pat] = n
            text = new_text
    return text, removed_counts

def _normalize_sentence(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[\s]+", " ", s)
    s = s.rstrip(".,!?;:")
    return s

def _limit_global_sentence_frequency(text: str, max_global: int = 4) -> tuple[str, dict]:
    """Limit occurrences of identical (normalized) sentences in entire document.
    Returns (cleaned_text, stats).
    """
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    counts = {}
    kept = []
    dropped = 0
    cap = max(1, int(os.environ.get("TRANSCRIBE_GLOBAL_SENTENCE_CAP", max_global)))
    for sent in sentences:
        norm = _normalize_sentence(sent)
        if not norm:
            continue
        prev = counts.get(norm, 0)
        if prev < cap:
            kept.append(sent)
        else:
            dropped += 1
        counts[norm] = prev + 1
    stats = {
        "unique_sentences": len(counts),
        "total_sentences": len(sentences),
        "dropped_sentences": dropped,
        "cap": cap,
        "top_repeated": sorted(((k, v) for k, v in counts.items() if v > 1), key=lambda x: -x[1])[:20],
    }
    return " ".join(kept), stats

def _detect_and_break_loops(text: str, window: int = 12, dup_ratio: float = 0.5) -> tuple[str, dict]:
    """Detect high repetition loops in sliding window and prune repeats beyond first occurrence.
    Returns (cleaned_text, stats).
    """
    win = max(4, int(os.environ.get("TRANSCRIBE_LOOP_WINDOW", window)))
    ratio = float(os.environ.get("TRANSCRIBE_LOOP_DUP_RATIO", dup_ratio))
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    cleaned = []
    loop_events = 0
    i = 0
    while i < len(sentences):
        chunk = sentences[i:i+win]
        norms = [_normalize_sentence(c) for c in chunk if c.strip()]
        if norms:
            most_common = max((norms.count(n) for n in set(norms)))
            if most_common / max(1, len(norms)) >= ratio and most_common > 1:
                # Loop detected: keep first unique order of sentences once
                loop_events += 1
                seen = set()
                for s in chunk:
                    n = _normalize_sentence(s)
                    if n not in seen:
                        cleaned.append(s)
                        seen.add(n)
                i += win
                continue
        if sentences[i].strip():
            cleaned.append(sentences[i])
        i += 1
    stats = {"loop_events": loop_events, "original_sentences": len(sentences), "final_sentences": len(cleaned)}
    return " ".join(cleaned), stats

def _summarize_quality(text: str, extra_stats: dict | None = None) -> dict:
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    norm_counts = {}
    for s in sentences:
        n = _normalize_sentence(s)
        if n:
            norm_counts[n] = norm_counts.get(n, 0) + 1
    summary = {
        "total_sentences": len([s for s in sentences if s.strip()]),
        "unique_sentences": len(norm_counts),
        "top_repeated": sorted(((k, v) for k, v in norm_counts.items() if v > 1), key=lambda x: -x[1])[:25],
    }
    if extra_stats:
        summary.update(extra_stats)
    return summary


def _fix_whisper_artifacts(text: str) -> str:
    """Fix common Whisper transcription artifacts found in analysis.
    
    Based on turbo/large-v3 comparison testing:
    - Removes double periods (. .)
    - Fixes dialogue punctuation inconsistencies
    - Fixes contraction spacing (let 's -> let's)
    - Improves sentence boundary detection
    """
    try:
        # Fix contraction spacing artifacts (very common in Whisper output)
        # Pattern: word + space + apostrophe + letters (e.g., "let 's" -> "let's")
        contraction_fixes = [
            (r"\b(\w+)\s+'s\b", r"\1's"),      # let 's -> let's, it 's -> it's
            (r"\b(\w+)\s+'d\b", r"\1'd"),      # I 'd -> I'd, we 'd -> we'd
            (r"\b(\w+)\s+'ll\b", r"\1'll"),    # I 'll -> I'll, we 'll -> we'll
            (r"\b(\w+)\s+'ve\b", r"\1've"),    # I 've -> I've, we 've -> we've
            (r"\b(\w+)\s+'re\b", r"\1're"),    # we 're -> we're, you 're -> you're
            (r"\b(\w+)\s+'m\b", r"\1'm"),      # I 'm -> I'm
            (r"\b(\w+)\s+n't\b", r"\1n't"),    # do n't -> don't, ca n't -> can't
            (r"\bI\s+'m\b", r"I'm"),           # Special case for I 'm
            (r"\bI\s+'d\b", r"I'd"),           # Special case for I 'd
            (r"\bI\s+'ll\b", r"I'll"),         # Special case for I 'll
            (r"\bI\s+'ve\b", r"I've"),         # Special case for I 've
        ]
        for pattern, replacement in contraction_fixes:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Fix double periods (common artifact in both models)
        text = re.sub(r'\.\s*\.+', '.', text)
        
        # Fix period-period spacing artifacts
        text = re.sub(r'\.\s+\.', '.', text)
        
        # Fix question mark patterns in quoted speech
        # "what am I doing now?" instead of what am I doing now?
        text = re.sub(r'\b(what|who|when|where|why|how)\s+([^?.!]+)\?', r'\1 \2?', text, flags=re.IGNORECASE)
        
        # Fix colon + lowercase after direct quote intro (should be lowercase for continuation)
        # He said: well, what -> He said: well, what (keep lowercase)
        # But: Right now I'm -> Right now I'm (keep capitals when appropriate)
        
        # Fix common conjunction drops at sentence starts
        # "All this energy" after period should check context
        # This is complex - leave for manual review for now
        
        # Clean up multiple spaces
        text = re.sub(r'  +', ' ', text)
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([,.!?;:])', r'\1', text)
        text = re.sub(r'([,.!?;:])\s+([,.!?;:])', r'\1 \2', text)
        
        return text.strip()
    except Exception as e:
        print(f"⚠️  Artifact fixing failed: {e}")
        return text


def _clean_repetitions_in_segment(text: str, max_phrase_repeats: int = 2) -> str:
    """Light, in-transcription de-repetition applied per segment.

    - Collapses immediate repeats of short phrases (1–5 words) within a single segment
    - Keeps at most `max_phrase_repeats` consecutive occurrences
    - Intended to run before any post-processing, preserving 'verbatim' intent while
      removing obvious decode loops that occur within one segment.
    """
    try:
        t = re.sub(r"\s*,\s*", ", ", text)
        pattern = r"\b((?:[A-Za-z']+\s+){0,4}[A-Za-z']+)\b(?:,?\s+\1\b){" + str(max_phrase_repeats) + ",}"

        def repl(m):
            phrase = m.group(1)
            return (phrase + ", ") * (max_phrase_repeats - 1) + phrase

        for _ in range(2):
            new_t = re.sub(pattern, repl, t, flags=re.IGNORECASE)
            if new_t == t:
                break
            t = new_t
        return t
    except Exception:
        return text


def _segments_to_paragraphs(segments: list, gap_threshold: float = 1.2) -> str:
    """Build coherent paragraphs from Whisper segments without altering words.

    Rules:
      - Start a new paragraph when the gap between segments >= gap_threshold seconds
      - Also break when a segment ends with terminal punctuation and the next segment starts a new sentence
      - Preserve original segment text (except for in-segment repetition cleanup)
    """
    paras: list[str] = []
    curr: list[str] = []
    prev_end = None
    for seg in segments:
        txt = str(seg.get("text", "")).strip()
        if not txt:
            continue
        # light per-segment repetition cleanup
        txt = _clean_repetitions_in_segment(txt)
        start = float(seg.get("start", 0.0) or 0.0)
        end = float(seg.get("end", 0.0) or 0.0)

        new_para = False
        if prev_end is not None and start - prev_end >= gap_threshold:
            new_para = True
        elif curr:
            # if previous chunk ends with .!? and current begins with capital letter
            if curr[-1][-1:] in ".!?":
                if txt and txt[0].isupper():
                    new_para = True

        if new_para and curr:
            paras.append(" ".join(curr).strip())
            curr = []

        curr.append(txt)
        prev_end = end

    if curr:
        paras.append(" ".join(curr).strip())
    return "\n\n".join(paras)


def _refine_capitalization(text: str) -> str:
    """Fix capitalization artifacts without changing word content.
    
    Fixes common Whisper artifacts:
    - Incorrectly capitalized words mid-sentence
    - Capital letters immediately after commas (not sentence starts)
    - Preserves proper nouns and acronyms
    """
    try:
        # Split into sentences while preserving structure
        sentences = re.split(r'([.!?]+\s+)', text)
        
        refined_sentences = []
        for i, part in enumerate(sentences):
            # Skip sentence delimiters
            if re.match(r'^[.!?]+\s+$', part):
                refined_sentences.append(part)
                continue
            
            # Process each sentence
            if part.strip():
                # Fix: Capital letter after comma mid-sentence
                # Pattern: ", Word" -> ", word" (unless it's a proper noun)
                part = re.sub(
                    r',\s+([A-Z])([a-z]+)',
                    lambda m: f', {m.group(1).lower()}{m.group(2)}' 
                    if m.group(1) + m.group(2) not in ['I', 'The', 'A', 'An'] 
                    else m.group(0),
                    part
                )
                
                # Fix: Mid-sentence capitalization not following punctuation
                # Split on spaces to check each word
                words = part.split()
                if words:
                    # First word of sentence should be capitalized
                    if words[0] and words[0][0].islower():
                        words[0] = words[0][0].upper() + words[0][1:]
                    
                    # Check remaining words
                    for j in range(1, len(words)):
                        word = words[j]
                        if not word:
                            continue
                        
                        # Check if previous word ended with sentence-ending punctuation
                        prev_ends_sentence = j > 0 and words[j-1] and words[j-1][-1] in '.!?'
                        
                        # If word is capitalized but not after punctuation
                        if word[0].isupper() and not prev_ends_sentence:
                            # Check if it's likely a proper noun (all caps, or starts uppercase and has uppercase later)
                            is_acronym = word.isupper() and len(word) > 1
                            has_internal_caps = len(word) > 1 and any(c.isupper() for c in word[1:])
                            
                            # Preserve known proper nouns and acronyms
                            if not (is_acronym or has_internal_caps or word in ['I']):
                                # Lowercase the first character
                                words[j] = word[0].lower() + word[1:]
                    
                    part = ' '.join(words)
                
                refined_sentences.append(part)
        
        return ''.join(refined_sentences)
    except Exception as e:
        print(f"⚠️  Capitalization refinement failed: {e}")
        return text


def transcribe_with_dataset_optimization(input_path: str, output_dir=None, threads_override: Optional[int] = None):
    """
    Transcribe audio using dataset-based GPU pipeline optimization.

    This function implements the GPU efficiency improvements suggested by PyTorch:
    - Uses PyTorch Dataset for batch processing
    - Leverages DataLoader for optimized data loading
    - Implements overlapping segments for better context
    - Utilizes GPU pipelines for maximum efficiency
    """
    _ensure_torch_available()
    torch_api = cast(Any, torch)

    # Lazy imports
    import whisper
    from docx import Document
    from deepmultilingualpunctuation import PunctuationModel
    from transcribe import (
        get_media_duration, split_into_paragraphs, format_duration, format_duration_minutes_only, format_duration_hms,
    )

    start_time = time.time()
    print("\n" + "="*80)
    print("🚀 GPU-ACCELERATED TRANSCRIPTION")
    print("="*80)
    print(f"📁 Input: {os.path.basename(input_path)}")

    # Preprocess audio with silence padding to prevent missed words
    preprocessed_path = preprocess_audio_with_padding(input_path)
    preprocessing_used = preprocessed_path != input_path
    
    if preprocessing_used:
        print(f"✅ Preprocessed with silence padding (prevents missed words)")
        # Use the preprocessed file for all subsequent operations
        working_input_path = preprocessed_path
    else:
        working_input_path = input_path

    # Get hardware config
    max_perf = os.environ.get("TRANSCRIBE_MAX_PERF", "").strip() in ("1", "true", "True")
    config = get_maximum_hardware_config(max_perf=max_perf)

    if not output_dir:
        output_dir = os.path.join(os.path.expanduser("~"), "Downloads")

    duration = get_media_duration(working_input_path)
    if duration:
        print(f"⏱️  Duration: {format_duration(duration)}")

    # Pre-run cleanup
    force_gpu_memory_cleanup()

    # Load model
    device_name = "CPU"
    model = None
    chosen_device = "cpu"
    
    # Check for model selection from environment variable (set by GUI)
    selected_model_name = os.environ.get("TRANSCRIBE_MODEL_NAME", "large-v3")
    # Respect user choice exactly; no automatic model size downgrades.
    # We will attempt smarter GPU loading (FP32 then FP16 fallback) before CPU fallback.

    try:
        avail = set(whisper.available_models())
        requested_available = (selected_model_name in avail)
        if not requested_available:
            print(f"⚠️  Requested model '{selected_model_name}' not available, falling back...")
            for cand in ("large-v3", "large-v2", "large"):
                if cand in avail:
                    selected_model_name = cand
                    break
        print(f"🎯 Model: {selected_model_name}")
    except Exception as e:
        print(f"⚠️  Could not query whisper.available_models(): {e}")

    # Load model on best available device
    try:
        if "cuda" in config["devices"] and torch_api.cuda.is_available():
            chosen_device = "cuda"
            device_name = f"CUDA GPU ({torch_api.cuda.get_device_name(0)})"
            print(f"🎯 Device: {device_name}")

            # Enable GPU optimizations
            if hasattr(torch_api.backends, "cudnn"):
                torch_api.backends.cudnn.benchmark = True
            if hasattr(torch_api.backends, "cuda") and hasattr(torch_api.backends.cuda, "matmul"):
                try:
                    torch_api.backends.cudnn.matmul.allow_tf32 = True
                except Exception:
                    pass
            try:
                torch_api.set_float32_matmul_precision("high")
            except Exception:
                pass

            # Load model for GPU
            model = whisper.load_model(selected_model_name, device="cuda")
            # Note: FP16 conversion currently disabled due to dtype compatibility issues
            # Will use FP32 for stability
            model_is_fp16 = False
            print(f"✅ Model loaded in FP32 (stable for parallel processing)")
        else:
            chosen_device = "cpu"
            device_name = f"CPU ({multiprocessing.cpu_count()} cores)"
            print(f"🎯 Device: {device_name}")
            model = whisper.load_model(selected_model_name, device="cpu")
            model_is_fp16 = False  # CPU doesn't use FP16
    except Exception as e:
        print(f"❌ Model load failed: {e}")
        raise

    # Set CPU threads
    if isinstance(threads_override, int) and threads_override > 0:
        config["cpu_threads"] = max(1, min(64, threads_override))

    torch_api.set_num_threads(config["cpu_threads"])
    interop = max(2, min(16, config["cpu_threads"] // 4))
    try:
        torch_api.set_num_interop_threads(interop)
    except Exception:
        pass

    print(f"⚙️  Threads: {config['cpu_threads']} CPU, {interop} interop")

    # Build optional initial_prompt from special terms
    awkward_terms = load_awkward_terms(input_path)
    initial_prompt = build_initial_prompt(awkward_terms)

    # Create dataset and dataloader for efficient processing
    try:
        # Use batch_size=1 in DataLoader since we handle parallelism with ThreadPoolExecutor
        # This avoids tensor batching complications
        batch_size_dataloader = 1  # DataLoader batching disabled - we do parallel processing manually
        # Increase worker threads for better data pipeline
        num_workers = min(4, config["cpu_threads"] // 2) if config["cpu_threads"] > 4 else 0

        dataloader = create_efficient_dataloader(
            working_input_path,
            batch_size=batch_size_dataloader,
            num_workers=num_workers
        )

        print(f"📊 Dataset: {len(dataloader.dataset)} segments")
        print("")  # Blank line for readability

    except Exception as e:
        print(f"⚠️  Dataset creation failed: {e} - falling back to standard processing")
        # Cleanup preprocessed file before fallback
        try:
            if preprocessing_used and os.path.exists(working_input_path):
                os.remove(working_input_path)
                print("🧹 Removed temporary preprocessed audio file")
        except Exception:
            pass
        return transcribe_file_simple_auto(input_path, output_dir, threads_override=threads_override)

    # Process segments with dataset optimization
    all_segments = []
    segment_count = 0

    print("🔄 Transcribing audio...")
    print("─" * 80)
    
    # Configure transcription parameters once
    seg_kwargs = dict(
        language="en",  # Optimized for English language
        compression_ratio_threshold=2.4,
        logprob_threshold=-1.0,
        no_speech_threshold=0.3,
        # Disable conditioning to prevent repetition loops between segments
        condition_on_previous_text=False,
        temperature=0.0,
        verbose=False,  # Reduce verbosity for batch processing
    )
    
    # Model-specific tuning for accuracy
    if selected_model_name == "large-v3-turbo":
        # Turbo-specific: tighter thresholds for better accuracy
        seg_kwargs["compression_ratio_threshold"] = 2.2
        seg_kwargs["logprob_threshold"] = -1.5
        seg_kwargs["no_speech_threshold"] = 0.4
    elif selected_model_name == "large-v3":
        # Large-v3: slightly more permissive for natural speech flow
        seg_kwargs["compression_ratio_threshold"] = 2.6
        seg_kwargs["logprob_threshold"] = -2.2
    
    # Apply quality mode if enabled
    quality_mode = os.environ.get("TRANSCRIBE_QUALITY_MODE", "").strip() in ("1", "true", "True")
    if quality_mode:
        if selected_model_name == "large-v3-turbo":
            seg_kwargs["beam_size"] = 10
            seg_kwargs["patience"] = 3.0
            seg_kwargs["best_of"] = 10
            seg_kwargs["temperature"] = 0.0
        else:
            seg_kwargs["beam_size"] = 10
            seg_kwargs["patience"] = 3.0
            seg_kwargs["best_of"] = 10
            seg_kwargs["temperature"] = 0.0
    
    # Only allow domain bias when explicitly enabled
    if initial_prompt and str(os.environ.get("TRANSCRIBE_ALLOW_PROMPT", "0")).lower() in ("1","true","yes"):
        seg_kwargs["initial_prompt"] = initial_prompt
    
    # Add FP16 flag if model was converted to half precision
    if chosen_device == "cuda" and model_is_fp16:
        seg_kwargs["fp16"] = True

    # Process in true batches for better GPU utilization
    if chosen_device == "cuda":
        import concurrent.futures
        
        # Collect all segments first for parallel processing
        all_batch_segments = []
        for batch in dataloader:
            # With batch_size=1, batch is a dict with single items (possibly tensors)
            # Extract the values - they might be tensors or lists of length 1
            audio = batch['audio']
            start_time = batch['start_time']
            end_time = batch['end_time']
            
            # Handle tensor/list wrapping from DataLoader
            if hasattr(audio, '__getitem__') and not isinstance(audio, np.ndarray):
                # It's a batched tensor or list - extract first element
                audio = audio[0] if len(audio) > 0 else audio
                start_time = start_time[0] if hasattr(start_time, '__getitem__') else start_time
                end_time = end_time[0] if hasattr(end_time, '__getitem__') else end_time
            
            # Convert to numpy if needed
            segment_audio = audio.numpy() if hasattr(audio, 'numpy') else audio
            
            all_batch_segments.append({
                'audio': segment_audio,
                'start_time': float(start_time),
                'end_time': float(end_time)
            })
        
        # Process segments with thread-safe model access
        
        # Create a lock to prevent concurrent model access (prevents tensor dimension mismatches)
        model_lock = threading.Lock()
        
        # Process multiple segments in parallel using ThreadPoolExecutor
        # Lock ensures model.transcribe() is called sequentially while still benefiting from
        # async I/O, preprocessing, and postprocessing parallelization
        max_workers = min(4, len(all_batch_segments))  # Up to 4 parallel transcriptions
        
        def process_segment(seg_data):
            """Process a single segment and return adjusted timestamps"""
            try:
                # Use lock to ensure only one thread accesses the model at a time
                # This prevents race conditions in the attention mechanism
                with model_lock:
                    # Ensure float32 dtype to avoid float/double mismatches in PyTorch
                    audio_arr = seg_data['audio']
                    try:
                        import numpy as _np
                    except Exception:
                        _np = None
                    if _np is not None and isinstance(audio_arr, _np.ndarray) and audio_arr.dtype != _np.float32:
                        audio_arr = audio_arr.astype(_np.float32)
                    result = _compatible_transcribe_call(model, audio_arr, seg_kwargs)
                    result = _as_result_dict(result)
                
                processed_segs = []
                
                if isinstance(result, dict) and "segments" in result:
                    for seg in result["segments"]:
                        seg_copy = dict(seg)
                        seg_copy["start"] = seg_data['start_time'] + seg.get("start", 0)
                        seg_copy["end"] = seg_data['start_time'] + seg.get("end", 0)
                        processed_segs.append(seg_copy)
                        
                        # Print transcribed text as we go (shows progress)
                        # Filter out prompt text artifacts and repetitive hallucinations
                        text = seg.get("text", "").strip()
                        
                        # Skip prompt artifacts and obvious hallucinations
                        skip_phrases = [
                            "Maintain capitalization",
                            "maintain capitalization", 
                            "Maintain capitalization or overuse these terms"
                        ]
                        
                        if text:
                            # Check if text is mostly repetitive prompt garbage
                            is_prompt_artifact = any(phrase in text for phrase in skip_phrases)
                            
                            # Don't print if it's a prompt artifact
                            if not is_prompt_artifact:
                                print(f"   {text}")
                
                return processed_segs
            except Exception as e:
                print(f"⚠️  Segment processing error: {e}")
                return []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(process_segment, seg) for seg in all_batch_segments]
            
            completed_count = 0
            for future in concurrent.futures.as_completed(futures):
                try:
                    segments = future.result()
                    all_segments.extend(segments)
                    completed_count += 1
                    
                    if completed_count % 10 == 0:
                        print(f"📊 Completed {completed_count}/{len(all_batch_segments)} segments (parallel - will sort by timestamp)...")
                except Exception as e:
                    print(f"⚠️  Future completion error: {e}")
            
            segment_count = len(all_batch_segments)
            print(f"\n✅ Transcription complete: {completed_count} segments")
    
    else:
        # CPU fallback: sequential processing
        print("🔄 Using sequential processing for CPU")
        for batch in dataloader:
            try:
                for segment_data in batch:
                    segment_audio = segment_data['audio'].numpy() if hasattr(segment_data['audio'], 'numpy') else segment_data['audio']
                    try:
                        import numpy as _np
                    except Exception:
                        _np = None
                    if _np is not None and isinstance(segment_audio, _np.ndarray) and segment_audio.dtype != _np.float32:
                        segment_audio = segment_audio.astype(_np.float32)
                    start_time_seg = segment_data['start_time']
                    result = _compatible_transcribe_call(model, segment_audio, seg_kwargs)
                    result = _as_result_dict(result)

                    if isinstance(result, dict) and "segments" in result:
                        for seg in result["segments"]:
                            seg_copy = dict(seg)
                            seg_copy["start"] = start_time_seg + seg.get("start", 0)
                            seg_copy["end"] = start_time_seg + seg.get("end", 0)
                            all_segments.append(seg_copy)

                    segment_count += 1
                    if segment_count % 10 == 0:
                        print(f"📊 Processed {segment_count} segments...")

            except Exception as e:
                print(f"⚠️  Batch processing error: {e} - continuing with next batch")

    print("─" * 80)

    # ROBUST SORTING: Critical for parallel processing where segments may complete out of order
    # Sort by start time (primary) and end time (secondary) to ensure correct chronological order
    if all_segments:
        # Multi-key sort: start time first, then end time for segments starting at same time
        all_segments.sort(key=lambda x: (x.get("start", 0), x.get("end", 0)))
        
        # Validation: Check for overlaps or out-of-order segments
        for i in range(1, len(all_segments)):
            prev_end = all_segments[i-1].get("end", 0)
            curr_start = all_segments[i].get("start", 0)
            
            # Log warning if segments are significantly out of order (gap > 5 seconds backwards)
            if curr_start < prev_end - 5.0:
                print(f"⚠️  Warning: Segment {i} may overlap (prev ends at {prev_end:.2f}s, current starts at {curr_start:.2f}s)")

    # Combine results
    full_text = ""
    if all_segments:

        # Extract and combine text
        texts = []
        for seg in all_segments:
            text = seg.get("text", "").strip()
            if text:
                texts.append(text)

        full_text = " ".join(texts).strip()

    if not full_text:
        print("⚠️  Warning: No transcription text generated")
        full_text = "[No speech detected or transcription failed]"

    print(f"⚡ Hardware utilised: {device_name} (Dataset Optimized)")

    # Post-processing (same as original)
    try:
        if _is_verbatim_mode():
            # Verbatim: keep Whisper output as-is (minimal trim only)
            quality_stats = {"verbatim": True}
            print("🧷 Verbatim mode: skipping punctuation/capitalization/repetition/artifact passes")
        else:
            # Enhanced pipeline
            full_text = _collapse_repetitions(full_text, max_repeats=3)
            full_text = _remove_prompt_artifacts(full_text)
            full_text, early_artifact_stats = _remove_extended_artifacts(full_text)
            pm = PunctuationModel()
            t0 = time.time()
            full_text = pm.restore_punctuation(full_text)
            t1 = time.time()
            full_text = pm.restore_punctuation(full_text)
            t2 = time.time()
            print(f"✅ Punctuation restoration completed (passes: 2 | {t1 - t0:.1f}s + {t2 - t1:.1f}s)")
            full_text = _fix_whisper_artifacts(full_text)
            full_text = _refine_capitalization(full_text)
            full_text = _collapse_sentence_repetitions(full_text, max_repeats=3)
            full_text, global_freq_stats = _limit_global_sentence_frequency(full_text)
            full_text, loop_stats = _detect_and_break_loops(full_text)
            full_text, late_artifact_stats = _remove_extended_artifacts(full_text)
            quality_stats = {
                "early_artifacts": early_artifact_stats,
                "global_frequency": global_freq_stats,
                "loop_detection": loop_stats,
                "late_artifacts": late_artifact_stats,
            }
            print("✅ Capitalization & artifact refinement completed")
    except Exception as e:
        print(f"⚠️  Post-processing failed: {e}")
        quality_stats = {}

    try:
        if _is_verbatim_mode():
            formatted_text = full_text
            print("✅ Verbatim formatting: preserved original model text")
        else:
            formatted = split_into_paragraphs(full_text, max_length=500)
            if isinstance(formatted, list):
                formatted_text = "\n\n".join(formatted)
            else:
                formatted_text = full_text
            print("✅ Text formatting completed")
    except Exception as e:
        print(f"⚠️  Text formatting failed: {e}")
        formatted_text = full_text

    # Save files (same as original)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    txt_path = os.path.join(output_dir, f"{base_name}.txt")
    docx_path = os.path.join(output_dir, f"{base_name}.docx")

    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(formatted_text)
        print(f"✅ Text file saved: {txt_path}")
    except Exception as e:
        print(f"❌ Failed to save text file: {e}")
        txt_path = None

    try:
        doc = Document()
        doc.add_heading(f'{base_name}', 0)
        
        # Add model and location info
        parent_folder = os.path.basename(os.path.dirname(input_path))
        doc.add_paragraph(f'Model: {selected_model_name}')
        doc.add_paragraph(f'Folder: {parent_folder}')
        doc.add_paragraph('')

        elapsed_total = time.time() - start_time
        if os.environ.get("TRANSCRIBE_HIDE_TIME", "").lower() not in ("1", "true", "yes"):
            doc.add_paragraph(f'Transcription time: {format_duration_hms(elapsed_total)}')
            doc.add_paragraph('')

        if _is_verbatim_mode():
            # Preserve raw text as a single block to avoid reflow changes
            doc.add_paragraph(formatted_text)
        else:
            for para in formatted_text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(docx_path)
        print(f"✅ Word document saved: {docx_path}")
    except Exception as e:
        print(f"❌ Failed to create Word document: {e}")
        try:
            doc = Document()
            doc.add_heading(f'Transcription: {base_name}', 0)
            
            # Add model and location info (fallback)
            parent_folder = os.path.basename(os.path.dirname(input_path))
            doc.add_paragraph(f'Model: {selected_model_name}')
            doc.add_paragraph(f'Folder: {parent_folder}')
            doc.add_paragraph('')
            
            doc.add_paragraph(formatted_text[:5000])
            doc.save(docx_path)
            print(f"✅ Basic Word document saved: {docx_path}")
        except Exception as e2:
            print(f"❌ Failed to save even basic Word document: {e2}")
            docx_path = None

    # Final stats
    elapsed = time.time() - start_time
    print("\n🎉 DATASET-OPTIMIZED TRANSCRIPTION COMPLETE!")
    print(f"📄 Text file: {txt_path}")
    print(f"📄 Word document: {docx_path}")
    print(f"⏱️  Total time: {format_duration(elapsed)}")

    # Cleanup
    force_gpu_memory_cleanup()

    import psutil as _ps
    mem = _ps.virtual_memory()
    if torch_api.cuda.is_available():
        try:
            gpu_after = torch_api.cuda.memory_allocated() / (1024 ** 3)
            print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available, GPU {gpu_after:.1f}GB used")
        except Exception:
            print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available")
    else:
        print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available")

    # Cleanup preprocessed file on completion
    try:
        if preprocessing_used and os.path.exists(working_input_path):
            os.remove(working_input_path)
            print("🧹 Removed temporary preprocessed audio file")
    except Exception:
        pass

    return txt_path


def get_maximum_hardware_config(max_perf: bool = False):
    """Detect hardware and return a conservative, stable config dict."""
    _ensure_torch_available()
    torch_api = cast(Any, torch)
    cpu_cores = max(multiprocessing.cpu_count(), 1)
    vm = psutil.virtual_memory()
    total_ram_gb = vm.total / (1024 ** 3)
    available_ram_gb = vm.available / (1024 ** 3)
    # Default: plan to use 98% of currently available RAM (ULTRA OPTIMISED)
    usable_ram_gb = max(available_ram_gb * 0.98, 1.0)
    # RAM overrides via env: prefer absolute GB then fraction
    try:
        env_ram_gb = float(os.environ.get("TRANSCRIBE_RAM_GB", "") or 0)
    except Exception:
        env_ram_gb = 0.0
    try:
        env_ram_frac = float(os.environ.get("TRANSCRIBE_RAM_FRACTION", "") or 0)
    except Exception:
        env_ram_frac = 0.0
    if env_ram_gb > 0:
        # Absolute cap (do not exceed physical total)
        usable_ram_gb = max(1.0, min(total_ram_gb, env_ram_gb))
    elif 0.05 <= env_ram_frac <= 1.0:
        # Fraction of currently available RAM
        usable_ram_gb = max(1.0, available_ram_gb * env_ram_frac)

    devices = ["cpu"]
    has_cuda = False
    cuda_total_vram_gb = 0.0
    try:
        has_cuda = torch_api.cuda.is_available()
        if has_cuda:
            devices.insert(0, "cuda")
            try:
                cuda_total_vram_gb = torch_api.cuda.get_device_properties(0).total_memory / (1024 ** 3)
            except Exception:
                cuda_total_vram_gb = 0.0
    except Exception:
        has_cuda = False

    # DirectML availability (optional)
    dml_available = False
    try:
        import torch_directml  # type: ignore
        _ = torch_directml.device()
        dml_available = True
        if not has_cuda:
            devices.insert(0, "dml")
    except Exception:
        dml_available = False

    # Threads: target ~90% by default; in max_perf mode, use 100% of logical cores
    import math
    if max_perf:
        cpu_threads = max(1, min(64, cpu_cores))  # Use ALL cores in max perf mode
    else:
        cpu_threads = max(1, min(64, math.ceil(cpu_cores * 0.95)))  # Increased from 90% to 95%

    # Environment override for threads
    try:
        env_threads = int(os.environ.get("TRANSCRIBE_THREADS", "") or 0)
    except Exception:
        env_threads = 0
    if env_threads > 0:
        cpu_threads = max(1, min(64, env_threads))

    # For ULTRA optimised utilization, use more GPU workers for maximum utilization
    if has_cuda:
        try:
            gpu_count = torch_api.cuda.device_count()
            # Use more workers to take advantage of additional GPU shared memory
            gpu_workers = min(gpu_count * 3, 8)  # Increased from 6 to 8 max, 3x GPU count
        except Exception:
            gpu_workers = 3  # Increased fallback from 2 to 3
    else:
        gpu_workers = 0

    # VRAM overrides via env
    try:
        env_vram_gb = float(os.environ.get("TRANSCRIBE_VRAM_GB", "") or 0)
    except Exception:
        env_vram_gb = 0.0
    try:
        env_vram_frac = float(os.environ.get("TRANSCRIBE_VRAM_FRACTION", "") or 0)
    except Exception:
        env_vram_frac = 0.0
    allowed_vram_gb = cuda_total_vram_gb
    if cuda_total_vram_gb > 0:
        if env_vram_gb > 0:
            allowed_vram_gb = max(0.5, min(cuda_total_vram_gb, env_vram_gb))
        elif 0.05 <= env_vram_frac <= 1.0:
            allowed_vram_gb = max(0.5, cuda_total_vram_gb * env_vram_frac)

    cfg = {
        "cpu_cores": cpu_cores,
        "cpu_threads": cpu_threads,
        "total_ram_gb": total_ram_gb,
    "available_ram_gb": available_ram_gb,
        "usable_ram_gb": usable_ram_gb,
        "devices": devices,
        "gpu_workers": gpu_workers,
        "total_workers": max(cpu_threads, gpu_workers),
        "dml_available": dml_available,
        "cuda_total_vram_gb": cuda_total_vram_gb,
        "allowed_vram_gb": allowed_vram_gb,
        "max_perf": bool(max_perf),
    }
    return cfg


def adjust_workers_for_model(config, model_name):
    """Optionally tweak worker counts based on model size. Keep it conservative."""
    cfg = dict(config)
    name = (model_name or "large").lower()
    # Larger models -> fewer CPU threads to reduce contention
    if name in ("large", "medium"):
        cfg["cpu_threads"] = min(cfg.get("cpu_threads", 8), 12)
    cfg["total_workers"] = max(cfg.get("cpu_threads", 1), cfg.get("gpu_workers", 0))
    return cfg


def force_gpu_memory_cleanup():
    """Clear GPU caches and model-related module caches without touching torch modules."""
    try:
        _ensure_torch_available()
        torch_api = cast(Any, torch)
        if torch_api.cuda.is_available():
            torch_api.cuda.empty_cache()
            try:
                torch_api.cuda.synchronize()
            except Exception:
                pass
    except Exception as e:
        print(f"⚠️  GPU cache clear warning: {e}")

    # Clear whisper/transformers modules only (keeps torch intact)
    to_clear = [
        name for name in list(sys.modules.keys())
        if name.startswith(("whisper", "transformers"))
    ]
    for name in to_clear:
        try:
            del sys.modules[name]
        except Exception:
            pass

    gc.collect()


def vad_segment_times_optimized(input_path, aggressiveness=2, frame_duration_ms=30, padding_ms=300):
    """Voice Activity Detection optimized for maximum performance with fallback."""
    # Check VAD availability within this function scope
    try:
        import webrtcvad
        _vad_available = True
    except ImportError:
        _vad_available = False
        webrtcvad = None

    if not _vad_available:
        print("⚠️  webrtcvad not available - using optimized duration-based segmentation")
        # Fallback: create segments based on duration (every 25 seconds for better performance)
        try:
            from moviepy.editor import AudioFileClip
            audio_clip = AudioFileClip(input_path)
            duration = audio_clip.duration
            audio_clip.close()

            segments = []
            segment_length = 25.0  # 25 second segments for optimized processing
            for i in range(0, int(duration), int(segment_length)):
                start = float(i)
                end = min(float(i + segment_length), duration)
                segments.append((start, end))

            print(f"📊 Created {len(segments)} optimized duration-based segments ({segment_length}s each)")
            return segments

        except Exception as e:
            print(f"❌ Error creating optimized fallback segments: {e}")
            # Last resort: single segment for entire audio
            return [(0.0, 60.0)]  # Assume 60s max, will be clipped later

    # Original webrtcvad implementation with optimized settings
    try:
        # Import required functions from transcribe.py
        from transcribe import get_pcm_from_file, frames_from_pcm

        pcm = get_pcm_from_file(input_path)
        vad = webrtcvad.Vad(aggressiveness)
        frames = list(frames_from_pcm(pcm, frame_duration_ms=frame_duration_ms))
        sample_rate = 16000
        in_speech = False
        segments = []
        speech_start = 0

        for i, frame in enumerate(frames):
            is_speech = False
            if len(frame) == int(sample_rate * 2 * (frame_duration_ms/1000.0)):
                is_speech = vad.is_speech(frame, sample_rate)
            t = (i * frame_duration_ms) / 1000.0
            if is_speech and not in_speech:
                in_speech = True
                speech_start = t
            elif not is_speech and in_speech:
                in_speech = False
                speech_end = t
                # Optimized padding for better performance
                start = max(0, speech_start - (padding_ms/1000.0))
                end = speech_end + (padding_ms/1000.0)
                segments.append((start, end))

        # Handle file ending while in speech
        if in_speech:
            speech_end = (len(frames) * frame_duration_ms) / 1000.0
            start = max(0, speech_start - (padding_ms/1000.0))
            end = speech_end + (padding_ms/1000.0)
            segments.append((start, end))

        return segments

    except Exception as e:
        print(f"❌ VAD segmentation failed: {e}")
        # Fallback to duration-based segmentation
        try:
            from moviepy.editor import AudioFileClip
            audio_clip = AudioFileClip(input_path)
            duration = audio_clip.duration
            audio_clip.close()

            segments = []
            segment_length = 25.0
            for i in range(0, int(duration), int(segment_length)):
                start = float(i)
                end = min(float(i + segment_length), duration)
                segments.append((start, end))

            print(f"📊 VAD failed, using {len(segments)} duration-based segments")
            return segments

        except Exception as fallback_e:
            print(f"❌ Fallback segmentation also failed: {fallback_e}")
            return [(0.0, 60.0)]


def transcribe_with_vad_parallel(input_path, vad_segments, model, base_transcribe_kwargs, config):
    """Transcribe audio using VAD segments processed in parallel for maximum performance."""
    import concurrent.futures
    import tempfile
    import subprocess

    print(f"🔄 Processing {len(vad_segments)} VAD segments in parallel...")

    # Try to import moviepy
    try:
        from moviepy.editor import AudioFileClip  # type: ignore
        moviepy_available = True
    except ImportError:
        moviepy_available = False
        AudioFileClip = None  # type: ignore
        print("⚠️  moviepy not available - falling back to ffmpeg for segment extraction")

    # Create temporary directory for segment files
    with tempfile.TemporaryDirectory() as temp_dir:
        segment_files = []
        segment_results = []

        # Extract audio segments
        def extract_segment(segment_idx, start_time, end_time):
            try:
                segment_path = os.path.join(temp_dir, f"segment_{segment_idx:03d}.wav")

                if moviepy_available and AudioFileClip is not None:
                    audio_clip = AudioFileClip(input_path)
                    segment_clip = audio_clip.subclip(start_time, end_time)
                    segment_clip.write_audiofile(segment_path, verbose=False, logger=None)
                    audio_clip.close()
                    segment_clip.close()
                else:
                    # Fallback to ffmpeg
                    duration = end_time - start_time
                    cmd = [
                        "ffmpeg", "-i", input_path,
                        "-ss", str(start_time), "-t", str(duration),
                        "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
                        "-y", segment_path
                    ]
                    subprocess.run(cmd, check=True, capture_output=True)

                return segment_path, (start_time, end_time)
            except Exception as e:
                print(f"⚠️  Failed to extract segment {segment_idx}: {e}")
                return None, None

        # Extract all segments SEQUENTIALLY for perfect temporal order
        print("🔧 Using sequential segment extraction for guaranteed order...")
        
        segment_files = []
        for i, (start, end) in enumerate(vad_segments):
            print(f"🎵 Extracting segment {i + 1}/{len(vad_segments)}: {start:.1f}s-{end:.1f}s")
            segment_path, time_range = extract_segment(i, start, end)
            if segment_path and time_range:
                segment_files.append((segment_path, time_range))
                print(f"✅ Extracted segment {i + 1}: {time_range[0]:.1f}s-{time_range[1]:.1f}s")
            else:
                print(f"⚠️  Failed to extract segment {i + 1}")
        
        print(f"✅ Extracted {len(segment_files)} audio segments in perfect temporal order")

        # Transcribe segments in parallel
        def transcribe_segment(segment_path, time_range):
            try:
                # Create a copy of transcribe kwargs for this segment
                segment_kwargs = base_transcribe_kwargs.copy()
                # Remove vad_filter since we're already using segmented audio
                segment_kwargs.pop("vad_filter", None)

                result = _compatible_transcribe_call(model, segment_path, segment_kwargs)
                result = _as_result_dict(result)

                # Add timing information to segments
                if isinstance(result, dict) and "segments" in result:
                    for segment in result["segments"]:
                        segment["start"] += time_range[0]
                        segment["end"] += time_range[0]

                return result
            except Exception as e:
                print(f"⚠️  Failed to transcribe segment {os.path.basename(segment_path)}: {e}")
                return None

        # Transcribe all segments in parallel with enhanced CPU utilization
        if config.get('max_perf'):
            # Ultra optimised: use up to 75% of CPU cores for VAD parallel processing
            max_workers = min(len(segment_files), max(1, int(config.get("cpu_threads", 4) * 0.75)))
        else:
            # Conservative: use up to 50% of CPU cores
            max_workers = min(len(segment_files), config.get("cpu_threads", 4) // 2)
        max_workers = max(1, min(max_workers, 12))  # Cap at 12 workers for stability

        # SEQUENTIAL TRANSCRIPTION: Process segments one by one for guaranteed order
        print("🔧 Using sequential transcription processing for perfect order...")
        
        segment_results = []
        for i, (seg_path, time_range) in enumerate(segment_files):
            print(f"🎯 Processing segment {i + 1}/{len(segment_files)}: {time_range[0]:.1f}s-{time_range[1]:.1f}s")
            try:
                result = transcribe_segment(seg_path, time_range)
                if result:
                    segment_results.append(result)
                    print(f"✅ Segment {i + 1} completed successfully")
                else:
                    print(f"⚠️  Segment {i + 1} returned empty result")
            except Exception as e:
                print(f"❌ Segment {i + 1} failed: {e}")
        
        print(f"📊 Successfully processed {len(segment_results)}/{len(segment_files)} segments in perfect temporal order")

        # Combine results from all segments IN TEMPORAL ORDER
        combined_result = {"text": "", "segments": []}
        text_parts = []

        # Process results in temporal order, skipping None entries
        for i, result in enumerate(segment_results):
            if result is not None and isinstance(result, dict):
                # Combine text in order
                text_content = result.get("text", "")
                if text_content and isinstance(text_content, str):
                    text_parts.append(text_content.strip())
                    print(f"📝 Added segment {i+1} text: '{text_content.strip()[:50]}...'")

                # Combine segments in order
                segments_data = result.get("segments", [])
                if segments_data and isinstance(segments_data, list):
                    combined_result["segments"].extend(segments_data)

                # Copy metadata from first valid result
                if not combined_result.get("language") and result.get("language"):
                    # Copy common metadata fields
                    for key in ["language", "language_probability", "duration"]:
                        if key in result and key not in combined_result:
                            combined_result[key] = result[key]
            else:
                print(f"⚠️  Skipping empty segment {i+1}")

        # Assemble text in proper temporal order
        combined_result["text"] = " ".join(text_parts).strip()
        
        # Debug: Show first part of combined text
        if combined_result["text"]:
            first_part = combined_result["text"][:100]
            print(f"🔍 Combined text starts with: '{first_part}...'")
        else:
            print("⚠️  Combined text is empty!")

        # Ensure segments are sorted by start time (double-check)
        if combined_result["segments"]:
            combined_result["segments"].sort(key=lambda x: x.get("start", 0))
            
            # Verify segment order and report any issues
            prev_end = 0
            for i, seg in enumerate(combined_result["segments"]):
                seg_start = seg.get("start", 0)
                if seg_start < prev_end - 1:  # Allow 1 second tolerance for overlaps
                    print(f"⚠️  Segment order issue detected at segment {i}: start={seg_start:.1f}s, prev_end={prev_end:.1f}s")
                prev_end = seg.get("end", seg_start)

        print(f"✅ Combined transcription from {len(segment_results)} segments in temporal order")
        print(f"📝 Total text length: {len(combined_result['text'])} characters")
        print(f"🎯 Total segments: {len(combined_result.get('segments', []))}")
        
        # Debug: Show first few characters to verify beginning is preserved
        if combined_result["text"]:
            preview = combined_result["text"][:100]
            print(f"🔍 Transcript begins: '{preview}...'")
        
        return combined_result


def transcribe_file_simple_auto(input_path, output_dir=None, threads_override: Optional[int] = None):
    """
    High-quality, simplified single-file transcription on best available device.
    - Device selection: CUDA > DirectML > CPU
    - No VAD; transcribe the entire file
    - Robust DOCX save with fallback
    - Safe cleanup that avoids torch re-import problems
    Returns path to the .txt file.
    """
    # Initialize all variables at the beginning to ensure they're always accessible
    use_vad = False
    enable_speakers = False
    use_dataset = False
    max_perf = True
    transcription_complete = False
    transcription_result = None
    transcription_error = None
    using_fw = False
    using_distil = False
    using_insanely_fast = False

    _ensure_torch_available()
    torch_api = cast(Any, torch)

    # Lazy imports (after torch is imported) to avoid docstring errors
    import whisper  # type: ignore
    from docx import Document  # type: ignore
    try:
        from text_processor_enhanced import create_enhanced_processor
        _enhanced_processor_available = True
    except ImportError:
        _enhanced_processor_available = False
        create_enhanced_processor = None  # Define it to avoid unbound variable error
    from transcribe import (
        get_media_duration, split_into_paragraphs, format_duration, format_duration_minutes_only, format_duration_hms,
    )

    # Speaker identification imports
    try:
        import webrtcvad
        _vad_available = True
    except ImportError:
        _vad_available = False

    # Import VAD functions from transcribe.py
    try:
        from transcribe import vad_segment_times
        _vad_functions_available = True
    except ImportError:
        _vad_functions_available = False
        # Define fallback VAD function
        def vad_segment_times(input_path):
            """Fallback VAD function when transcribe.py is not available"""
            try:
                from moviepy import AudioFileClip
                audio_clip = AudioFileClip(input_path)
                duration = audio_clip.duration
                audio_clip.close()

                segments = []
                segment_length = 30.0  # 30 second segments
                for i in range(0, int(duration), int(segment_length)):
                    start = float(i)
                    end = min(float(i + segment_length), duration)
                    segments.append((start, end))

                print(f"📊 Created {len(segments)} duration-based segments ({segment_length}s each)")
                return segments
            except Exception as e:
                print(f"❌ Error creating fallback segments: {e}")
                return [(0.0, 60.0)]  # Single segment fallback

    start_time = time.time()

    # Check if speaker identification should be enabled
    enable_speakers = False
    try:
        # Speaker identification is disabled - comment out to re-enable if needed
        # enable_speakers = os.environ.get("TRANSCRIBE_SPEAKER_ID", "").strip() in ("1", "true", "True")
        if enable_speakers:
            print("� Speaker identification enabled")
    except Exception:
        enable_speakers = False

    # Check if dataset optimization should be used - DEFAULT TO OFF for quality/stability
    # Enable via TRANSCRIBE_USE_DATASET=1 if you want the segmented GPU path.
    use_dataset = False
    try:
        env_dataset = os.environ.get("TRANSCRIBE_USE_DATASET", "").strip()
        if env_dataset in ("1", "true", "True"):
            use_dataset = True
            print("✅ Dataset optimization ENABLED via TRANSCRIBE_USE_DATASET=1")
    except Exception:
        use_dataset = False

    # Check if VAD segmentation should be used
    use_vad = False
    try:
        use_vad = os.environ.get("TRANSCRIBE_VAD", "").strip() in ("1", "true", "True")
        if use_vad:
            print("🎯 VAD segmentation enabled for performance optimization")
    except Exception:
        use_vad = False
    
    # Use dataset optimization (parallel GPU processing) for all files when enabled
    if use_dataset:
        try:
            result = transcribe_with_dataset_optimization(input_path, output_dir, threads_override)
            # If we got a valid result, return it immediately
            if result and isinstance(result, str) and os.path.exists(result):
                return result
            # If result is invalid, fall through to standard processing
            print(f"⚠️  Dataset optimization returned invalid result, falling back to standard processing")
            use_dataset = False
        except Exception as e:
            print(f"⚠️  Dataset optimization failed: {e} - falling back to standard processing")
            use_dataset = False

    # Decide max performance mode from env - DEFAULT TO MAX PERF FOR BETTER CPU UTILIZATION
    max_perf = True  # Default to maximum performance for better CPU utilization
    try:
        env_max_perf = os.environ.get("TRANSCRIBE_MAX_PERF", "").strip()
        if env_max_perf in ("0", "false", "False"):
            max_perf = False
    except Exception:
        max_perf = True
    
    # VAD CONTROL: Allow disabling VAD via environment variable
    disable_vad = False
    try:
        env_disable_vad = os.environ.get("TRANSCRIBE_DISABLE_VAD", "").strip()
        if env_disable_vad in ("1", "true", "True"):
            disable_vad = True
            print("🚫 VAD processing disabled via TRANSCRIBE_DISABLE_VAD environment variable")
    except Exception:
        disable_vad = False
    
    config = get_maximum_hardware_config(max_perf=max_perf)
    config['disable_vad'] = disable_vad
    
    # Compact hardware config summary
    try:
        gpu_info = ""
        if float(config.get('allowed_vram_gb') or 0) > 0:
            gpu_info = f" | GPU: {float(config['allowed_vram_gb']):.1f}GB VRAM"
        print(f"⚙️  Config: {config['cpu_threads']} threads, {config['usable_ram_gb']:.1f}GB RAM{gpu_info}")
    except Exception:
        pass
        
    if not output_dir:
        output_dir = os.path.dirname(input_path)

    # Preprocess audio with silence padding to prevent missed words
    try:
        preprocessed_path = preprocess_audio_with_padding(input_path)
    except Exception as _pre_e:
        print(f"⚠️  Preprocessing step failed early: {_pre_e} - using original file")
        preprocessed_path = input_path
    preprocessing_used = preprocessed_path != input_path
    working_input_path = preprocessed_path if preprocessing_used else input_path

    duration = get_media_duration(working_input_path)
    if duration:
        print(f"⏱️  Duration: {format_duration(duration)}")

    # Build optional initial prompt from awkward words
    awkward_terms = load_awkward_terms(input_path)
    initial_prompt = build_initial_prompt(awkward_terms)
    if initial_prompt:
        preview = initial_prompt[:100]
        print(f"🧩 Using domain terms bias (initial_prompt): '{preview}...'")

    # Pre-run cleanup
    force_gpu_memory_cleanup()

    # Choose device and load one model only
    device_name = "CPU"
    model: Any = None
    chosen_device = "cpu"
    
    # Check for model selection from environment variable (set by GUI)
    selected_model_name = os.environ.get("TRANSCRIBE_MODEL_NAME", "large-v3")
    
    # Parse model name to determine backend and actual model
    # Format: "backend-modelname" or just "modelname" (default to native whisper)
    backend = "native"  # Default backend
    actual_model_name = selected_model_name
    
    if selected_model_name.startswith("faster-whisper-"):
        backend = "faster-whisper"
        actual_model_name = selected_model_name.replace("faster-whisper-", "")
        print(f"🚀 Backend: Faster-Whisper (CTranslate2) - 4x faster")
    elif selected_model_name.startswith("distil-whisper-"):
        backend = "distil-whisper"
        actual_model_name = "distil-whisper/distil-large-v3"  # HuggingFace model ID
        print(f"🚀 Backend: Distil-Whisper (HuggingFace) - 6x faster, English-only")
    elif selected_model_name == "insanely-fast-whisper":
        backend = "insanely-fast-whisper"
        actual_model_name = "openai/whisper-large-v3"  # Uses HF pipeline with optimizations
        print(f"🚀 Backend: Insanely-Fast-Whisper (Flash Attention + Batching)")
    else:
        print(f"🎯 Backend: Native OpenAI Whisper")
    
    # Store original for logging
    original_model_selection = selected_model_name
    selected_model_name = actual_model_name

    # Prefer selected model; if it's not listed as available, fall back to the next best
    # (Only applies to native whisper backend)
    if backend == "native":
        try:
            import whisper  # type: ignore
            avail = set(whisper.available_models())
            requested_available = (selected_model_name in avail)
            if not requested_available:
                print(f"⚠️  Requested model '{selected_model_name}' not available, falling back...")
                for cand in ("large-v3", "large-v2", "large"):
                    if cand in avail:
                        selected_model_name = cand
                        break
            print(f"🧩 Requested model available: {requested_available}")
            print(f"🗂️  Selecting model: {selected_model_name}")
            print(f"🎯 Model Source: Environment variable TRANSCRIBE_MODEL_NAME = '{os.environ.get('TRANSCRIBE_MODEL_NAME', 'NOT SET')}'")
        except Exception as e:
            print(f"⚠️  Could not query whisper.available_models(): {e}. Proceeding with '{selected_model_name}'.")
    else:
        print(f"🗂️  Model: {selected_model_name} (via {backend} backend)")

    try:
        # Elevate process priority on Windows for max perf
        if config.get('max_perf'):
            try:
                import psutil
                p = psutil.Process(os.getpid())
                if hasattr(psutil, 'HIGH_PRIORITY_CLASS'):
                    p.nice(psutil.HIGH_PRIORITY_CLASS)
                    print("🚀 Process priority set to HIGH")
            except Exception as e:
                print(f"⚠️  Could not raise process priority: {e}")

        if "cuda" in config["devices"] and torch_api.cuda.is_available():
            chosen_device = "cuda"
            device_name = f"CUDA GPU ({torch_api.cuda.get_device_name(0)})"
            print("🎯 Device: CUDA GPU")
            # Apply CUDA per-process memory fraction if an allowed VRAM cap is set
            try:
                total_vram = float(config.get("cuda_total_vram_gb") or 0.0)
                allowed_vram = float(config.get("allowed_vram_gb") or 0.0)
                # Preflight available (current used) memory BEFORE load
                try:
                    used_vram = torch_api.cuda.memory_allocated() / (1024**3)
                except Exception:
                    used_vram = 0.0
                free_est = max(0.0, total_vram - used_vram)
                # Heuristic model VRAM footprint (base weights + buffers)
                MODEL_VRAM_GB_EST = {
                    "large-v3": 7.4,
                    "large-v3-turbo": 5.2,
                    "large": 6.8,
                    "medium": 3.2,
                    "small": 1.4,
                    "base": 0.9,
                    "tiny": 0.6,
                }
                est_need = MODEL_VRAM_GB_EST.get(selected_model_name.lower(), 5.0) + 0.4  # + overhead margin
                # Allow override to force attempt
                force_gpu = os.environ.get("TRANSCRIBE_FORCE_GPU", "").lower() in ("1","true","yes")
                if free_est and free_est < est_need and not force_gpu:
                    print(f"⚠️  Preflight: estimated free VRAM {free_est:.2f}GB < required ~{est_need:.2f}GB for '{selected_model_name}'. Using CPU to avoid OOM.")
                    raise RuntimeError("Preflight GPU memory insufficient")
                if total_vram > 0:
                    if 0.5 <= allowed_vram < total_vram:
                        frac = max(0.05, min(0.95, allowed_vram / total_vram))
                        torch_api.cuda.set_per_process_memory_fraction(frac, device=0)
                        print(f"🧩 Limiting CUDA allocator to ~{frac*100:.0f}% of VRAM ({allowed_vram:.1f}GB)")
                    elif config.get('max_perf'):
                        # Leave a safety margin unless explicitly overridden
                        # Adjust default for low-VRAM GPUs with large models
                        default_frac = "0.85" if total_vram <= 8 and selected_model_name in ["large-v3", "large-v3-turbo"] else "0.92"
                        high_frac = os.environ.get("TRANSCRIBE_GPU_FRACTION", default_frac)
                        try:
                            frac_val = float(high_frac)
                        except Exception:
                            frac_val = float(default_frac)
                        try:
                            torch_api.cuda.set_per_process_memory_fraction(min(0.99, max(0.5, frac_val)), device=0)
                            print(f"🧩 Allowing CUDA allocator to use ~{min(0.99, max(0.5, frac_val))*100:.0f}% of VRAM (safety margin enabled)")
                        except Exception:
                            pass
            except Exception as e:
                print(f"⚠️  Could not set CUDA memory fraction: {e}")
            # Enable ULTRA performance knobs for maximum GPU utilization
            try:
                if hasattr(torch_api.backends, "cudnn"):
                    torch_api.backends.cudnn.benchmark = True
                    # Enable deterministic mode for reproducibility but max performance
                    torch_api.backends.cudnn.deterministic = False
                # TF32 can speed up matmul on Ampere+; harmless elsewhere
                if hasattr(torch_api.backends, "cuda") and hasattr(torch_api.backends.cuda, "matmul"):
                    try:
                        torch_api.backends.cuda.matmul.allow_tf32 = True
                        # Enable cuBLAS optimizations
                        torch_api.backends.cuda.matmul.allow_bf16_reduced_precision_reduction = True
                    except Exception:
                        pass
                try:
                    torch_api.set_float32_matmul_precision("high")
                except Exception:
                    pass
                # Enable GPU memory pooling for better shared memory utilization
                try:
                    torch_api.cuda.set_per_process_memory_fraction(0.99, device=0)
                    print("🧩 ULTRA GPU Memory Pooling: Enabled for 15GB+ shared memory utilization")
                except Exception:
                    pass
            except Exception:
                pass
            # Final pre-load cache clear to reduce fragmentation
            try:
                torch_api.cuda.empty_cache()
                torch_api.cuda.synchronize()
            except Exception:
                pass
            alloc_conf = os.environ.get("PYTORCH_CUDA_ALLOC_CONF", "")
            if not alloc_conf:
                # Provide sane defaults with expandable segments unless user overrides
                os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "expandable_segments:True,max_split_size_mb:64"
            
            # Use the backend determined from model name parsing above
            # Auto-switch to faster-whisper for large models on low-VRAM GPUs (only if native)
            if backend == "native" and total_vram <= 8 and selected_model_name in ["large-v3", "large-v3-turbo"]:
                backend = "faster-whisper"
                print(f"🎯 Auto-switching to faster-whisper backend for {selected_model_name} on {total_vram:.1f}GB GPU")
            
            using_fw = False
            using_distil = False
            using_insanely_fast = False
            
            # === DISTIL-WHISPER BACKEND ===
            if backend == "distil-whisper":
                try:
                    from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
                    
                    print(f"🔄 Loading Distil-Whisper model: {selected_model_name}")
                    
                    # Determine compute type
                    torch_dtype = torch_api.float16 if torch_api.cuda.is_available() else torch_api.float32
                    
                    distil_model = AutoModelForSpeechSeq2Seq.from_pretrained(
                        selected_model_name,
                        torch_dtype=torch_dtype,
                        low_cpu_mem_usage=True,
                        use_safetensors=True
                    )
                    distil_model.to("cuda" if torch_api.cuda.is_available() else "cpu")
                    
                    distil_processor = AutoProcessor.from_pretrained(selected_model_name)
                    
                    # Create pipeline
                    model = pipeline(
                        "automatic-speech-recognition",
                        model=distil_model,
                        tokenizer=distil_processor.tokenizer,
                        feature_extractor=distil_processor.feature_extractor,
                        torch_dtype=torch_dtype,
                        device="cuda" if torch_api.cuda.is_available() else "cpu",
                    )
                    using_distil = True
                    print(f"✅ Distil-Whisper loaded successfully")
                except Exception as distil_e:
                    print(f"⚠️  Distil-Whisper load failed: {distil_e}")
                    print("🔄 Falling back to native Whisper...")
                    backend = "native"
            
            # === INSANELY-FAST-WHISPER BACKEND ===
            elif backend == "insanely-fast-whisper":
                try:
                    from transformers import pipeline
                    
                    print(f"🔄 Loading Insanely-Fast-Whisper: {selected_model_name}")
                    
                    # Use Flash Attention 2 if available
                    model_kwargs = {"use_flash_attention_2": True} if torch_api.cuda.is_available() else {}
                    
                    model = pipeline(
                        "automatic-speech-recognition",
                        model=selected_model_name,
                        torch_dtype=torch_api.float16 if torch_api.cuda.is_available() else torch_api.float32,
                        device="cuda" if torch_api.cuda.is_available() else "cpu",
                        model_kwargs=model_kwargs,
                    )
                    using_insanely_fast = True
                    print(f"✅ Insanely-Fast-Whisper loaded with Flash Attention")
                except Exception as isf_e:
                    print(f"⚠️  Insanely-Fast-Whisper load failed: {isf_e}")
                    # Try without flash attention
                    try:
                        model = pipeline(
                            "automatic-speech-recognition",
                            model=selected_model_name,
                            torch_dtype=torch_api.float16 if torch_api.cuda.is_available() else torch_api.float32,
                            device="cuda" if torch_api.cuda.is_available() else "cpu",
                        )
                        using_insanely_fast = True
                        print(f"✅ Insanely-Fast-Whisper loaded (without Flash Attention)")
                    except Exception as isf_e2:
                        print(f"⚠️  Insanely-Fast-Whisper fallback also failed: {isf_e2}")
                        print("🔄 Falling back to native Whisper...")
                        backend = "native"
            
            # === FASTER-WHISPER BACKEND ===
            elif backend == "faster-whisper":
                try:
                    from faster_whisper import WhisperModel  # type: ignore
                except Exception as fw_imp_e:
                    print(f"⚠️  faster-whisper import failed ({fw_imp_e}); falling back to native backend")
                    backend = "native"
                    
            if backend == "faster-whisper":
                pref_raw = os.environ.get("TRANSCRIBE_FW_COMPUTE_TYPES", "auto,int8,float16")
                compute_order = [c.strip() for c in pref_raw.split(',') if c.strip()]
                load_success = False
                for idx, ctype in enumerate(compute_order, start=1):
                    try:
                        torch_api.cuda.empty_cache(); torch_api.cuda.synchronize()
                        print(f"🔁 FW Attempt {idx}: compute_type={ctype}")
                        fw_model = WhisperModel(selected_model_name, device="cuda", compute_type=ctype)
                        model = fw_model  # type: ignore
                        using_fw = True
                        load_success = True
                        print(f"🎯 faster-whisper model '{selected_model_name}' loaded (compute_type={ctype})")
                        break
                    except RuntimeError as rte:
                        if 'out of memory' in str(rte).lower():
                            print(f"⚠️  FW CUDA OOM on compute_type={ctype}: {rte}")
                            continue
                        print(f"⚠️  FW load error (compute_type={ctype}): {rte}")
                        continue
                    except Exception as e_fw:
                        print(f"⚠️  FW general error (compute_type={ctype}): {e_fw}")
                        continue
                if not load_success:
                    raise RuntimeError(f"All faster-whisper load attempts failed for '{selected_model_name}'")
            else:
                force_fp16_env = os.environ.get("TRANSCRIBE_FORCE_FP16", "").lower() in ("1","true","yes")
                attempts = [
                    {"fp16": False, "adjust_fraction": None},
                    {"fp16": True, "adjust_fraction": 0.90},
                    {"fp16": True, "adjust_fraction": 0.85},
                    {"fp16": True, "adjust_fraction": 0.80},
                    {"fp16": True, "adjust_fraction": 0.75},
                    {"fp16": True, "adjust_fraction": 0.70},
                ] if not force_fp16_env else [
                    {"fp16": True, "adjust_fraction": None},
                    {"fp16": True, "adjust_fraction": 0.90},
                    {"fp16": True, "adjust_fraction": 0.85},
                    {"fp16": True, "adjust_fraction": 0.80},
                    {"fp16": True, "adjust_fraction": 0.75},
                    {"fp16": True, "adjust_fraction": 0.70},
                ]
                load_success = False
                for idx, att in enumerate(attempts, start=1):
                    try:
                        if att["adjust_fraction"] is not None:
                            try:
                                torch_api.cuda.set_per_process_memory_fraction(att["adjust_fraction"], device=0)
                                print(f"🔁 Attempt {idx}: fraction={att['adjust_fraction']:.2f}, fp16={att['fp16']}")
                            except Exception:
                                print(f"🔁 Attempt {idx}: fp16={att['fp16']} (fraction adjust failed)")
                        else:
                            print(f"🔁 Attempt {idx}: fp16={att['fp16']}")
                        torch_api.cuda.empty_cache(); torch_api.cuda.synchronize()
                        model = whisper.load_model(selected_model_name, device="cuda")
                        if att["fp16"]:
                            try:
                                model.to(torch_api.float16)
                                print("🎯 Converted to FP16 for reduced VRAM footprint")
                            except Exception as fp16_e:
                                print(f"⚠️  FP16 conversion failed (attempt {idx}): {fp16_e}")
                        load_success = True
                        print(f"🎯 Model '{selected_model_name}' loaded on CUDA (attempt {idx})")
                        break
                    except RuntimeError as rte:
                        msg = str(rte).lower()
                        if "out of memory" in msg or "cuda error" in msg:
                            print(f"⚠️  CUDA OOM on attempt {idx}: {rte}")
                            continue
                        else:
                            print(f"⚠️  Non-OOM CUDA load error (attempt {idx}): {rte}")
                            break
                    except Exception as gen_e:
                        print(f"⚠️  General CUDA load error (attempt {idx}): {gen_e}")
                        break
                if not load_success:
                    raise RuntimeError(f"All CUDA load attempts failed for '{selected_model_name}'")
            try:
                used_after = torch_api.cuda.memory_allocated() / (1024**3)
                try:
                    total_vram_lookup = torch_api.cuda.get_device_properties(0).total_memory / (1024**3)
                except Exception:
                    total_vram_lookup = config.get("cuda_total_vram_gb") or 0.0
                print(f"📊 VRAM in use after load: {used_after:.2f} GB / {total_vram_lookup:.2f} GB")
            except Exception:
                pass
        elif config.get("dml_available", False):
            try:
                import torch_directml  # type: ignore
                dml_device = torch_directml.device()
                chosen_device = dml_device
                device_name = "DirectML GPU"
                print("🎯 Device: DirectML GPU")
                model = whisper.load_model(selected_model_name, device=dml_device)
            except Exception as e:
                print(f"⚠️  DirectML unavailable, falling back to CPU: {e}")
                model = None
        if model is None:
            chosen_device = "cpu"
            device_name = f"CPU ({multiprocessing.cpu_count()} cores)"
            print(f"🎯 Device: {device_name}")
            model = whisper.load_model(selected_model_name, device="cpu")
    except Exception as load_e:
        print(f"❌ Model load failed on preferred device: {load_e}")
        print("🔄 Falling back to CPU...")
        chosen_device = "cpu"
        device_name = f"CPU ({multiprocessing.cpu_count()} cores)"
        try:
            import whisper  # type: ignore
            model = whisper.load_model(selected_model_name, device="cpu")
        except Exception:
            # Last resort
            model = whisper.load_model("large", device="cpu")

    # Apply explicit threads override if provided
    if isinstance(threads_override, int) and threads_override > 0:
        config["cpu_threads"] = max(1, min(64, threads_override))

    # Set CPU threads (& interop) with enhanced threading for text processing
    torch_api.set_num_threads(config["cpu_threads"])
    
    # Calculate interop threads with enhanced settings
    if config.get('max_perf'):
        interop = max(4, min(24, config["cpu_threads"] // 2))  # More optimised interop
    else:
        interop = max(2, min(16, config["cpu_threads"] // 4))
    
    # Explicit thread configuration logging
    print(f"🔧 Thread Configuration:")
    print(f"   • CPU Threads: {config['cpu_threads']}")
    print(f"   • Interop Threads: {interop}")
    print(f"   • Device: {device_name}")
    print(f"   • Model: {selected_model_name}")
    
    try:
        torch_api.set_num_interop_threads(interop)
    except Exception:
        pass
    
    # Enhanced MKL/OMP configuration for better CPU utilization
    try:
        os.environ.setdefault("MKL_NUM_THREADS", str(config["cpu_threads"]))
        os.environ.setdefault("OMP_NUM_THREADS", str(config["cpu_threads"]))
        # Additional performance tunings
        os.environ.setdefault("MKL_DYNAMIC", "TRUE")  # Dynamic thread adjustment
        os.environ.setdefault("OMP_DYNAMIC", "TRUE")  # Dynamic thread adjustment
        os.environ.setdefault("NUMEXPR_MAX_THREADS", str(min(config["cpu_threads"], 16)))  # NumPy/SciPy threading
        
        # Enable optimised CPU utilization
        os.environ.setdefault("OMP_WAIT_POLICY", "ACTIVE")  # Keep threads active (don't sleep)
        os.environ.setdefault("OMP_PROC_BIND", "TRUE")  # Bind threads to cores
        os.environ.setdefault("KMP_BLOCKTIME", "0")  # Immediate response (Intel specific)
        os.environ.setdefault("KMP_AFFINITY", "granularity=fine,compact,1,0")  # Core affinity
        
        # Text processing specific threading
        text_threads = max(2, min(8, config["cpu_threads"] // 2))
        os.environ.setdefault("NLTK_NUM_THREADS", str(text_threads))
        os.environ.setdefault("SPACY_NUM_THREADS", str(text_threads))
        
    except Exception:
        pass
    
    text_workers = max(2, min(8, config["cpu_threads"] // 2))
    print(f"🧵 Enhanced threading: PyTorch {config['cpu_threads']} threads, interop {interop}, text processing up to {text_workers} workers")
    print(f"🔧 CPU optimization: MKL/OMP dynamic threading enabled for maximum utilization")

    # Note: batch size is not passed to Whisper to ensure broad compatibility across versions

    # Optional: NVML for GPU utilisation logging
    nvml = None
    try:
        import pynvml  # type: ignore
        pynvml.nvmlInit()
        nvml = pynvml
    except Exception:
        nvml = None

    # Run transcription in a watchdog thread
    import threading
    transcription_complete = False
    transcription_result = None
    transcription_error = None

    def _run_transcribe():
        nonlocal transcription_complete, transcription_result, transcription_error, use_vad, using_fw, using_distil, using_insanely_fast
        try:
            print("🔄 Starting transcription process...")
            if model is None:
                raise RuntimeError("Transcription model is not loaded")

            # === DISTIL-WHISPER or INSANELY-FAST-WHISPER (HuggingFace Pipeline) ===
            if using_distil or using_insanely_fast:
                backend_name = "Distil-Whisper" if using_distil else "Insanely-Fast-Whisper"
                print(f"🎯 Using {backend_name} pipeline for transcription...")
                
                # HuggingFace pipeline transcription
                pipeline_kwargs = {
                    "return_timestamps": True,
                    "chunk_length_s": 30,
                    "batch_size": 16 if using_insanely_fast else 8,  # Larger batches for insanely-fast
                }
                
                # Add language hint
                generate_kwargs = {"language": "en", "task": "transcribe"}
                
                result = model(working_input_path, generate_kwargs=generate_kwargs, **pipeline_kwargs)
                
                # Convert to standard format
                if isinstance(result, dict):
                    text = result.get("text", "")
                    chunks = result.get("chunks", [])
                    segments = []
                    for chunk in chunks:
                        if isinstance(chunk, dict):
                            seg_text = chunk.get("text", "")
                            timestamps = chunk.get("timestamp", (0, 0))
                            if timestamps and len(timestamps) >= 2:
                                segments.append({
                                    "text": seg_text,
                                    "start": timestamps[0] or 0,
                                    "end": timestamps[1] or 0,
                                })
                    transcription_result = {"text": text, "segments": segments}
                else:
                    transcription_result = {"text": str(result), "segments": []}
                
                print(f"✅ {backend_name} transcription completed")
                transcription_complete = True
                return

            # === NATIVE WHISPER / FASTER-WHISPER ===
            # Apply VAD segmentation if enabled
            transcribe_kwargs = {
                "language": "en",  # Optimized for English language
                # Conservative thresholds with mild rejection of low-confidence / compressed gibberish
                "compression_ratio_threshold": 2.4,
                "logprob_threshold": -0.5,
                "no_speech_threshold": 0.4,
                # Disable cross-segment conditioning to prevent repetition loops
                "condition_on_previous_text": False,
                "temperature": 0.0,
                # Note: temperature_increment_on_fallback not supported in some Whisper builds
                "verbose": False,
            }
            
            # Model-specific tuning for accuracy
            if selected_model_name == "large-v3-turbo":
                # Turbo-specific: slightly stricter compression ratio, same logprob guard
                transcribe_kwargs["compression_ratio_threshold"] = 2.2
                transcribe_kwargs["no_speech_threshold"] = 0.4
                print("🎯 Turbo model: repetition-guard thresholds applied")
            elif selected_model_name == "large-v3":
                # Large-v3: allow a bit more compression ratio, keep logprob guard
                transcribe_kwargs["compression_ratio_threshold"] = 2.6
                print("🎯 Large-v3 model: repetition-guard thresholds applied")
            
            # Apply quality mode if enabled
            quality_mode = os.environ.get("TRANSCRIBE_QUALITY_MODE", "").strip() in ("1", "true", "True")
            if quality_mode:
                if selected_model_name == "large-v3-turbo":
                    # Turbo with MAXIMUM quality mode: aggressive beam search for best accuracy
                    transcribe_kwargs["beam_size"] = 10  # Maximum beam width
                    transcribe_kwargs["patience"] = 3.0  # Maximum patience
                    transcribe_kwargs["best_of"] = 10  # Try many candidates
                    transcribe_kwargs["temperature"] = 0.0  # Greedy decoding for consistency
                    transcribe_kwargs["compression_ratio_threshold"] = 2.8  # More lenient
                    transcribe_kwargs["no_speech_threshold"] = 0.4  # Less aggressive silence detection
                    print("🎯 ULTRA Quality mode (turbo): beam_size=10, patience=3.0, best_of=10, temp=0")
                else:
                    # Large-v3 with maximum quality mode
                    transcribe_kwargs["beam_size"] = 10
                    transcribe_kwargs["patience"] = 3.0
                    transcribe_kwargs["best_of"] = 10
                    transcribe_kwargs["temperature"] = 0.0
                    transcribe_kwargs["compression_ratio_threshold"] = 2.8
                    transcribe_kwargs["no_speech_threshold"] = 0.4
                    print("🎯 ULTRA Quality mode (large-v3): beam_size=10, patience=3.0, best_of=10, temp=0")
            
            # Tune defaults when using faster-whisper to avoid over-pruning
            if using_fw:
                # Greedy capture but still apply repetition guard by disabling conditioning
                transcribe_kwargs["beam_size"] = 1
                transcribe_kwargs.pop("patience", None)
                transcribe_kwargs["best_of"] = 1
                transcribe_kwargs["no_speech_threshold"] = 0.25
                # Keep compression ratio guard off to avoid over-pruning empty results for FW
                transcribe_kwargs["compression_ratio_threshold"] = None
                transcribe_kwargs["vad_filter"] = False
                transcribe_kwargs["chunk_length"] = 30
                print("🎯 FW tuning: greedy decode, repetition guard (no conditioning), chunk_length=30")
            
            # Gate initial prompt behind explicit opt-in to preserve strict verbatim neutrality
            if initial_prompt and str(os.environ.get("TRANSCRIBE_ALLOW_PROMPT", "0")).lower() in ("1","true","yes"):
                transcribe_kwargs["initial_prompt"] = initial_prompt

            if use_vad:
                try:
                    # Get VAD segments for the audio file
                    # VAD CONTROL: Allow disabling VAD when ordering issues occur
                    if config.get('disable_vad', False):
                        print("⚠️  VAD processing DISABLED - using full file transcription for guaranteed temporal order")
                        print("   This ensures perfect segment ordering but may be slower for long files")
                    elif _vad_functions_available:
                        vad_segments = vad_segment_times(working_input_path)
                        if vad_segments and len(vad_segments) > 0:
                            print(f"🎯 VAD detected {len(vad_segments)} speech segments - processing in parallel")
                            print("💡 If transcript beginnings are missing, set 'disable_vad': True in config")
                            # Use actual VAD segmentation with parallel processing
                            result = transcribe_with_vad_parallel(working_input_path, vad_segments, model, transcribe_kwargs, config)
                            transcription_result = result
                            print("✅ VAD parallel transcription completed successfully")
                            transcription_complete = True
                            return  # Exit early since we processed with VAD
                        else:
                            print("⚠️  VAD enabled but no segments detected, proceeding without VAD")
                    else:
                        print("⚠️  VAD functions not available, proceeding without VAD")
                except Exception as vad_e:
                    print(f"⚠️  VAD segmentation failed: {vad_e} - proceeding without VAD")
                    use_vad = False  # Disable for this run

            # Call transcribe with optimized parameters
            result = _compatible_transcribe_call(model, working_input_path, transcribe_kwargs)
            transcription_result = _as_result_dict(result)
            try:
                segs_dbg = transcription_result.get('segments') if isinstance(transcription_result, dict) else None
                segs_count = len(segs_dbg) if isinstance(segs_dbg, list) else 0
                txt_dbg = transcription_result.get('text') if isinstance(transcription_result, dict) else ''
                print(f"🔎 Transcribe result: segments={segs_count}, text_len={len(txt_dbg) if isinstance(txt_dbg, str) else 0}")
                if segs_count == 0 and (not isinstance(txt_dbg, str) or len(txt_dbg.strip()) == 0):
                    # Retry 1: conservative greedy, no VAD, short chunks
                    print("⚠️  Empty result detected — retrying with conservative decode parameters (no VAD, greedy)")
                    fw_retry1 = {
                        'language': transcribe_kwargs.get('language'),  # keep if provided
                        'task': 'transcribe',
                        'temperature': 0.0,
                        'beam_size': 1,
                        'best_of': 1,
                        'condition_on_previous_text': False,
                        'vad_filter': False,
                        'compression_ratio_threshold': None,
                        'log_prob_threshold': None,
                        'no_speech_threshold': 0.5,
                        'without_timestamps': False,
                        'chunk_length': 30,
                    }
                    # Drop initial_prompt on FW retries to avoid bias blocking detection
                    result2 = _compatible_transcribe_call(model, working_input_path, fw_retry1)
                    transcription_result = _as_result_dict(result2)
                    segs_dbg2 = transcription_result.get('segments') if isinstance(transcription_result, dict) else None
                    segs_count2 = len(segs_dbg2) if isinstance(segs_dbg2, list) else 0
                    txt_dbg2 = transcription_result.get('text') if isinstance(transcription_result, dict) else ''
                    print(f"🔎 Retry result: segments={segs_count2}, text_len={len(txt_dbg2) if isinstance(txt_dbg2, str) else 0}")

                    # Retry 2 (optional): VAD-filtered greedy capture; OFF by default due to runtime
                    if segs_count2 == 0 and (not isinstance(txt_dbg2, str) or len(txt_dbg2.strip()) == 0):
                        if str(os.environ.get('TRANSCRIBE_FW_RETRY2', '0')).lower() in ('1','true','yes'):
                            print("⚠️  Retry still empty — trying VAD-filtered greedy decode with low no_speech threshold (FW_RETRY2=on)")
                            fw_retry2 = {
                                'language': None,  # auto-detect language
                                'task': 'transcribe',
                                'temperature': 0.0,
                                'beam_size': 1,
                                'best_of': 1,
                                'condition_on_previous_text': False,
                                'vad_filter': True,
                                'vad_parameters': {'min_silence_duration_ms': 250},
                                'compression_ratio_threshold': None,
                                'log_prob_threshold': None,
                                'no_speech_threshold': 0.1,
                                'without_timestamps': False,
                                'chunk_length': 30,
                            }
                            try:
                                result3 = _compatible_transcribe_call(model, working_input_path, fw_retry2)
                                transcription_result = _as_result_dict(result3)
                                segs_dbg3 = transcription_result.get('segments') if isinstance(transcription_result, dict) else None
                                segs_count3 = len(segs_dbg3) if isinstance(segs_dbg3, list) else 0
                                txt_dbg3 = transcription_result.get('text') if isinstance(transcription_result, dict) else ''
                                print(f"🔎 Retry2 result: segments={segs_count3}, text_len={len(txt_dbg3) if isinstance(txt_dbg3, str) else 0}")
                                if segs_count3 > 0 or (isinstance(txt_dbg3, str) and len(txt_dbg3.strip()) > 0):
                                    # success; skip CPU fallback
                                    pass
                                else:
                                    raise RuntimeError('Retry2 still empty')
                            except Exception as r2e:
                                print(f"⚠️  Retry2 error/empty: {r2e}")
                                # Fall through to CPU fallback
                        # Final fallback: native Whisper on CPU
                        try:
                            print("🔁 All FW attempts empty — falling back to native Whisper on CPU for this file")
                            import whisper as _wh
                            _cpu_model = _wh.load_model(selected_model_name, device='cpu')
                            # Decode-time guard rails to mitigate repetition/hallucination
                            cpu_kwargs = {
                                'language': transcribe_kwargs.get('language'),
                                # Try multiple temperatures progressively to escape repetitive paths
                                'temperature': [0.0, 0.2, 0.4],
                                'beam_size': 5,
                                'patience': 2.0,
                                # Disable conditioning on previous text to prevent feedback loops
                                'condition_on_previous_text': False,
                                # Enable stricter thresholds to reject low-confidence gibberish
                                'compression_ratio_threshold': 2.0,
                                'logprob_threshold': -0.5,
                                'no_speech_threshold': 0.3,
                            }
                            # Avoid bias during fallback; only pass prompt if explicitly requested via env
                            if initial_prompt and str(os.environ.get('TRANSCRIBE_FALLBACK_ALLOW_PROMPT', '0')).lower() in ('1','true','yes'):
                                cpu_kwargs['initial_prompt'] = initial_prompt
                            cpu_res = _cpu_model.transcribe(working_input_path, **cpu_kwargs)
                            transcription_result = cpu_res if isinstance(cpu_res, dict) else _as_result_dict(cpu_res)
                        except Exception as cpu_e:
                            print(f"❌ Native CPU fallback failed: {cpu_e}")
            except Exception as dbg_e:
                print(f"⚠️  Debug/Retry flow error: {dbg_e}")
            print("✅ Whisper transcription completed successfully")
        except Exception as e:
            transcription_error = e
        finally:
            transcription_complete = True

    transcribe_thread = threading.Thread(target=_run_transcribe, daemon=True)
    transcribe_thread.start()

    start_watch = time.time()
    print(f"🎙️  Transcribing audio...")

    # Initialize CPU/RAM monitoring
    import psutil
    process = psutil.Process(os.getpid())
    last_status_time = 0
    status_interval = 10  # Update status every 10 seconds

    while not transcription_complete:
        time.sleep(2)
        elapsed = time.time() - start_watch
        
        # Only print status updates periodically to reduce log spam
        if elapsed - last_status_time >= status_interval:
            last_status_time = elapsed
            
            # Format elapsed time
            if elapsed >= 60:
                mins = int(elapsed // 60)
                secs = int(elapsed % 60)
                time_str = f"{mins}:{secs:02d}"
            else:
                time_str = f"{elapsed:.0f}s"
            
            if torch_api.cuda.is_available() and chosen_device == "cuda":
                try:
                    used = torch_api.cuda.memory_allocated() / (1024 ** 3)
                    total = torch_api.cuda.get_device_properties(0).total_memory / (1024 ** 3)
                    gpu_pct = (used / total) * 100
                    
                    # Build compact status line
                    status = f"⏳ Transcribing: {time_str} | GPU: {gpu_pct:.0f}%"
                    
                    # Add GPU utilization if available
                    if nvml is not None:
                        try:
                            h = nvml.nvmlDeviceGetHandleByIndex(0)
                            util = nvml.nvmlDeviceGetUtilizationRates(h)
                            status += f" (util: {util.gpu}%)"
                        except:
                            pass
                    
                    print(status)
                    
                    if gpu_pct > 95:
                        torch_api.cuda.empty_cache()
                except Exception as e:
                    print(f"⏳ Transcribing: {time_str}")
            else:
                # CPU mode - simpler status
                try:
                    cpu_pct = process.cpu_percent(interval=None)
                    print(f"⏳ Transcribing: {time_str} | CPU: {cpu_pct:.0f}%")
                except:
                    print(f"⏳ Transcribing: {time_str}")

    # No timeout fallback - transcription completes naturally
    if transcription_error:
        raise Exception(f"Transcription failed: {transcription_error}")

    # Extract text (with artifact suppression around music)
    result = transcription_result
    full_text = ""
    segments_with_speakers = []
    removed_segments = []
    kept_count = 0
    if isinstance(result, dict):
        segments = result.get("segments")
        if isinstance(segments, list) and segments:
            def _is_suspicious_music_artifact(seg_text: str) -> bool:
                t = (seg_text or "").lower()
                if "©" in seg_text or "(c)" in t or "copyright" in t:
                    return True
                markers = [
                    "bf-watch", "watch tv", "bfwatch", "all rights reserved",
                    "www.", "http://", "https://",
                ]
                return any(m in t for m in markers)

            cleaned_parts = []
            cleaned_segments = []
            
            # Sort segments by start time FIRST to ensure proper order
            segments.sort(key=lambda x: x.get("start", 0))
            print(f"🔍 Processing {len(segments)} segments in temporal order")
            
            for i, seg in enumerate(segments):
                seg_text = str(seg.get("text", "")).strip()
                # Light, in-segment de-repetition to curb decode loops without changing wording intent
                try:
                    seg_text = _clean_repetitions_in_segment(seg_text)
                except Exception as _e:
                    pass
                avg_logprob = seg.get("avg_logprob", 0.0)
                no_speech_prob = seg.get("no_speech_prob", 0.0)
                seg_start = seg.get("start", 0)
                seg_end = seg.get("end", 0)

                suspicious = _is_suspicious_music_artifact(seg_text)
                # Make artifact filtering MUCH more conservative - only remove if BOTH suspicious AND very low confidence
                very_low_confidence = (isinstance(avg_logprob, (int, float)) and avg_logprob < -1.0) and \
                                     (isinstance(no_speech_prob, (int, float)) and no_speech_prob > 0.8)

                # CONSERVATIVE FILTERING: Only remove if clearly suspicious AND very low confidence
                should_keep = seg_text and not (suspicious and very_low_confidence)

                if should_keep:
                    cleaned_parts.append(seg_text)
                    # store a copy with cleaned text for paragraph assembly
                    seg_copy = dict(seg)
                    seg_copy["text"] = seg_text
                    cleaned_segments.append(seg_copy)
                    kept_count += 1
                    
                    # Debug: Show first few segments to verify beginning preservation
                    if i < 5:
                        print(f"  ✅ Segment {i+1}: [{seg_start:.1f}s-{seg_end:.1f}s] '{seg_text[:50]}...' (logprob: {avg_logprob:.2f})")
                else:
                    removed_segments.append({
                        "text": seg_text[:120],
                        "avg_logprob": avg_logprob,
                        "no_speech_prob": no_speech_prob,
                        "start": seg_start,
                        "end": seg_end,
                    })
                    print(f"  🧽 Filtered segment {i+1}: [{seg_start:.1f}s-{seg_end:.1f}s] '{seg_text[:30]}...' (suspicious={suspicious}, low_conf={very_low_confidence})")
                    
            # Assemble text. In verbatim mode, prefer coherent paragraphs built from segment timings.
            para_text = None
            try:
                if _is_verbatim_mode() and cleaned_segments:
                    gap = float(os.environ.get("TRANSCRIBE_PARAGRAPH_GAP", "1.2"))
                    para_text = _segments_to_paragraphs(cleaned_segments, gap_threshold=gap)
            except Exception as _e:
                para_text = None

            full_text = (para_text if (para_text and para_text.strip()) else " ".join(cleaned_parts)).strip()
            
            # Additional debugging
            if cleaned_parts:
                first_part = cleaned_parts[0][:100] if cleaned_parts[0] else "N/A"
                print(f"🔍 First segment text: '{first_part}...'")
                
            print(f"📊 Segment filtering: kept {kept_count}/{len(segments)} segments")

            # Speaker identification removed as requested
        else:
            text_result = result.get("text", "")
            full_text = text_result.strip() if isinstance(text_result, str) else str(text_result).strip()
    elif isinstance(result, list) and result:
        full_text = str(result[0]).strip()
    else:
        full_text = ""

    if removed_segments:
        print(f"🧽 Artifact filter removed {len(removed_segments)} low-confidence watermark-like segment(s) during music.")
        # Show a brief preview for diagnostics
        sample = removed_segments[0]
        print(f"   e.g., '{sample['text']}' (avg_logprob={sample['avg_logprob']}, no_speech_prob={sample['no_speech_prob']})")

    if not full_text:
        print("⚠️  Warning: No transcription text generated")
        full_text = "[No speech detected or transcription failed]"

    print(f"⚡ Hardware utilised: {device_name}")

    # Post-processing with ULTRA/enhanced processors (skipped in verbatim)
    try:
        if _is_verbatim_mode():
            formatted_text = full_text
            quality_stats = {"verbatim": True}
        else:
            try:
                from text_processor_ultra import create_ultra_processor, create_advanced_paragraph_formatter
                text_workers = max(2, min(8, config.get("cpu_threads", 4) // 2))
                print(f"🚀 Using ULTRA text processor with {text_workers} workers and 6 specialized passes")
                t0 = time.time()
                ultra_processor = create_ultra_processor(max_workers=text_workers)
                full_text = _collapse_repetitions(full_text, max_repeats=3)
                # Remove music/silence hallucinations early (before punctuation restoration)
                full_text, music_removed = _remove_music_hallucinations(full_text)
                if music_removed:
                    print(f"🎵 Removed {music_removed} music/hallucination pattern(s)")
                try:
                    from deepmultilingualpunctuation import PunctuationModel
                    pm_ultra = PunctuationModel()
                    full_text = pm_ultra.restore_punctuation(full_text)
                except Exception as punc_e:
                    print(f"⚠️  Punctuation pre-pass unavailable: {punc_e} — proceeding with ULTRA-only punctuation fixes")
                full_text = ultra_processor.process_text_ultra(full_text, passes=6)
                t1 = time.time()
                print(f"✅ Ultra text processing completed ({t1 - t0:.1f}s)")
                full_text, early_artifact_stats = _remove_extended_artifacts(full_text)
                full_text = _refine_capitalization(_fix_whisper_artifacts(full_text))
                full_text = _collapse_sentence_repetitions(full_text, max_repeats=3)
                full_text, global_freq_stats = _limit_global_sentence_frequency(full_text)
                full_text, loop_stats = _detect_and_break_loops(full_text)
                full_text, late_artifact_stats = _remove_extended_artifacts(full_text)
                quality_stats = {
                    "early_artifacts": early_artifact_stats,
                    "global_frequency": global_freq_stats,
                    "loop_detection": loop_stats,
                    "late_artifacts": late_artifact_stats,
                    "music_hallucinations_removed": music_removed,
                }
                try:
                    paragraph_formatter = create_advanced_paragraph_formatter(max_workers=text_workers)
                    formatted_text = paragraph_formatter.format_paragraphs_advanced(full_text, target_length=600)
                except Exception as pf_e:
                    print(f"⚠️  Advanced paragraph formatter failed: {pf_e} — using basic split_into_paragraphs")
                    formatted = split_into_paragraphs(full_text, max_length=600)
                    formatted_text = "\n\n".join(formatted) if isinstance(formatted, list) else full_text
            except ImportError:
                print("⚠️  Ultra processor not available, falling back to enhanced processor")
                if _enhanced_processor_available and create_enhanced_processor is not None:
                    processor = create_enhanced_processor(use_spacy=True, use_transformers=False)
                    t0 = time.time()
                    full_text = processor.restore_punctuation(full_text)
                    full_text = processor.restore_punctuation(full_text)
                    full_text = processor.restore_punctuation(full_text)
                    t1 = time.time()
                    print(f"✅ Enhanced punctuation restoration completed (3 passes | {t1 - t0:.1f}s)")
                    full_text, early_artifact_stats = _remove_extended_artifacts(full_text)
                    full_text = _refine_capitalization(_fix_whisper_artifacts(full_text))
                    full_text = _collapse_sentence_repetitions(full_text, max_repeats=3)
                    full_text, global_freq_stats = _limit_global_sentence_frequency(full_text)
                    full_text, loop_stats = _detect_and_break_loops(full_text)
                    full_text, late_artifact_stats = _remove_extended_artifacts(full_text)
                    quality_stats = {
                        "early_artifacts": early_artifact_stats,
                        "global_frequency": global_freq_stats,
                        "loop_detection": loop_stats,
                        "late_artifacts": late_artifact_stats,
                    }
                    formatted = split_into_paragraphs(full_text, max_length=600)
                    if isinstance(formatted, list):
                        formatted_text = "\n\n".join(formatted)
                    else:
                        formatted_text = full_text
                else:
                    from deepmultilingualpunctuation import PunctuationModel
                    pm = PunctuationModel()
                    t0 = time.time()
                    full_text = pm.restore_punctuation(full_text)
                    full_text = pm.restore_punctuation(full_text)
                    full_text = pm.restore_punctuation(full_text)
                    full_text = pm.restore_punctuation(full_text)
                    t1 = time.time()
                    print(f"✅ Basic punctuation restoration completed (4 passes | {t1 - t0:.1f}s)")
                    full_text, early_artifact_stats = _remove_extended_artifacts(full_text)
                    full_text = _refine_capitalization(_fix_whisper_artifacts(full_text))
                    full_text = _collapse_sentence_repetitions(full_text, max_repeats=3)
                    full_text, global_freq_stats = _limit_global_sentence_frequency(full_text)
                    full_text, loop_stats = _detect_and_break_loops(full_text)
                    full_text, late_artifact_stats = _remove_extended_artifacts(full_text)
                    quality_stats = {
                        "early_artifacts": early_artifact_stats,
                        "global_frequency": global_freq_stats,
                        "loop_detection": loop_stats,
                        "late_artifacts": late_artifact_stats,
                    }
                    formatted = split_into_paragraphs(full_text, max_length=600)
                    if isinstance(formatted, list):
                        formatted_text = "\n\n".join(formatted)
                    else:
                        formatted_text = full_text
    except Exception as e:
        print(f"⚠️  Text processing failed: {e}")
        formatted_text = full_text
        quality_stats = {}

    # Refine capitalization to fix artifacts (applies to all processing paths)
    try:
        if not _is_verbatim_mode():
            # Fix Whisper-specific artifacts and capitalization only in enhanced mode
            formatted_text = _fix_whisper_artifacts(formatted_text)
            formatted_text = _refine_capitalization(formatted_text)
            print("✅ Capitalization & artifact refinement completed")
    except Exception as e:
        print(f"⚠️  Capitalization refinement failed: {e}")

    # Final quality check and validation (skip in verbatim)
    try:
        if not _is_verbatim_mode():
            if formatted_text and len(formatted_text) > 10:
                if not formatted_text.endswith(('.', '!', '?')):
                    formatted_text += '.'
                if formatted_text and not formatted_text[0].isupper():
                    formatted_text = formatted_text[0].upper() + formatted_text[1:]
                print("✅ Final text validation completed")
            else:
                print("⚠️  Formatted text too short, using original")
                formatted_text = full_text
    except Exception as e:
        print(f"⚠️  Text validation failed: {e}")
        formatted_text = full_text

    base_name = os.path.splitext(os.path.basename(input_path))[0]
    txt_path = os.path.join(output_dir, f"{base_name}.txt")
    docx_path = os.path.join(output_dir, f"{base_name}.docx")
    quality_path = os.path.join(output_dir, f"{base_name}_quality_report.json")

    # Save TXT
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(formatted_text)
        print(f"✅ Text file saved: {txt_path}")
    except Exception as e:
        print(f"❌ Failed to save text file: {e}")
        txt_path = None

    # Save DOCX with fallback
    try:
        doc = Document()
        doc.add_heading(f'{base_name}', 0)
        
        # Add model and location info (use original selection for clarity)
        parent_folder = os.path.basename(os.path.dirname(input_path))
        doc.add_paragraph(f'Model: {original_model_selection}')
        doc.add_paragraph(f'Folder: {parent_folder}')
        doc.add_paragraph('')
        
        elapsed_total = time.time() - start_time
        if os.environ.get("TRANSCRIBE_HIDE_TIME", "").lower() not in ("1", "true", "yes"):
            doc.add_paragraph(f'Transcription time: {format_duration_hms(elapsed_total)}')
            doc.add_paragraph('')
        for para in formatted_text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(docx_path)
        print(f"✅ Word document saved: {docx_path}")
    except Exception as e:
        print(f"❌ Failed to create Word document: {e}")
        try:
            doc = Document()
            doc.add_heading(f'Transcription: {base_name}', 0)
            
            # Add model and location info (fallback - use original selection)
            parent_folder = os.path.basename(os.path.dirname(input_path))
            doc.add_paragraph(f'Model: {original_model_selection}')
            doc.add_paragraph(f'Folder: {parent_folder}')
            doc.add_paragraph('')
            
            doc.add_paragraph(formatted_text[:5000])
            doc.save(docx_path)
            print(f"✅ Basic Word document saved: {docx_path}")
        except Exception as e2:
            print(f"❌ Failed to save even basic Word document: {e2}")
            docx_path = None

    # Final stats
    elapsed = time.time() - start_time
    print("\n🎉 TRANSCRIPTION COMPLETE!")
    print(f"📄 Text file: {txt_path}")
    print(f"📄 Word document: {docx_path}")
    print(f"⏱️  Total time: {format_duration(elapsed)}")
    try:
        base_quality = _summarize_quality(formatted_text, {"pipeline": quality_stats})
        with open(quality_path, "w", encoding="utf-8") as qf:
            json.dump(base_quality, qf, indent=2)
        print(f"📊 Quality report saved: {quality_path}")
    except Exception as qerr:
        print(f"⚠️  Failed to save quality report: {qerr}")

    # Cleanup caches (do not touch torch modules)
    force_gpu_memory_cleanup()

    # Print memory status after cleanup
    import psutil as _ps
    mem = _ps.virtual_memory()
    if torch_api is not None and torch_api.cuda.is_available():
        try:
            gpu_after = torch_api.cuda.memory_allocated() / (1024 ** 3)
            print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available, GPU {gpu_after:.1f}GB used")
        except Exception:
            print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available")
    else:
        print(f"📊 Memory after cleanup: RAM {mem.available / (1024**3):.1f}GB available")

    # Remove temporary preprocessed file if used
    try:
        if 'preprocessing_used' in locals() and preprocessing_used and working_input_path != input_path:
            if os.path.exists(working_input_path):
                os.remove(working_input_path)
                print("🧹 Removed temporary preprocessed audio file")
    except Exception as _cleanup_e:
        print(f"⚠️  Failed to remove temporary file: {_cleanup_e}")

    return txt_path


def transcribe_file_optimised(input_path, model_name="medium", output_dir=None, force_optimised=True, *, threads_override: Optional[int] = None):
    """Compatibility wrapper. Uses the simple auto path."""
    return transcribe_file_simple_auto(input_path, output_dir=output_dir, threads_override=threads_override)


def main():
    parser = argparse.ArgumentParser(description="Simplified auto-detected transcription")
    parser.add_argument("--input", help="Input audio/video file (omit when using --postprocess-only)")
    parser.add_argument("--output-dir", help="Output directory (default: same directory as input file)")
    parser.add_argument("--threads", type=int, help="Override CPU threads for PyTorch/OMP/MKL")
    parser.add_argument("--ram-gb", type=float, help="Cap usable system RAM in GB (env TRANSCRIBE_RAM_GB)")
    parser.add_argument("--ram-frac", "--ram-fraction", dest="ram_fraction", type=float, help="Cap usable system RAM as fraction 0-1 (env TRANSCRIBE_RAM_FRACTION)")
    parser.add_argument("--vram-gb", type=float, help="Cap usable CUDA VRAM in GB (env TRANSCRIBE_VRAM_GB)")
    parser.add_argument("--vram-frac", "--vram-fraction", dest="vram_fraction", type=float, help="Cap usable CUDA VRAM as fraction 0-1 (env TRANSCRIBE_VRAM_FRACTION)")
    parser.add_argument("--vad", action="store_true", help="Enable VAD segmentation for parallel processing performance boost (env TRANSCRIBE_VAD)")
    parser.add_argument("--postprocess-only", help="Existing transcript TXT file to post-process (skips audio decoding)")
    args = parser.parse_args()

    # Post-process only mode -------------------------------------------------
    if args.postprocess_only:
        src = args.postprocess_only
        if not os.path.isfile(src):
            print(f"Error: postprocess file not found: {src}")
            return 1
        out_dir = args.output_dir or os.path.dirname(src)
        os.makedirs(out_dir, exist_ok=True)
        raw_text = open(src, "r", encoding="utf-8", errors="ignore").read()
        before_quality = _summarize_quality(raw_text, {"stage": "before"})

        if _is_verbatim_mode():
            processed = raw_text
            after_quality = _summarize_quality(processed, {"stage": "verbatim"})
        else:
            # Apply pipeline (mirror transcription post-processing w/out punctuation restoration model)
            processed = raw_text
            processed = _collapse_repetitions(processed, max_repeats=3)
            processed = _remove_prompt_artifacts(processed)
            processed, early_artifacts = _remove_extended_artifacts(processed)
            processed = _fix_whisper_artifacts(processed)
            processed = _refine_capitalization(processed)
            processed = _collapse_sentence_repetitions(processed, max_repeats=3)
            processed, global_freq = _limit_global_sentence_frequency(processed)
            processed, loop_stats = _detect_and_break_loops(processed)
            processed, late_artifacts = _remove_extended_artifacts(processed)
            after_quality = _summarize_quality(processed, {
                "stage": "after",
                "early_artifacts": early_artifacts,
                "global_frequency": global_freq,
                "loop_detection": loop_stats,
                "late_artifacts": late_artifacts,
            })
        base = os.path.splitext(os.path.basename(src))[0]
        out_txt = os.path.join(out_dir, f"{base}_postprocessed.txt")
        out_json = os.path.join(out_dir, f"{base}_quality_compare.json")
        try:
            with open(out_txt, "w", encoding="utf-8") as fpp:
                fpp.write(processed)
            with open(out_json, "w", encoding="utf-8") as fj:
                json.dump({"before": before_quality, "after": after_quality}, fj, indent=2)
            print(f"✅ Post-processed transcript saved: {out_txt}")
            print(f"📊 Quality comparison saved: {out_json}")
        except Exception as e:
            print(f"Error saving postprocess outputs: {e}")
            return 1
        return 0

    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        return 1

    # Apply env overrides for RAM/VRAM if provided
    try:
        if getattr(args, "ram_gb", None) is not None:
            os.environ["TRANSCRIBE_RAM_GB"] = str(max(1.0, float(args.ram_gb)))
        if getattr(args, "ram_fraction", None) is not None:
            os.environ["TRANSCRIBE_RAM_FRACTION"] = str(max(0.05, min(1.0, float(args.ram_fraction))))
        if getattr(args, "vram_gb", None) is not None:
            os.environ["TRANSCRIBE_VRAM_GB"] = str(max(0.5, float(args.vram_gb)))
        if getattr(args, "vram_fraction", None) is not None:
            os.environ["TRANSCRIBE_VRAM_FRACTION"] = str(max(0.05, min(1.0, float(args.vram_fraction))))
    except Exception:
        pass

    # Apply VAD override if provided
    try:
        if getattr(args, "vad", False):
            os.environ["TRANSCRIBE_VAD"] = "1"
    except Exception:
        pass

    try:
        transcribe_file_simple_auto(
            args.input,
            output_dir=args.output_dir,
            threads_override=args.threads,
        )
        return 0
    except Exception as e:
        print(f"Error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())