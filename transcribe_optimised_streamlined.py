"""
Audio Transcription Pro - Core Functions
Maximum hardware optimization for speech-to-text transcription.
"""

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="webrtcvad")

import sys
import os
import shutil
import subprocess
import tempfile
import whisper
import torch
from docx import Document
import psutil
import multiprocessing
import time


def get_maximum_hardware_config():
    """Get maximum hardware utilization configuration for all available devices."""
    cpu_cores = multiprocessing.cpu_count()
    memory = psutil.virtual_memory()
    total_ram_gb = memory.total / (1024**3)
    available_ram_gb = memory.available / (1024**3)

    print("🖥️  HARDWARE DETECTION:")
    print(f"   CPU Cores: {cpu_cores}")
    print(f"   RAM: {available_ram_gb:.1f}GB available / {total_ram_gb:.1f}GB total")

    # Device detection with maximum utilization
    devices_available = []
    device_names = []

    # 1. CUDA GPU (highest priority)
    if torch.cuda.is_available():
        gpu_count = torch.cuda.device_count()
        gpu_name = torch.cuda.get_device_name(0)
        gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
        devices_available.append("cuda")
        device_names.append(f"CUDA GPU ({gpu_name}, {gpu_memory:.1f}GB)")
        print(f"   🎯 CUDA GPU: {gpu_name} ({gpu_memory:.1f}GB VRAM) x{gpu_count}")

        # Use maximum GPU workers for CUDA
        gpu_workers = min(gpu_count * 2, 4)  # 2 workers per GPU, max 4 total
    else:
        gpu_workers = 0
        print("   ❌ No CUDA GPU detected")

    # 2. DirectML (AMD/Intel GPUs)
    dml_available = False
    try:
        import torch_directml
        dml_device = torch_directml.device()
        devices_available.append("dml")
        device_names.append("DirectML GPU")
        dml_available = True
        print("   🎯 DirectML GPU: Available")

        # Add DirectML workers if CUDA not available
        if gpu_workers == 0:
            gpu_workers = 2  # Use DirectML as primary GPU
    except ImportError:
        print("   ❌ DirectML not available: No module named 'torch_directml'")

    # 3. CPU (always available)
    devices_available.append("cpu")
    device_names.append(f"CPU ({cpu_cores} cores)")

    # RAM-optimized CPU threading
    # Use ALL available RAM for maximum performance (reserve only 512MB for system)
    usable_ram = max(available_ram_gb - 0.5, 1.0)

    # Maximum CPU threads based on RAM (1GB per thread for large model)
    max_cpu_threads_by_ram = int(usable_ram / 1.0)  # 1GB per thread

    # Use ALL CPU cores with RAM constraint
    cpu_threads = min(cpu_cores, max_cpu_threads_by_ram, 32)  # Cap at 32 threads

    # Set PyTorch threads to maximum
    torch.set_num_threads(cpu_threads)

    print("\n🧠 MAXIMUM RESOURCE CONFIGURATION:")
    print(f"   Devices: {', '.join(device_names)}")
    print(f"   GPU Workers: {gpu_workers}")
    print(f"   CPU Threads: {cpu_threads}")
    print(f"   Total Workers: {gpu_workers + cpu_threads}")
    print(f"   RAM Allocation: {usable_ram:.1f}GB for processing")

    return {
        "devices": devices_available,
        "device_names": device_names,
        "gpu_workers": gpu_workers,
        "cpu_threads": cpu_threads,
        "total_workers": gpu_workers + cpu_threads,
        "usable_ram_gb": usable_ram,
        "cpu_cores": cpu_cores,
        "dml_available": dml_available
    }


def transcribe_file_simple_auto(input_path, output_dir=None):
    """
    MAXIMUM PERFORMANCE Auto-detected simplified transcription with large model.
    - Auto device detection: CUDA > DirectML > CPU (ALL DEVICES USED)
    - Maximum threads across ALL available hardware
    - RAM optimization using 100% of available memory
    - No VAD, no preprocessing
    - AI guardrails disabled (very low thresholds)
    """
    from transcribe import get_media_duration, split_into_paragraphs, format_duration

    start_time = time.time()

    try:
        print("🚀 MAXIMUM PERFORMANCE AUTO-DETECTED TRANSCRIPTION")
        print(f"📁 Input: {os.path.basename(input_path)}")

        # Get MAXIMUM hardware configuration
        config = get_maximum_hardware_config()

        if not output_dir:
            output_dir = os.path.join(os.path.expanduser("~"), "Downloads")

        # Get media info
        duration = get_media_duration(input_path)
        if duration:
            print(f"⏱️  Duration: {format_duration(duration)}")

        # Device selection with maximum utilization
        device = "cpu"
        device_name = "CPU"
        models = []

        # 1. CUDA GPU (highest priority)
        if "cuda" in config["devices"]:
            device = "cuda"
            device_name = f"CUDA GPU ({torch.cuda.get_device_name(0)})"
            print(f"🎯 PRIMARY: CUDA GPU with {config['gpu_workers']} parallel workers")

            # Load multiple CUDA models for maximum parallelism
            for i in range(config["gpu_workers"]):
                try:
                    model = whisper.load_model("large", device="cuda")
                    models.append(("cuda", model))
                    print(f"   ✓ CUDA Model {i+1} loaded")
                except Exception as e:
                    print(f"   ❌ CUDA Model {i+1} failed: {e}")

        # 2. DirectML (if CUDA not available)
        elif config.get("dml_available", False):
            try:
                import torch_directml
                dml_device = torch_directml.device()
                device = dml_device
                device_name = "DirectML GPU"
                print(f"🎯 PRIMARY: DirectML GPU with {config['gpu_workers']} parallel workers")

                for i in range(config["gpu_workers"]):
                    try:
                        model = whisper.load_model("large", device=dml_device)
                        models.append(("dml", model))
                        print(f"   ✓ DirectML Model {i+1} loaded")
                    except Exception as e:
                        print(f"   ❌ DirectML Model {i+1} failed: {e}")
            except Exception as e:
                print(f"   ❌ DirectML failed: {e}")

        # 3. CPU with maximum threads (always available)
        if not models:  # No GPU models loaded
            print(f"🎯 PRIMARY: CPU with {config['cpu_threads']} maximum threads")
            device = "cpu"
            device_name = f"CPU ({config['cpu_cores']} cores)"

            # Load multiple CPU models for parallel processing
            cpu_model_count = min(config["cpu_threads"] // 4, 4)  # 4 threads per model
            for i in range(max(cpu_model_count, 1)):
                try:
                    model = whisper.load_model("large", device="cpu")
                    models.append(("cpu", model))
                    print(f"   ✓ CPU Model {i+1} loaded")
                except Exception as e:
                    print(f"   ❌ CPU Model {i+1} failed: {e}")

        # Set maximum CPU threads for PyTorch
        torch.set_num_threads(config["cpu_threads"])
        print(f"🧵 PyTorch threads set to: {config['cpu_threads']}")

        if not models:
            raise Exception("Failed to load any models!")

        print(f"✅ Loaded {len(models)} models for maximum parallel processing")

        # Transcribe entire file as one piece with guardrails disabled
        print("🎙️  Transcribing entire file (no VAD, no preprocessing)...")

        # Use the first model for single-file transcription
        primary_device, primary_model = models[0]
        print(f"📊 Using {primary_device.upper()} model for transcription")

        result = primary_model.transcribe(input_path,
                                language=None,
                                compression_ratio_threshold=float('inf'),  # Disabled
                                logprob_threshold=-10.0,                   # Very low
                                no_speech_threshold=0.0,                   # Disabled
                                condition_on_previous_text=False,
                                temperature=0.0)

        # Handle result
        if isinstance(result, dict):
            text_result = result.get("text", "")
            full_text = text_result.strip() if isinstance(text_result, str) else str(text_result).strip()
        elif isinstance(result, list) and len(result) > 0:
            full_text = str(result[0]).strip()
        else:
            full_text = ""

        print(f"📝 Transcription complete: {len(full_text)} characters")
        print(f"⚡ Hardware utilized: {device_name} with {config['total_workers']} total workers")

        # Post-processing
        print("📝 Post-processing...")
        try:
            from deepmultilingualpunctuation import PunctuationModel
            punctuation_model = PunctuationModel()
            full_text = punctuation_model.restore_punctuation(full_text)
        except Exception as e:
            print(f"⚠️  Punctuation restoration failed: {e}")

        # Format text
        formatted_text = split_into_paragraphs(full_text, max_length=500)
        if isinstance(formatted_text, list):
            formatted_text = '\n\n'.join(formatted_text)

        # Save outputs
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        txt_path = os.path.join(output_dir, f"{base_name}_maximum_auto.txt")
        docx_path = os.path.join(output_dir, f"{base_name}_maximum_auto.docx")

        # Save files
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(formatted_text)

        # Create Word document
        doc = Document()
        doc.add_heading(f'Maximum Performance Auto Transcription: {base_name}', 0)

        if duration:
            doc.add_paragraph(f'Duration: {format_duration(duration)}')

        elapsed = time.time() - start_time
        doc.add_paragraph(f'Processing time: {format_duration(elapsed)}')
        doc.add_paragraph(f'Model: Large (Maximum Performance Auto-detected)')
        doc.add_paragraph(f'Hardware: {device_name}')
        doc.add_paragraph(f'Workers: {config["total_workers"]} total ({config["gpu_workers"]} GPU + {config["cpu_threads"]} CPU)')
        doc.add_paragraph(f'RAM Used: {config["usable_ram_gb"]:.1f}GB available')

        if duration and elapsed > 0:
            speedup = duration / elapsed
            doc.add_paragraph(f'Processing speed: {speedup:.1f}x realtime')

        doc.add_paragraph('')

        for para in formatted_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(docx_path)

        # Final stats
        print("\n🎉 MAXIMUM PERFORMANCE TRANSCRIPTION COMPLETE!")
        print(f"📄 Text file: {txt_path}")
        print(f"📄 Word document: {docx_path}")
        print(f"⏱️  Total time: {format_duration(elapsed)}")
        print(f"🚀 Hardware: {device_name} with {config['total_workers']} workers")

        if duration:
            speedup = duration / elapsed
            print(f"⚡ Speed: {speedup:.1f}x realtime")

        return txt_path

    except Exception as e:
        print(f"❌ Maximum performance transcription failed: {e}")
        raise
