"""Speech to Text Transcription Tool v1.0Beta - WhisperCPP Version

A clean, professional GUI for converting audio and video files to text using WhisperCPP.
- Turbo v3 model for fast CPU transcription
- No GPU required, runs on any CPU
- Clean, modern interface
"""
import argparse
import os
import threading
import queue
import sys
import subprocess
from typing import Optional, List

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
except Exception:
    tk = None

# Supported file extensions for batch processing
SUPPORTED_EXTS = (
    ".mp3", ".wav", ".flac", ".m4a", ".aac", ".ogg", ".wma",
    ".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm"
)

# Default Downloads folder retained for compatibility
DEFAULT_DOWNLOADS = os.path.normpath(os.path.join(os.path.expanduser("~"), "Downloads"))

# WhisperCPP Configuration
WHISPER_CPP_EXE = "whisper-cli.exe"
MODEL_PATH = "models/ggml-large-v3-turbo.bin"
PUNCTUATION_MODEL = "punctuation.onnx"

def run_whispercpp_transcription(input_file: str, outdir: Optional[str], output_queue: queue.Queue):
    """Run transcription for a single file using WhisperCPP with turbo v3 model."""
    try:
        # Determine output directory
        target_outdir = outdir if outdir else os.path.dirname(input_file)
        base_name = os.path.splitext(os.path.basename(input_file))[0]

        output_queue.put("Starting WhisperCPP transcription for: {}\n".format(os.path.basename(input_file)))
        output_queue.put("Using Turbo v3 model (fast CPU processing)\n")

        # Build command for WhisperCPP
        cmd = [
            WHISPER_CPP_EXE,
            "-m", MODEL_PATH,
            "-f", input_file,
            "-otxt",
            "-of", os.path.join(target_outdir, base_name)
        ]

        # Run WhisperCPP
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=os.getcwd())

        if result.returncode == 0:
            txt_file = os.path.join(target_outdir, f"{base_name}.txt")
            if os.path.exists(txt_file):
                output_queue.put("Transcription complete! Files saved.\n")
                output_queue.put("  Text file: {}\n".format(os.path.basename(txt_file)))

                # Try to create DOCX version
                try:
                    from docx import Document
                    docx_file = os.path.join(target_outdir, f"{base_name}.docx")
                    doc = Document()
                    doc.add_heading("Speech to Text Transcription", 0)

                    with open(txt_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                        doc.add_paragraph(content)

                    doc.save(docx_file)
                    output_queue.put("  Word document: {}\n".format(os.path.basename(docx_file)))
                except Exception as e:
                    output_queue.put("  (DOCX creation failed: {})\n".format(e))

            else:
                output_queue.put("Transcription completed but output file not found.\n")
        else:
            output_queue.put("Transcription failed: {}\n".format(result.stderr))

    except Exception as e:
        output_queue.put("Transcription failed: {}\n".format(e))


def run_batch_whispercpp_transcription(paths: List[str], outdir_override: Optional[str], output_queue: queue.Queue):
    """Run batch transcription sequentially over provided file paths using WhisperCPP."""
    total = len(paths)
    output_queue.put("Batch mode: {} eligible files queued.\n".format(total))
    successful = 0
    failed = 0

    for idx, p in enumerate(paths, start=1):
        output_queue.put("\n[{} / {}] Processing: {}\n".format(idx, total, os.path.basename(p)))
        try:
            run_whispercpp_transcription(p, outdir_override, output_queue)
            successful += 1
        except Exception as e:
            failed += 1
            output_queue.put("Failed to process '{}': {}\n".format(os.path.basename(p), str(e)))
            output_queue.put("Continuing with next file...\n")

    output_queue.put("\nBatch processing complete!\n")
    output_queue.put("Successfully processed: {} files\n".format(successful))
    if failed > 0:
        output_queue.put("Failed: {} files\n".format(failed))


class QueueWriter:
    """A file-like writer that puts text into a queue."""
    def __init__(self, q: queue.Queue):
        self.q = q

    def write(self, msg):
        if msg:
            self.q.put(str(msg))

    def flush(self):
        pass


def launch_whispercpp_gui(default_outdir: Optional[str] = None):
    if tk is None:
        print("Tkinter is not available in this Python build.")
        return

    # Check if WhisperCPP is available
    if not os.path.exists(WHISPER_CPP_EXE):
        print(f"WhisperCPP executable not found: {WHISPER_CPP_EXE}")
        print("Please run the installer again to download WhisperCPP binaries.")
        return

    if not os.path.exists(MODEL_PATH):
        print(f"WhisperCPP model not found: {MODEL_PATH}")
        print("Please run the installer again to download the model.")
        return

    root = tk.Tk()
    root.title("Speech to Text Transcription Tool v1.0Beta - WhisperCPP")
    root.geometry("800x750")
    root.minsize(700, 650)
    root.resizable(True, True)
    root.configure(bg='#f8f9fa')

    # Center the window on screen
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry('{}x{}+{}+{}'.format(width, height, x, y))

    # Configure grid weights for proper resizing
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    # Clean modern styling
    style = ttk.Style()
    style.configure('Clean.TFrame', background='#f8f9fa')
    style.configure('Clean.TLabel', background='#f8f9fa', foreground='#2c3e50', font=('Segoe UI', 11))
    style.configure('Title.TLabel', font=('Segoe UI', 20, 'bold'), foreground='#1a365d', background='#f8f9fa')
    style.configure('Subtitle.TLabel', font=('Segoe UI', 10), foreground='#64748b', background='#f8f9fa')
    style.configure('Section.TLabel', font=('Segoe UI', 13, 'bold'), foreground='#374151', background='#f8f9fa')

    mainframe = ttk.Frame(root, style='Clean.TFrame', padding="30 30 30 30")
    mainframe.grid(column=0, row=0, sticky="nsew")
    mainframe.columnconfigure(1, weight=1)
    mainframe.rowconfigure(6, weight=1)
    mainframe.rowconfigure(7, minsize=60)

    # Clean title section
    title_frame = ttk.Frame(mainframe, style='Clean.TFrame')
    title_frame.grid(column=0, row=0, columnspan=3, pady=(0, 30), sticky="ew")
    title_frame.columnconfigure(1, weight=1)

    title_label = ttk.Label(title_frame, text="Speech to Text Transcription Tool", style='Title.TLabel')
    title_label.grid(column=0, row=0, sticky="w")

    subtitle_label = ttk.Label(title_frame,
                              text="v1.0Beta - WhisperCPP Turbo v3 (CPU)",
                              style='Subtitle.TLabel')
    subtitle_label.grid(column=0, row=1, sticky="w", pady=(8, 0))

    # Clean file/folder selection section
    file_section_label = ttk.Label(mainframe, text="File Selection", style='Section.TLabel')
    file_section_label.grid(column=0, row=1, columnspan=3, sticky="w", pady=(0, 15))

    # File input area with subtle background
    file_frame = tk.Frame(mainframe, bg='white', relief='flat', borderwidth=1)
    file_frame.grid(column=0, row=2, columnspan=3, sticky="ew", pady=(0, 25))
    file_frame.columnconfigure(1, weight=1)

    ttk.Label(file_frame, text="Audio/Video File or Folder:", background='white', foreground='#374151',
             font=('Segoe UI', 10, 'bold')).grid(column=0, row=0, sticky="w", padx=20, pady=(20, 10))

    input_var = tk.StringVar()
    input_entry = tk.Entry(file_frame, textvariable=input_var, font=('Segoe UI', 10),
                          relief='flat', borderwidth=1, bg='white', fg='#2c3e50',
                          insertbackground='#2c3e50')
    input_entry.grid(column=0, row=1, columnspan=2, sticky="ew", padx=20, pady=(0, 20))

    def browse_input():
        filetypes = [
            ("Audio files", "*.mp3 *.wav *.flac *.m4a *.aac *.ogg *.wma"),
            ("Video files", "*.mp4 *.avi *.mkv *.mov *.wmv *.flv *.webm"),
            ("All files", "*.*")
        ]
        p = filedialog.askopenfilename(title="Select audio/video file", filetypes=filetypes)
        if p:
            input_var.set(p)
            status_label.config(text="Selected file: {}".format(os.path.basename(p)), foreground='#059669')

    def browse_folder():
        d = filedialog.askdirectory(title="Select folder for batch processing")
        if d:
            input_var.set(d)
            try:
                files = [f for f in os.listdir(d) if os.path.splitext(f)[1].lower() in SUPPORTED_EXTS]
                status_label.config(
                    text="Selected folder: {} ({} eligible files)".format(os.path.basename(d), len(files)),
                    foreground='#059669'
                )
            except Exception:
                status_label.config(text="Selected folder: {}".format(os.path.basename(d)), foreground='#059669')

    browse_btn = tk.Button(file_frame, text="Browse File", command=browse_input,
                          font=('Segoe UI', 10, 'bold'), bg='#007acc', fg='white',
                          relief='flat', borderwidth=0, padx=20, pady=8,
                          activebackground='#0056b3', activeforeground='white')
    browse_btn.grid(column=2, row=1, sticky="e", padx=(10, 10))

    browse_folder_btn = tk.Button(file_frame, text="Browse Folder", command=browse_folder,
                          font=('Segoe UI', 10, 'bold'), bg='#0ea5e9', fg='white',
                          relief='flat', borderwidth=0, padx=20, pady=8,
                          activebackground='#0284c7', activeforeground='white')
    browse_folder_btn.grid(column=2, row=2, sticky="e", padx=(10, 20), pady=(0, 20))

    # Status label
    status_label = ttk.Label(file_frame, text="No file or folder selected",
                            font=('Segoe UI', 9), foreground='#6b7280', background='white')
    status_label.grid(column=0, row=3, columnspan=3, sticky="w", padx=20, pady=(0, 10))

    # Clean settings section
    settings_section_label = ttk.Label(mainframe, text="Settings", style='Section.TLabel')
    settings_section_label.grid(column=0, row=3, columnspan=3, sticky="w", pady=(0, 15))

    # Settings area with subtle background
    settings_frame = tk.Frame(mainframe, bg='white', relief='flat', borderwidth=1)
    settings_frame.grid(column=0, row=4, columnspan=3, sticky="ew", pady=(0, 25))

    settings_text = """WhisperCPP Turbo v3 Model
CPU-Only Processing (No GPU Required)
Fast Transcription Speeds
Output Directory: Same as source file(s)
Full Audio Capture (All Content Preserved)"""

    settings_label = tk.Text(settings_frame, height=6, wrap=tk.WORD, font=('Segoe UI', 10),
                            bg='white', fg='#374151', borderwidth=0, highlightthickness=0,
                            padx=20, pady=20)
    settings_label.insert(tk.END, settings_text)
    settings_label.config(state=tk.DISABLED)
    settings_label.grid(column=0, row=0, sticky="ew")

    # Log area with modern design
    log_section_label = ttk.Label(mainframe, text="Activity Log", style='Section.TLabel')
    log_section_label.grid(column=0, row=5, columnspan=3, sticky="w", pady=(0, 15))

    log_frame = tk.Frame(mainframe, bg='white', relief='flat', borderwidth=1)
    log_frame.grid(column=0, row=6, columnspan=3, sticky="nsew", pady=(0, 25))
    log_frame.columnconfigure(0, weight=1)
    log_frame.rowconfigure(0, weight=1)

    log_text = tk.Text(log_frame, wrap=tk.WORD, height=25, font=('Segoe UI', 9),
                      bg='white', fg='#2c3e50', borderwidth=0, highlightthickness=0,
                      insertbackground='#2c3e50')
    log_text.grid(column=0, row=0, sticky="nsew")

    log_scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=log_text.yview)
    log_scroll.grid(column=1, row=0, sticky="ns")
    log_text['yscrollcommand'] = log_scroll.set
    log_text.configure(state='disabled')

    # Add initial welcome message
    log_text.configure(state='normal')
    log_text.insert('end', "Welcome to Speech to Text Transcription Tool v1.0Beta!\n", "title")
    log_text.insert('end', "WhisperCPP Turbo v3 - Fast CPU transcription\n\n")
    log_text.insert('end', "Select an audio or video file and click 'Start Transcription' to begin.\n\n")
    log_text.insert('end', "Tip: This version runs entirely on CPU and is optimized for speed.\n\n")
    log_text.configure(state='disabled')

    # Tag configuration for colored text
    log_text.tag_configure("title", foreground="#007acc", font=('Segoe UI', 9, 'bold'))
    log_text.tag_configure("success", foreground="#4ade80")
    log_text.tag_configure("error", foreground="#ef4444")
    log_text.tag_configure("info", foreground="#60a5fa")

    q = queue.Queue()

    def poll_queue():
        try:
            while True:
                msg = q.get_nowait()
                formatted_msg = format_log_message(msg)
                log_text.configure(state='normal')
                log_text.insert('end', formatted_msg)
                log_text.see('end')
                log_text.configure(state='disabled')
        except queue.Empty:
            pass
        except Exception as e:
            import sys
            print("DEBUG: poll_queue exception: {}".format(e), file=sys.__stderr__)
        root.after(200, poll_queue)

    def format_log_message(msg):
        """Format log messages for better readability in the GUI."""
        if not msg.strip():
            return msg

        formatted = msg

        if any(msg.startswith(prefix) for prefix in ['Starting', 'Using', 'Transcription']):
            formatted = "{}\n".format(msg)
        elif any(word in msg for word in ['Config:', 'RAM:', 'GPU:', 'CPU:', 'Workers:', 'System:']):
            formatted = "  {}".format(msg)
        elif any(char in msg for char in ['%']) or 'segments' in msg or 'Speed:' in msg:
            formatted = "    {}".format(msg)
        elif any(prefix in msg for prefix in ['Transcription failed', 'Error']):
            formatted = "\nError: {}\n".format(msg)
        elif any(prefix in msg for prefix in ['complete', 'Complete']) or 'saved' in msg.lower():
            formatted = "\nSuccess: {}".format(msg)
        elif any(ext in msg for ext in ['.txt', '.docx', '.mp3', '.wav']):
            formatted = "    {}".format(msg)
        elif any(word in msg for word in ['Model', 'Loaded', 'Loading']):
            formatted = "      {}".format(msg)
        elif any(word in msg for word in ['Duration:', 'Time:', 'seconds', 'minutes']):
            formatted = "    {}".format(msg)

        return formatted

    def start_transcription_thread():
        inp = input_var.get().strip()
        if not inp or not os.path.exists(inp):
            error_msg = "Please select a valid file or folder.\nPath: '{}'\nExists: {}".format(inp, os.path.exists(inp) if inp else False)
            messagebox.showerror("Input Required", error_msg)
            return

        # Clear log and show starting message
        log_text.configure(state='normal')
        log_text.delete(1.0, 'end')
        log_text.insert('end', "Starting WhisperCPP transcription...\n")
        log_text.see('end')
        log_text.configure(state='disabled')

        # Disable run button during processing
        run_btn.configure(state='disabled')

        def worker():
            try:
                q.put("Initializing WhisperCPP transcription process...\n")

                if os.path.isdir(inp):
                    files = []
                    try:
                        for name in sorted(os.listdir(inp)):
                            full = os.path.join(inp, name)
                            if os.path.isfile(full) and os.path.splitext(name)[1].lower() in SUPPORTED_EXTS:
                                files.append(full)
                    except Exception as e:
                        q.put("Failed to list folder '{}': {}\n".format(inp, e))

                    if not files:
                        q.put("No supported media files found in the selected folder.\n")
                    else:
                        run_batch_whispercpp_transcription(files, outdir_override=None, output_queue=q)
                else:
                    run_whispercpp_transcription(inp, outdir=None, output_queue=q)

                q.put("WhisperCPP transcription process finished!\n")

            except Exception as e:
                import traceback
                error_msg = "Worker error: {}\n{}".format(e, traceback.format_exc())
                q.put(error_msg)
            finally:
                def enable_btn():
                    run_btn.configure(state='normal')
                root.after(0, enable_btn)

        t = threading.Thread(target=worker, daemon=True)
        t.start()

    # Run button
    run_btn = tk.Button(mainframe, text="Start Transcription", command=start_transcription_thread,
                       font=('Segoe UI', 12, 'bold'), bg='#007acc', fg='white',
                       relief='flat', borderwidth=0, padx=30, pady=12,
                       activebackground='#0056b3', activeforeground='white',
                       cursor='hand2')
    run_btn.grid(column=1, row=7, pady=(10, 0))

    poll_queue()
    root.mainloop()


def main():
    parser = argparse.ArgumentParser(description="Speech to Text Transcription Tool v1.0Beta - WhisperCPP Version")
    parser.add_argument("--input", help="input audio/video file OR folder (if provided, runs headless)")
    parser.add_argument("--outdir", help="optional output folder override; defaults to saving next to source file(s)")
    args = parser.parse_args()

    outdir = args.outdir or None

    if args.input:
        # Run headless
        q = queue.Queue()

        def runner():
            p = args.input
            if os.path.isdir(p):
                files = [
                    os.path.join(p, name)
                    for name in sorted(os.listdir(p))
                    if os.path.isfile(os.path.join(p, name)) and os.path.splitext(name)[1].lower() in SUPPORTED_EXTS
                ]
                if not files:
                    q.put("No supported media files found in the provided folder.\n")
                else:
                    run_batch_whispercpp_transcription(files, outdir_override=outdir, output_queue=q)
            else:
                run_whispercpp_transcription(p, outdir, q)

        t = threading.Thread(target=runner)
        t.start()

        while t.is_alive() or not q.empty():
            try:
                msg = q.get(timeout=0.2)
                print(msg, end="")
            except queue.Empty:
                pass
        return

    # Launch GUI in detached process
    try:
        subprocess.run(["cmd", "/c", "start", "python", __file__, "--gui"],
                      creationflags=subprocess.DETACHED_PROCESS | subprocess.CREATE_NEW_PROCESS_GROUP)
        print("WhisperCPP GUI launched in background. You can now use the terminal for other commands.")
    except Exception as e:
        print(f"Failed to launch GUI in background: {e}")
        print("Launching GUI in foreground instead...")
        launch_whispercpp_gui(default_outdir=outdir)


if __name__ == "__main__":
    main()
