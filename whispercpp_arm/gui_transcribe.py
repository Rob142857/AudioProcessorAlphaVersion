import os
import sys
import subprocess
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
import threading

try:
    import docx as _docx_module  # lazy actual class import inside function
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Note: punctuation and ONNX runtime disabled for now to reduce startup noise on ARM
# import onnxruntime as ort
# import numpy as np
# from deepmultilingualpunctuation import PunctuationModel


# Use the in-place built binary which has its required DLLs alongside
BASE_DIR = Path(__file__).parent.resolve()
WHISPER_CPP_PATH = str(BASE_DIR / "whisper.cpp" / "build" / "bin" / "whisper-cli.exe")
MODEL_PATH = str(BASE_DIR / "models" / "ggml-large-v3-turbo.bin")  # Absolute model path
PUNCTUATION_MODEL_PATH = "./models/punctuation.onnx"  # Path to ONNX punctuation model

# Helper functions
def run_whisper_cpp(audio_file, log_callback=None):
    import tempfile
    import shutil
    
    # Copy file to temp directory to avoid permission issues
    with tempfile.NamedTemporaryFile(suffix=audio_file.suffix, delete=False) as temp_file:
        shutil.copy2(str(audio_file), temp_file.name)
        temp_path = Path(temp_file.name)
    
    # Outputs: whisper.cpp expects -of as a prefix (without extension)
    base_dir = Path(__file__).parent
    out_dir = base_dir / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_prefix = str(out_dir / audio_file.stem)
    out_txt = str(out_dir / f"{audio_file.stem}.txt")

    cmd = [WHISPER_CPP_PATH, "-m", MODEL_PATH, "-f", str(temp_path), "-otxt", "-of", out_prefix, "--language", "en", "--suppress-non-speech", "--no-speech-thresh", "0.0"]
    exe_dir = Path(WHISPER_CPP_PATH).parent
    env = os.environ.copy()
    # Ensure the binary's folder is in PATH so dependent DLLs (ggml*.dll) are found
    env["PATH"] = str(exe_dir) + os.pathsep + env.get("PATH", "")
    try:
        # Stream stdout live for progress/verbose logs
        proc = subprocess.Popen(
            cmd,
            cwd=str(exe_dir),
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            universal_newlines=True,
        )

        if proc.stdout is not None:
            for line in proc.stdout:
                if not line:
                    continue
                if log_callback:
                    try:
                        log_callback(line.rstrip())
                    except Exception:
                        # Fallback to printing if GUI callback fails
                        print(line.rstrip())
                else:
                    print(line.rstrip())
        code = proc.wait()

        if Path(out_txt).exists() and Path(out_txt).stat().st_size > 0:
            # Consider it a success even if returncode != 0 (some builds return 1 after success)
            return out_txt
        else:
            raise RuntimeError(f"whisper-cli failed with exit code {code} and no output file")
    finally:
        # Clean up temp file
        try:
            temp_path.unlink(missing_ok=True)
        except Exception:
            pass

def run_whisper_python(audio_file):
    raise RuntimeError("Python Whisper fallback disabled on ARM due to missing deps. Use whisper.cpp.")

def restore_punctuation(text_file):
    # For now, skip punctuation restoration as model is not available
    return text_file

def transcribe_files(files, log_callback):
    os.makedirs("output", exist_ok=True)
    for audio_file in files:
        try:
            log_callback(f"Transcribing: {audio_file}")
            txt_file = Path(run_whisper_cpp(audio_file, log_callback=log_callback))
            log_callback(f"Restoring punctuation...")
            punct_file = restore_punctuation(str(txt_file))

            # Create a Word document next to the source file with plain <basename>.docx
            try:
                if DOCX_AVAILABLE and txt_file.exists():
                    docx_path = audio_file.with_suffix(".docx")
                    _create_docx_from_txt(txt_file, docx_path)
                    log_callback(f"Saved Word document: {docx_path}")
                elif not DOCX_AVAILABLE:
                    log_callback("python-docx not installed; skipping .docx export (pip install python-docx)")
                else:
                    log_callback("No text output found; skipping .docx export")
            except Exception as e:
                log_callback(f"DOCX export failed: {e}")

            log_callback(f"Done: {punct_file}")
        except Exception as e:
            log_callback(f"Error: {e}")


import re


def _rich_paragraphs(text: str,
                     max_chars_per_para: int = 800,
                     max_sentences_per_para: int = 5) -> list[str]:
    """Richer paragraphing heuristics:
    - Treat blank lines as hard paragraph breaks.
    - Within a block, split into sentences (., !, ?) while respecting common abbreviations and initials.
    - Group 2–5 sentences or until ~800 chars per paragraph.
    - Preserve bullet/numbered lines as their own paragraphs.
    """
    if not text:
        return []

    # Hard breaks by blank lines
    blocks = re.split(r"\n\s*\n+", text)

    # Common abbreviations to avoid splitting on
    abbreviations = {
        "Mr.", "Mrs.", "Ms.", "Dr.", "Prof.", "Sr.", "Jr.", "St.",
        "vs.", "etc.", "e.g.", "i.e.", "cf.", "Co.", "Corp.", "Inc.", "Ltd.",
        "U.S.", "U.K.", "No.", "Mt.", "Rd.", "Ave.", "Jan.", "Feb.", "Mar.",
        "Apr.", "Aug.", "Sept.", "Oct.", "Nov.", "Dec."
    }

    bullet_re = re.compile(r"^\s*(?:[-*•]\s+|\d+[\.)]\s+)")
    # Sentence split (rough): break after ., !, ? followed by whitespace and a capital/quote/open paren
    sentence_split_re = re.compile(r"(?<=[.!?])[\)]?\"?'?\s+(?=[\"'\(A-Z0-9])")

    def is_abbrev_end(s: str) -> bool:
        s = s.strip()
        if not s:
            return False
        last = s.split()[-1]
        if last in abbreviations:
            return True
        # Single-letter initial (e.g., "J.")
        if re.match(r"^[A-Z]\.$", last):
            return True
        # Ellipses
        if s.endswith("..."):
            return True
        return False

    paragraphs: list[str] = []

    for block in blocks:
        if not block.strip():
            continue

        # If block looks like bullets or numbered list, keep each line as a paragraph
        lines = block.splitlines()
        if any(bullet_re.match(ln) for ln in lines):
            for ln in lines:
                ln = ln.strip()
                if ln:
                    paragraphs.append(ln)
            continue

        # Collapse internal newlines to single spaces for sentence processing
        block_text = re.sub(r"\s+", " ", block.strip())
        # Initial naive split
        parts = sentence_split_re.split(block_text)
        # Merge sentences that were split after abbreviations or initials
        sentences: list[str] = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if sentences and (is_abbrev_end(sentences[-1]) or len(sentences[-1]) <= 2):
                sentences[-1] = (sentences[-1] + " " + part).strip()
            else:
                sentences.append(part)

        # Group into paragraphs by sentence count and char budget
        cur: list[str] = []
        cur_len = 0
        for s in sentences:
            # Start new para if adding would exceed limits
            if cur and (len(s) + cur_len > max_chars_per_para or len(cur) >= max_sentences_per_para):
                paragraphs.append(" ".join(cur).strip())
                cur = []
                cur_len = 0
            cur.append(s)
            cur_len += len(s) + 1
        if cur:
            paragraphs.append(" ".join(cur).strip())

    return paragraphs


def _create_docx_from_txt(txt_path: Path, docx_path: Path):
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read().strip()

    # Richer paragraphing for better readability
    paragraphs = _rich_paragraphs(content, max_chars_per_para=800, max_sentences_per_para=5)
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    # Import inside to avoid static analyser unbound warnings
    from docx import Document as _DocxDocument
    doc = _DocxDocument()
    doc.add_heading("Transcription", 0)
    for p in paragraphs:
        if p:
            doc.add_paragraph(p)
    # Ensure parent exists (should be the source file's folder)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))

# GUI Implementation
class TranscribeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Whisper.cpp ARM64 Transcription")
        self.root.geometry("600x400")  # Set a specific size
        self.file_list = []
        self.create_widgets()
        print("GUI window should now be visible!")

    def create_widgets(self):
        self.frame = tk.Frame(self.root)
        self.frame.pack(padx=10, pady=10)

        self.select_btn = tk.Button(self.frame, text="Select File(s)", command=self.select_files)
        self.select_btn.grid(row=0, column=0, padx=5, pady=5)

        self.select_folder_btn = tk.Button(self.frame, text="Select Folder", command=self.select_folder)
        self.select_folder_btn.grid(row=0, column=1, padx=5, pady=5)

        self.transcribe_btn = tk.Button(self.frame, text="Transcribe", command=self.start_transcription)
        self.transcribe_btn.grid(row=0, column=2, padx=5, pady=5)

        self.log_text = tk.Text(self.frame, width=60, height=20)
        self.log_text.grid(row=1, column=0, columnspan=3, padx=5, pady=5)

    def select_files(self):
        files = filedialog.askopenfilenames(filetypes=[("Audio Files", "*.wav *.mp3 *.m4a")])
        self.file_list = [Path(f) for f in files]
        self.log(f"Selected files: {', '.join(files)}")

    def select_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            p = Path(folder)
            self.file_list = list(p.glob("*.wav")) + list(p.glob("*.mp3")) + list(p.glob("*.m4a"))
            self.log(f"Selected folder: {folder}")
            self.log(f"Found files: {', '.join(str(f) for f in self.file_list)}")

    def start_transcription(self):
        if not self.file_list:
            messagebox.showerror("Error", "No files selected.")
            return
        threading.Thread(target=transcribe_files, args=(self.file_list, self.log), daemon=True).start()

    def log(self, msg):
        self.log_text.insert(tk.END, msg + "\n")
        self.log_text.see(tk.END)

if __name__ == "__main__":
    print("Starting GUI application...")
    try:
        root = tk.Tk()
        print("Tkinter root created")
        app = TranscribeApp(root)
        print("App initialized")
        print("Starting main loop...")
        root.mainloop()
        print("Main loop ended")
    except Exception as e:
        print(f"Error starting GUI: {e}")
        import traceback
        traceback.print_exc()
