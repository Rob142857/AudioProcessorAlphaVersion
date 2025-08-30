"""
Optimised GPU+CPU hybrid transcription with maximum hardware utilisation.
This module provides the most efficient transcription by utilising all available
hardware resources simultaneously for maximum performance.
"""
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="webrtcvad")

import sys
import os
import shutil
import subprocess
import tempfile
import argparse
import whisper
import torch
from docx import Document
from moviepy import VideoFileClip, AudioFileClip
import imageio_ffmpeg as iio_ffmpeg
from deepmultilingualpunctuation import PunctuationModel
try:
    import webrtcvad
    WEBRTCVAD_AVAILABLE = True
except ImportError:
    WEBRTCVAD_AVAILABLE = False
    print("⚠️  webrtcvad not available - Voice Activity Detection disabled")
import re
import time
import json
import multiprocessing
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor, as_completed
import queue
import threading
import psutil
import gc


def adjust_workers_for_model(config, model_name):
    """Adjust worker counts based on actual model size - OPTIMISED RAM utilisation."""
    model_ram_usage = {
        'tiny': 0.5, 'base': 0.8, 'small': 1.5, 'medium': 2.5, 
        'large': 4.0, 'large-v2': 4.0, 'large-v3': 4.0
    }
    
    actual_ram_per_model = model_ram_usage.get(model_name, 2.5)
    # OPTIMISED: Use almost all RAM - reserve only 1GB (was 4GB)
    usable_ram = max(config["available_ram_gb"] - 1.0, 2.0)
    
    # Recalculate CPU workers based on actual model RAM usage
    max_cpu_workers = int(usable_ram / actual_ram_per_model)
    # OPTIMISED: Allow more workers (was limited to 8)
    safe_cpu_workers = min(config["cpu_workers"], max_cpu_workers, 20)
    
    if safe_cpu_workers != config["cpu_workers"]:
        print(f"🧠 Model '{model_name}' requires {actual_ram_per_model}GB RAM per instance")
        print(f"   OPTIMISED adjustment: {config['cpu_workers']} → {safe_cpu_workers} (using {usable_ram:.1f}GB)")
        
        config["cpu_workers"] = safe_cpu_workers
        config["total_workers"] = config["gpu_workers"] + safe_cpu_workers
        config["actual_ram_per_model"] = actual_ram_per_model
    
    return config


def check_memory_safety(min_gb=1.5):
    """Check if system has enough memory to continue processing."""
    memory = psutil.virtual_memory()
    available_gb = memory.available / (1024**3)
    
    if available_gb < min_gb:
        print(f"🚨 CRITICAL: Only {available_gb:.1f}GB RAM available (minimum {min_gb}GB required)")
        print("   System may become unstable. Consider reducing worker count or using smaller batches.")
        return False
    
    if available_gb < 2.0:
        print(f"⚠️  WARNING: Low memory ({available_gb:.1f}GB available)")
    
    return True


def extract_segments_aggressive(audio_path, segments, temp_dir, config):
    """Ultra-fast parallel segment extraction using all available CPU cores with memory safeguards."""
    print(f"🎵 Extracting {len(segments)} segments with memory-safe batch processing...")
    
    # Memory-aware batch processing to prevent overflow
    total_segments = len(segments)
    batch_size = min(100, max(20, total_segments // 10))  # Adaptive batch size
    memory = psutil.virtual_memory()
    available_gb = memory.available / (1024**3)
    
    # Reduce batch size if memory is low
    if available_gb < 4.0:
        batch_size = min(batch_size, 20)
        print(f"⚠️  Low memory detected ({available_gb:.1f}GB), using smaller batches")
    
    print(f"   Batch size: {batch_size} segments | Total batches: {(total_segments + batch_size - 1) // batch_size}")
    
    # Send initial progress update for extraction phase
    try:
        import sys
        if hasattr(sys.stdout, 'output_queue'):
            print(f"PROGRESS:0.0|Extracting audio segments...|2|")
    except:
        pass
    
    extraction_start = time.time()
    successful_segments = []
    total_processed = 0
    
    # Process segments in batches to prevent memory overflow
    for batch_start in range(0, total_segments, batch_size):
        batch_end = min(batch_start + batch_size, total_segments)
        batch_segments = segments[batch_start:batch_end]
        batch_args = [(i + batch_start, start, end, audio_path, temp_dir) for i, (start, end) in enumerate(batch_segments)]
        
        print(f"   Processing batch {batch_start//batch_size + 1}/{(total_segments + batch_size - 1) // batch_size} ({len(batch_args)} segments)")
        
        # Use conservative parallel processing for this batch
        max_workers = min(3, len(batch_args))  # Limit to 3 workers max per batch
        
        try:
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                future_to_segment = {}
                for args in batch_args:
                    future = executor.submit(extract_single_segment_worker, args)
                    future_to_segment[future] = args[0]
                
                # Collect results for this batch
                for future in as_completed(future_to_segment, timeout=180):  # 3 minute timeout per batch
                    segment_index = future_to_segment[future]
                    try:
                        result = future.result(timeout=60)
                        if result:
                            successful_segments.append(result)
                            total_processed += 1
                            if total_processed % 50 == 0:  # Progress update every 50 segments
                                progress = (total_processed / total_segments) * 100
                                print(f"   Progress: {total_processed}/{total_segments} segments ({progress:.1f}%)")
                    except Exception as e:
                        print(f"   Failed segment {segment_index}: {e}")
                        continue
        
        except Exception as e:
            print(f"   Batch processing failed, falling back to sequential: {e}")
            # Fallback to sequential processing for this batch
            for args in batch_args:
                try:
                    result = extract_single_segment_worker(args)
                    if result:
                        successful_segments.append(result)
                        total_processed += 1
                except Exception as seg_e:
                    print(f"   Failed segment {args[0]}: {seg_e}")
                    continue
        
        # Memory cleanup between batches
        import gc
        gc.collect()
        
        # Check memory usage and pause if needed
        current_memory = psutil.virtual_memory()
        if current_memory.available / (1024**3) < 1.0:  # Less than 1GB available
            print(f"   ⚠️  Low memory ({current_memory.available / (1024**3):.1f}GB), pausing for cleanup...")
            time.sleep(2)  # Give system time to recover
    
    extraction_time = time.time() - extraction_start
    print(f"\n⚡ Extracted {len(successful_segments)}/{total_segments} segments in {extraction_time:.1f}s")
    if len(successful_segments) > 0:
        print(f"   Speed: {len(successful_segments) / extraction_time:.1f} segments/second")
    
    return [(path, start, end) for _, path, start, end in successful_segments]


def get_optimal_worker_counts():
    """Determine optimal worker distribution for maximum hardware utilization with RAM constraints."""
    cpu_cores = multiprocessing.cpu_count()
    has_cuda = torch.cuda.is_available()
    
    # Get system memory info
    memory = psutil.virtual_memory()
    total_ram_gb = memory.total / (1024**3)
    available_ram_gb = memory.available / (1024**3)
    
    print(f"🖥️  System Resources:")
    print(f"   RAM: {available_ram_gb:.1f}GB available / {total_ram_gb:.1f}GB total")
    print(f"   CPU cores: {cpu_cores}")
    
    # Estimate RAM usage per model (conservative estimates)
    model_ram_usage = {
        'tiny': 0.5,      # ~500MB
        'base': 0.8,      # ~800MB  
        'small': 1.5,     # ~1.5GB
        'medium': 2.5,    # ~2.5GB
        'large': 4.0,     # ~4GB
        'large-v2': 4.0,  # ~4GB
        'large-v3': 4.0   # ~4GB
    }
    
    if has_cuda:
        gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
        print(f"   GPU: NVIDIA {torch.cuda.get_device_name(0)} ({gpu_memory:.1f}GB VRAM)")
        
        # Conservative GPU+CPU configuration with RAM awareness
        gpu_workers = 2  # Keep GPU workers low to save VRAM
        
        # Calculate OPTIMISED CPU workers based on available RAM
        # Reserve only 1GB for system + other processes (was 4GB - now truly optimised!)
        usable_ram = max(available_ram_gb - 1.0, 2.0)
        
        # Estimate CPU workers based on RAM per model instance
        # We'll determine model size later, so use medium as baseline
        estimated_ram_per_worker = model_ram_usage.get('medium', 2.5)
        max_cpu_workers_by_ram = int(usable_ram / estimated_ram_per_worker)
        
        # OPTIMISED: Use conservative worker counts to prevent memory overflow
        cpu_workers = min(
            cpu_cores // 2,  # Use half CPU cores (was cpu_cores - safer)
            max_cpu_workers_by_ram,  # RAM constraint
            8  # Lower maximum (was 32) - prevent memory overflow
        )
        
        print(f"🚀 MEMORY-SAFE Hybrid Config:")
        print(f"   GPU Workers: {gpu_workers} (CUDA parallel streams)")
        print(f"   CPU Workers: {cpu_workers} (Conservative - memory safe)")
        print(f"   Total Workers: {gpu_workers + cpu_workers}")
        print(f"   RAM Usage: {usable_ram:.1f}GB available for models")
        
        return {
            "gpu_workers": gpu_workers,
            "cpu_workers": cpu_workers,
            "total_workers": gpu_workers + cpu_workers,
            "segment_extraction_workers": min(cpu_cores // 4, 4),  # Conservative extraction workers
            "ram_constraint": True,
            "available_ram_gb": available_ram_gb,
            "estimated_ram_per_model": estimated_ram_per_worker
        }
    else:
        # OPTIMISED CPU-only mode with maximum RAM utilisation
        usable_ram = max(available_ram_gb - 0.5, 1.0)  # Reserve only 500MB (was 2GB)
        estimated_ram_per_worker = model_ram_usage.get('medium', 2.5)
        max_cpu_workers_by_ram = int(usable_ram / estimated_ram_per_worker)
        
        cpu_workers = min(
            cpu_cores // 2,  # Use half CPU cores (was cpu_cores - safer)
            max_cpu_workers_by_ram,
            6  # Lower maximum (was 24) - prevent memory overflow
        )
        
        print(f"💻 MEMORY-SAFE CPU-Only Config:")
        print(f"   CPU Workers: {cpu_workers} (Conservative - memory safe)")
        print(f"   RAM Usage: {usable_ram:.1f}GB available for models")
        
        return {
            "gpu_workers": 0,
            "cpu_workers": cpu_workers,
            "total_workers": cpu_workers,
            "segment_extraction_workers": min(cpu_cores // 4, 3),  # Conservative extraction workers
            "ram_constraint": True,
            "available_ram_gb": available_ram_gb,
            "estimated_ram_per_model": estimated_ram_per_worker
        }


def monitor_system_usage():
    """Monitor CPU and GPU usage during processing."""
    def monitor_worker():
        while True:
            cpu_percent = psutil.cpu_percent(interval=1)
            memory_percent = psutil.virtual_memory().percent
            
            # GPU monitoring if available
            gpu_info = ""
            if torch.cuda.is_available():
                try:
                    gpu_memory = torch.cuda.memory_allocated() / 1024**3  # GB
                    gpu_memory_max = torch.cuda.max_memory_allocated() / 1024**3  # GB
                    gpu_info = f" | GPU: {gpu_memory:.1f}/{gpu_memory_max:.1f}GB"
                except:
                    gpu_info = " | GPU: Active"
            
            print(f"📊 System Load: CPU {cpu_percent:5.1f}% | RAM {memory_percent:4.1f}%{gpu_info}", end="\r")
            time.sleep(2)
    
    monitor_thread = threading.Thread(target=monitor_worker, daemon=True)
    monitor_thread.start()
    return monitor_thread


def load_models_aggressive(model_name="medium", config=None):
    """Load multiple model instances for maximum parallel processing with RAM monitoring."""
    if config is None:
        config = {"gpu_workers": 0, "cpu_workers": 1}
    
    models = {}
    
    # Safety check: ensure minimum RAM available
    memory = psutil.virtual_memory()
    if memory.available / (1024**3) < 3.0:  # Less than 3GB available
        print(f"⚠️  WARNING: Only {memory.available / (1024**3):.1f}GB RAM available")
        print("   Switching to conservative single-model mode to prevent system crash")
        
        # Emergency fallback: single model only
        try:
            if torch.cuda.is_available():
                model = whisper.load_model(model_name, device="cuda")
                models["gpu_0"] = model
                config["gpu_workers"] = 1
                config["cpu_workers"] = 0
            else:
                model = whisper.load_model(model_name, device="cpu")
                models["cpu_0"] = model
                config["gpu_workers"] = 0
                config["cpu_workers"] = 1
            
            config["total_workers"] = 1
            print(f"🚨 Emergency mode: Single {list(models.keys())[0]} model loaded")
            return models
            
        except Exception as e:
            raise RuntimeError(f"Failed to load even a single model: {e}")
    # Monitor RAM usage during model loading
    initial_memory = psutil.virtual_memory()
    print(f"📊 Initial RAM: {initial_memory.available / (1024**3):.1f}GB available")
    
    if config["gpu_workers"] > 0:
        print(f"🧠 Loading {config['gpu_workers']} GPU model instances...")
        # Load multiple GPU instances for parallel CUDA streams
        for i in range(config["gpu_workers"]):
            try:
                model = whisper.load_model(model_name, device="cuda")
                # Force model to stay on GPU
                model = model.cuda()
                models[f"gpu_{i}"] = model
                
                # Monitor RAM after each model load
                current_memory = psutil.virtual_memory()
                ram_used = (initial_memory.available - current_memory.available) / (1024**3)
                print(f"   GPU Model {i}: Loaded on {next(model.parameters()).device} (RAM used: {ram_used:.1f}GB)")
                
                # Clear cache to prevent accumulation
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                    
            except Exception as e:
                print(f"   GPU Model {i}: Failed - {e}")
    
    if config["cpu_workers"] > 0:
        # Adjust CPU model count based on RAM constraints - MORE AGGRESSIVE
        max_cpu_models = min(
            config["cpu_workers"] // 2,  # Allow more models (was // 4)
            8,  # Higher maximum (was 4)
            int(config.get("available_ram_gb", 8) / config.get("actual_ram_per_model", 2.5))  # RAM constraint
        )
        
        print(f"🖥️  Loading {max_cpu_models} CPU model instances (RAM-limited)...")
        
        for i in range(max_cpu_models):
            try:
                current_memory = psutil.virtual_memory()
                if current_memory.available / (1024**3) < 2.0:  # Less than 2GB available
                    print(f"   ⚠️  Skipping CPU Model {i}: Insufficient RAM ({current_memory.available / (1024**3):.1f}GB)")
                    break
                    
                model = whisper.load_model(model_name, device="cpu")
                models[f"cpu_{i}"] = model
                
                # Monitor RAM usage
                new_memory = psutil.virtual_memory()
                ram_used = (current_memory.available - new_memory.available) / (1024**3)
                print(f"   CPU Model {i}: Loaded (RAM used: {ram_used:.1f}GB, {new_memory.available / (1024**3):.1f}GB available)")
                
            except Exception as e:
                print(f"   CPU Model {i}: Failed - {e}")
                break
    
    # Final memory check
    final_memory = psutil.virtual_memory()
    total_ram_used = (initial_memory.available - final_memory.available) / (1024**3)
    print(f"📊 Total models loaded: {len(models)} (Total RAM used: {total_ram_used:.1f}GB)")
    
    return models


def transcribe_segment_aggressive(model, audio_path, segment_info, worker_id):
    """Optimised transcription with performance monitoring."""
    start_time = time.time()
    segment_path, seg_start, seg_end = segment_info
    
    try:
        # Optimised Whisper parameters for better speech detection
        result = model.transcribe(
            segment_path,
            language=None,
            compression_ratio_threshold=2.4,  # Less restrictive content filtering
            logprob_threshold=-2.0,           # More lenient confidence threshold
            no_speech_threshold=0.3,          # Less sensitive speech detection (was 0.05)
            condition_on_previous_text=False, # No context dependency
            temperature=0.0,                  # Deterministic
            beam_size=5,                     # Higher quality beam search
            patience=2.0                     # More thorough processing
        )
        
        text = result.get("text", "").strip()
        duration = time.time() - start_time
        
        return {
            "text": text, 
            "start": seg_start, 
            "end": seg_end,
            "worker_id": worker_id,
            "processing_time": duration
        }
        
    except Exception as e:
        print(f"\n❌ Worker {worker_id} error: {e}")
        return {
            "text": "", 
            "start": seg_start, 
            "end": seg_end,
            "worker_id": worker_id,
            "processing_time": 0
        }


def extract_single_segment_worker(args):
    """Worker function for parallel segment extraction with multi-threading."""
    idx, start_time, end_time, audio_path, temp_dir = args
    output_path = os.path.join(temp_dir, f"seg_{idx:04d}.mp3")

    try:
        # Multi-threaded ffmpeg extraction for maximum CPU utilization
        cpu_cores = multiprocessing.cpu_count()
        ffmpeg_threads = min(cpu_cores // 2, 8)  # Use up to 8 threads for ffmpeg

        cmd = [
            "ffmpeg", "-y", "-v", "quiet",
            "-i", audio_path,
            "-ss", str(start_time),
            "-t", str(end_time - start_time),
            "-acodec", "libmp3lame",
            "-b:a", "128k",  # Higher bitrate for quality
            "-ar", "16000",
            "-ac", "1",
            "-threads", str(ffmpeg_threads),  # Multi-threaded extraction
            output_path
        ]
        subprocess.run(cmd, check=True, timeout=60)
        return (idx, output_path, start_time, end_time)
    except Exception as e:
        print(f"\n❌ Segment {idx} extraction failed: {e}")
        return None


def extract_single_segment_worker_thread(args):
    """Thread-safe worker function for segment extraction fallback."""
    idx, start_time, end_time, audio_path, temp_dir = args
    output_path = os.path.join(temp_dir, f"seg_{idx:04d}.mp3")
    
    try:
        # Conservative ffmpeg extraction for thread safety
        cmd = [
            "ffmpeg", "-y", "-v", "error",  # Show only errors
            "-i", audio_path,
            "-ss", str(start_time),
            "-t", str(end_time - start_time),
            "-acodec", "libmp3lame",
            "-b:a", "64k",   # Even lower bitrate for stability
            "-ar", "16000",
            "-ac", "1",
            "-threads", "1",
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0 and os.path.exists(output_path):
            return (idx, output_path, start_time, end_time)
        else:
            print(f"\n❌ Thread segment {idx} failed: {result.stderr}")
            return None
    except Exception as e:
        print(f"\n❌ Thread segment {idx} exception: {e}")
        return None


def transcribe_parallel_aggressive(models, segment_files, segments, config, model_name="medium"):
    """Maximum parallel transcription using all GPU and CPU resources with progress tracking."""

    # Safety check for None or empty segments
    if segments is None:
        print("❌ Error: segments parameter is None")
        return []

    if len(segments) == 0:
        print("⚠️  Warning: No segments to process")
        return []

    print(f"🗣️  Starting MEMORY-SAFE parallel transcription...")
    print(f"   Processing {len(segments)} segments with {config['total_workers']} workers")
    print(f"   Using ProcessPoolExecutor with memory monitoring")

    # Start system monitoring
    monitor_thread = monitor_system_usage()

    results = []
    segment_queue = multiprocessing.Manager().Queue()  # Use Manager for process-safe queue
    completed_segments = multiprocessing.Value('i', 0)  # Shared counter
    start_time = time.time()

    # Add all segments to processing queue
    for i, segment_file in enumerate(segment_files):
        orig_start, orig_end = segments[i]
        segment_queue.put((i, (segment_file, 0, orig_end - orig_start)))

    def send_progress_update():
        """Send progress update to GUI if running in GUI mode."""
        if completed_segments.value > 0 and len(segments) > 0:
            percentage = (completed_segments.value / len(segments)) * 100
            elapsed = time.time() - start_time
            active_processes = multiprocessing.active_children()

            status = f"Transcribing: {completed_segments.value}/{len(segments)} segments"
            elapsed_str = f"Elapsed: {elapsed:.0f}s"

            # Try to send to GUI queue if available
            try:
                # This will only work if called from GUI context
                import sys
                if hasattr(sys.stdout, 'output_queue'):
                    print(f"PROGRESS:{percentage:.1f}|{status}|{len(active_processes)}|{elapsed_str}")
            except:
                pass  # Not in GUI context, just continue

            # Also print to console
            print(f"📊 Progress: {percentage:.1f}% ({completed_segments.value}/{len(segments)}) - {len(active_processes)} processes active")

def gpu_worker_process(args):
    """GPU worker process for high-throughput transcription."""
    worker_id, model_key, queue, completed_counter, model_name = args

    # Load model in this process
    try:
        import torch
        import whisper
        model = whisper.load_model(model_name, device="cuda")  # Use actual model name
        print(f"🔄 GPU Worker {worker_id}: Model loaded on {next(model.parameters()).device}")
    except Exception as e:
        print(f"❌ GPU Worker {worker_id}: Failed to load model: {e}")
        return []
        worker_results = []
        processed_count = 0

        while True:
            try:
                seg_id, segment_info = queue.get_nowait()

                result = transcribe_segment_aggressive(model, None, segment_info, f"GPU-{worker_id}")
                result["segment_id"] = seg_id
                worker_results.append(result)
                processed_count += 1

                # Update progress
                with completed_counter:
                    completed_counter.value += 1
                    if completed_counter.value % 5 == 0:  # Update every 5 segments
                        send_progress_update()

                # Clear GPU cache periodically
                if processed_count % 5 == 0:
                    torch.cuda.empty_cache()
                    import gc
                    gc.collect()

            except Exception as e:
                # Check if it's an empty queue exception
                if "empty" in str(e).lower():
                    break
                else:
                    print(f"\n❌ GPU Worker {worker_id} error: {e}")

        return worker_results

def cpu_worker_process(args):
    """CPU worker process for parallel processing with multi-core support."""
    worker_id, model_key, queue, completed_counter, cores_per_worker, model_name = args

    # Load model in this process
    try:
        import torch
        import whisper
        model = whisper.load_model(model_name, device="cpu")  # Use actual model name
        print(f"🔄 CPU Worker {worker_id}: Model loaded with {cores_per_worker} cores available")
    except Exception as e:
        print(f"❌ CPU Worker {worker_id}: Failed to load model: {e}")
        return []
        worker_results = []
        processed_count = 0

        while True:
            try:
                seg_id, segment_info = queue.get_nowait()

                result = transcribe_segment_aggressive(model, None, segment_info, f"CPU-{worker_id}")
                result["segment_id"] = seg_id
                worker_results.append(result)
                processed_count += 1

                # Update progress
                with completed_counter:
                    completed_counter.value += 1
                    if completed_counter.value % 3 == 0:  # Update every 3 segments for CPU
                        send_progress_update()

            except Exception as e:
                # Check if it's an empty queue exception
                if "empty" in str(e).lower():
                    break
                else:
                    print(f"\n❌ CPU Worker {worker_id} error: {e}")

        return worker_results

    # Start all workers using ProcessPoolExecutor for true parallelism
    futures = []
    transcription_start = time.time()

    # Calculate cores per worker for optimal utilization
    cpu_cores = multiprocessing.cpu_count()
    gpu_workers = len([k for k in models.keys() if k.startswith('gpu')])
    cpu_workers = config["cpu_workers"]
    cores_per_cpu_worker = 1  # Default value

    # Distribute CPU cores optimally
    if cpu_workers > 0:
        cores_per_cpu_worker = max(1, cpu_cores // cpu_workers)
        print(f"⚡ CPU Distribution: {cpu_workers} workers × {cores_per_cpu_worker} cores each = MEMORY-SAFE UTILIZATION")

    with multiprocessing.Pool(processes=config["total_workers"]) as pool:
        # Start GPU workers
        gpu_model_keys = [k for k in models.keys() if k.startswith('gpu')]
        for i, model_key in enumerate(gpu_model_keys):
            args = (i, model_key, segment_queue, completed_segments, model_name)
            future = pool.apply_async(gpu_worker_process, (args,))
            futures.append(future)

        # Start CPU workers
        cpu_model_keys = [k for k in models.keys() if k.startswith('cpu')]
        if cpu_model_keys:
            for i in range(config["cpu_workers"]):
                model_idx = i % len(cpu_model_keys)
                model_key = cpu_model_keys[model_idx]
                cores_per_worker = max(1, cpu_cores // cpu_workers)
                args = (i, model_key, segment_queue, completed_segments, cores_per_worker, model_name)
                future = pool.apply_async(cpu_worker_process, (args,))
                futures.append(future)
        else:
            print("⚠️  No CPU models available, skipping CPU workers")

        # Collect all results
        all_results = []
        completed_workers = 0
        for future in futures:
            try:
                worker_results = future.get(timeout=3600)  # 1 hour timeout
                if worker_results:
                    all_results.extend(worker_results)
                completed_workers += 1
                print(f"\n✅ Worker {completed_workers}/{len(futures)} completed ({len(worker_results)} segments)")
            except Exception as e:
                print(f"\n❌ Worker completion error: {e}")

    transcription_time = time.time() - transcription_start

    # Sort results by segment ID
    all_results.sort(key=lambda x: x.get('segment_id', 0))

    # Send final progress update
    try:
        import sys
        if hasattr(sys.stdout, 'output_queue'):
            elapsed = time.time() - start_time
            print(f"PROGRESS:100.0|Transcription Complete!|0|Total: {elapsed:.0f}s")
    except:
        pass

    # Calculate performance stats
    total_audio_duration = sum(r.get('end', 0) - r.get('start', 0) for r in all_results)
    processing_speedup = total_audio_duration / transcription_time if transcription_time > 0 else 0

    print(f"\n🎯 MEMORY-SAFE Transcription Complete!")
    print(f"   Total segments: {len(all_results)}")
    print(f"   Processing time: {transcription_time:.1f}s")
    print(f"   Audio duration: {total_audio_duration:.1f}s")
    print(f"   Speed: {processing_speedup:.1f}x realtime")
    print(f"   CPU utilization: MEMORY-SAFE UTILIZATION")

    return all_results


def transcribe_file_aggressive(input_path, model_name="medium", output_dir=None, force_aggressive=True):
    """
    Ultra-optimised transcription using maximum available hardware resources.
    """
    from transcribe import (vad_segment_times, post_process_segments, 
                          preprocess_audio, get_media_duration, 
                          split_into_paragraphs, format_duration)
    
    start_time = time.time()
    temp_files = []
    
    try:
        print(f"🚀 OPTIMISED GPU+CPU HYBRID TRANSCRIPTION")
        print(f"📁 Input: {os.path.basename(input_path)}")
        
        # System info
        cpu_cores = multiprocessing.cpu_count()
        has_cuda = torch.cuda.is_available()
        gpu_name = torch.cuda.get_device_name(0) if has_cuda else "None"
        
        print(f"🖥️  System: {cpu_cores} CPU cores, GPU: {gpu_name}")
        
        # Get optimal configuration
        config = get_optimal_worker_counts()
        
        # Adjust worker counts based on actual model size to prevent RAM exhaustion
        config = adjust_workers_for_model(config, model_name)
        
        if not output_dir:
            output_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        
        # Create temp directory
        temp_dir = tempfile.mkdtemp(prefix="optimised_transcribe_")
        temp_files.append(temp_dir)
        
        # Get media info
        duration = get_media_duration(input_path)
        if duration:
            print(f"⏱️  Duration: {format_duration(duration)}")
        print()  # Add spacing for readability
        
        # Preprocess audio
        print("🔄 Preprocessing audio...")
        tf = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
        tf.close()
        pre_path = tf.name
        preprocess_audio(input_path, pre_path, bitrate="128k")  # Lower bitrate for speed
        audio_path = pre_path
        temp_files.append(pre_path)
        
        # Load models optimally
        models = load_models_aggressive(model_name, config)
        if not models:
            raise Exception("Failed to load any models!")
        print()  # Add spacing after model loading
        
        # VAD segmentation with LESS aggressive settings for better speech detection
        print("✂️  Running speech-sensitive VAD segmentation...")
        segments = vad_segment_times(audio_path, aggressiveness=1,  # Less aggressive (was 3)
                                   frame_duration_ms=30, padding_ms=200)  # More padding for context
        
        segments = post_process_segments(segments, min_duration=0.5,  # Longer minimum duration
                                       merge_gap=0.5, max_segments=1000)  # More merging to preserve speech
        
        print(f"📊 Final segments: {len(segments)}")
        if duration:
            print(f"   Average: {len(segments)/max(duration/60, 1):.1f} segments per minute)")
        
        # Debug: Show first few segments
        if segments:
            print(f"   Sample segments: {segments[:3]}")
        
        if len(segments) <= 1:
            # Single segment - use GPU if available
            print("🗣️  Single segment transcription...")
            gpu_model = next((v for k, v in models.items() if k.startswith('gpu')), None)
            model = gpu_model or list(models.values())[0]
            
            result = transcribe_segment_aggressive(model, audio_path, (audio_path, 0, duration or 0), "SINGLE")
            full_text = result.get("text", "")
            print(f"   Single segment transcription: '{full_text[:100]}...'")
        else:
            # Extract segments with maximum parallelism
            segment_files = extract_segments_aggressive(audio_path, segments, temp_dir, config)
            
            # Optimised parallel transcription
            results = transcribe_parallel_aggressive(models, segment_files, segments, config, model_name)
            
            # Debug: Check results
            print(f"   Transcription results: {len(results) if results else 0} segments processed")
            if results:
                sample_texts = [r.get("text", "") for r in results[:3] if r.get("text", "").strip()]
                print(f"   Sample transcriptions: {sample_texts}")
            
            # Combine results with safety check
            if results:
                full_text = " ".join([r.get("text", "") for r in results if r.get("text", "").strip()])
            else:
                full_text = ""
        
        # Check if we have meaningful content (not just punctuation or single characters)
        stripped_text = full_text.strip()
        if not stripped_text:
            full_text = "[No speech detected]"
        elif stripped_text in ['.', '!', '?', ',', ';', ':', '-', '_', '(', ')', '[', ']', '{', '}', '"', "'"]:
            full_text = "[No speech detected]"
        elif len(stripped_text.replace('.', '').replace('!', '').replace('?', '').replace(',', '').replace(';', '').replace(':', '').strip()) == 0:
            # If text contains only punctuation, treat as no speech
            full_text = "[No speech detected]"
        
        # Post-processing
        print("📝 Post-processing...")
        try:
            punctuation_model = PunctuationModel()
            full_text = punctuation_model.restore_punctuation(full_text)
        except Exception as e:
            print(f"⚠️  Punctuation restoration failed: {e}")
        
        # Format text
        formatted_text = split_into_paragraphs(full_text, max_length=500)
        if isinstance(formatted_text, list):
            formatted_text = '\n\n'.join(formatted_text)
        
        # Save outputs
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        txt_path = os.path.join(output_dir, f"{base_name}_optimised.txt")
        docx_path = os.path.join(output_dir, f"{base_name}_optimised.docx")
        
        # Save files
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(formatted_text)
        
        # Create Word document
        doc = Document()
        doc.add_heading(f'Optimised Transcription: {base_name}', 0)
        
        if duration:
            doc.add_paragraph(f'Duration: {format_duration(duration)}')
        
        elapsed = time.time() - start_time
        doc.add_paragraph(f'Processing time: {format_duration(elapsed)}')
        doc.add_paragraph(f'Model: {model_name} (Optimised GPU+CPU Hybrid)')
        doc.add_paragraph(f'Hardware: {gpu_name} + {cpu_cores} CPU cores')
        
        if duration and elapsed > 0:
            speedup = duration / elapsed
            doc.add_paragraph(f'Processing speed: {speedup:.1f}x realtime')
        
        doc.add_paragraph('')
        
        for para in formatted_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(docx_path)
        
        # Final stats
        elapsed = time.time() - start_time
        print(f"\n🎉 OPTIMISED TRANSCRIPTION COMPLETE!")
        print(f"📄 Text file: {txt_path}")
        print(f"📄 Word document: {docx_path}")
        print(f"⏱️  Total time: {format_duration(elapsed)}")
        
        if duration:
            speedup = duration / elapsed
            print(f"⚡ Final speed: {speedup:.1f}x realtime")
        
        return txt_path
        
    except Exception as e:
        print(f"❌ Optimised transcription failed: {e}")
        raise
        
    finally:
        # Cleanup
        for temp_file in temp_files:
            try:
                if os.path.isdir(temp_file):
                    shutil.rmtree(temp_file)
                elif os.path.isfile(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"⚠️  Cleanup warning: {e}")


def main():
    parser = argparse.ArgumentParser(description="Optimised GPU+CPU hybrid transcription")
    parser.add_argument("--input", required=True, help="Input audio/video file")
    parser.add_argument("--model", default="medium", choices=["tiny", "base", "small", "medium", "large"])
    parser.add_argument("--output-dir", help="Output directory (default: Downloads)")
    
    args = parser.parse_args()
    
    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        return 1
    
    try:
        transcribe_file_aggressive(
            args.input,
            model_name=args.model,
            output_dir=args.output_dir
        )
        return 0
    except Exception as e:
        print(f"Error: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
